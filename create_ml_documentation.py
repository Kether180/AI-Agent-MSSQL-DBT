"""
DataMigrate AI - ML Strategy & Data Collection Documentation Generator
Creates comprehensive PDF and Word documents for ML strategy and customer contracts

Author: Alexander Garcia Angus
Company: OKO Investments
"""

import os
from datetime import datetime
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
)
from reportlab.lib.units import inch

# Word document imports
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# =============================================================================
# ML STRATEGY DOCUMENTATION - PDF
# =============================================================================

def create_ml_strategy_pdf(output_dir: str = "docs"):
    """Generate comprehensive ML Strategy documentation PDF"""

    os.makedirs(output_dir, exist_ok=True)
    filename = f"{output_dir}/DATAMIGRATE_AI_ML_STRATEGY.pdf"

    doc = SimpleDocTemplate(
        filename,
        pagesize=A4,
        rightMargin=50,
        leftMargin=50,
        topMargin=50,
        bottomMargin=50
    )

    styles = getSampleStyleSheet()

    # Custom styles
    styles.add(ParagraphStyle(
        name='MainTitle',
        parent=styles['Heading1'],
        fontSize=28,
        spaceAfter=20,
        alignment=TA_CENTER,
        textColor=colors.HexColor('#1a365d')
    ))
    styles.add(ParagraphStyle(
        name='Subtitle',
        parent=styles['Heading2'],
        fontSize=16,
        spaceAfter=30,
        alignment=TA_CENTER,
        textColor=colors.HexColor('#4a5568')
    ))
    styles.add(ParagraphStyle(
        name='SectionTitle',
        parent=styles['Heading1'],
        fontSize=18,
        spaceBefore=25,
        spaceAfter=15,
        textColor=colors.HexColor('#2d3748')
    ))
    styles.add(ParagraphStyle(
        name='SubSection',
        parent=styles['Heading3'],
        fontSize=13,
        spaceBefore=15,
        spaceAfter=8,
        textColor=colors.HexColor('#4a5568')
    ))
    styles.add(ParagraphStyle(
        name='BodyCustom',
        parent=styles['Normal'],
        fontSize=10,
        spaceAfter=8,
        leading=14,
        alignment=TA_JUSTIFY
    ))
    styles.add(ParagraphStyle(
        name='BulletText',
        parent=styles['Normal'],
        fontSize=10,
        leftIndent=20,
        spaceAfter=5,
        leading=13
    ))
    styles.add(ParagraphStyle(
        name='CodeStyle',
        fontName='Courier',
        fontSize=8,
        backColor=colors.HexColor('#f7fafc'),
        leftIndent=10,
        spaceAfter=8
    ))

    story = []

    # ==========================================================================
    # COVER PAGE
    # ==========================================================================
    story.append(Spacer(1, 80))
    story.append(Paragraph("DataMigrate AI", styles['MainTitle']))
    story.append(Spacer(1, 10))
    story.append(Paragraph("Machine Learning Strategy", styles['Subtitle']))
    story.append(Spacer(1, 20))
    story.append(Paragraph("Fine-Tuning Open-Source Models for SQL Migration", styles['SubSection']))
    story.append(Paragraph("Technical Documentation & Business Case", styles['SubSection']))
    story.append(Spacer(1, 50))

    company_info = [
        ['Company:', 'OKO Investments'],
        ['Author:', 'Alexander Garcia Angus'],
        ['Date:', datetime.now().strftime('%B %Y')],
        ['Version:', '1.0'],
        ['Classification:', 'Confidential'],
    ]
    company_table = Table(company_info, colWidths=[100, 200])
    company_table.setStyle(TableStyle([
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('TEXTCOLOR', (0, 0), (0, -1), colors.HexColor('#4a5568')),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
    ]))
    story.append(company_table)
    story.append(PageBreak())

    # ==========================================================================
    # TABLE OF CONTENTS
    # ==========================================================================
    story.append(Paragraph("Table of Contents", styles['SectionTitle']))
    toc_items = [
        "1. Executive Summary",
        "2. Why Fine-Tune Open-Source Models?",
        "3. Recommended Model: SQLCoder",
        "4. Alternative Models Comparison",
        "5. Data Collection Strategy",
        "6. Privacy & Anonymization",
        "7. Training Pipeline",
        "8. Business Benefits & ROI",
        "9. Implementation Roadmap",
        "10. Infrastructure Requirements",
        "11. Risk Analysis",
    ]
    for item in toc_items:
        story.append(Paragraph(item, styles['BulletText']))
    story.append(PageBreak())

    # ==========================================================================
    # 1. EXECUTIVE SUMMARY
    # ==========================================================================
    story.append(Paragraph("1. Executive Summary", styles['SectionTitle']))
    story.append(Paragraph(
        "This document outlines DataMigrate AI's strategy for developing proprietary AI models through "
        "fine-tuning open-source language models. By collecting anonymized SQL migration patterns from "
        "customer engagements, we can create specialized models that outperform generic AI APIs while "
        "reducing operational costs by 70% and creating a defensible competitive moat.",
        styles['BodyCustom']
    ))
    story.append(Spacer(1, 10))

    summary_points = [
        "Primary Model: SQLCoder 15B - Best SQL generation accuracy",
        "Training Data: Anonymized schema patterns from customer migrations",
        "Cost Reduction: 70% savings vs API-based approach at scale",
        "Competitive Moat: Proprietary training data cannot be replicated",
        "Timeline: 12-18 months from data collection to production deployment",
    ]
    for point in summary_points:
        story.append(Paragraph(f"* {point}", styles['BulletText']))
    story.append(PageBreak())

    # ==========================================================================
    # 2. WHY FINE-TUNE?
    # ==========================================================================
    story.append(Paragraph("2. Why Fine-Tune Open-Source Models?", styles['SectionTitle']))

    story.append(Paragraph("Current State: API-Based Approach", styles['SubSection']))
    story.append(Paragraph(
        "DataMigrate AI currently uses Claude and GPT APIs for SQL generation. While effective, "
        "this approach has limitations at scale:",
        styles['BodyCustom']
    ))

    current_limitations = [
        "Cost: $0.01-0.10 per API call, scales linearly with usage",
        "Latency: 2-5 seconds per response",
        "Privacy: Customer data sent to third-party servers",
        "No Differentiation: Competitors can use same APIs",
        "No Offline: Requires internet connectivity",
    ]
    for item in current_limitations:
        story.append(Paragraph(f"- {item}", styles['BulletText']))
    story.append(Spacer(1, 15))

    story.append(Paragraph("Future State: Fine-Tuned Models", styles['SubSection']))
    future_benefits = [
        "Cost: ~$0.001 per query (self-hosted), 70% reduction",
        "Latency: 200-500ms per response (5-10x faster)",
        "Privacy: All processing on-premise, no data leaves",
        "Differentiation: Proprietary model trained on our data",
        "Offline: Full functionality without internet",
        "Specialization: Model learns MSSQL-to-dbt patterns specifically",
    ]
    for item in future_benefits:
        story.append(Paragraph(f"+ {item}", styles['BulletText']))
    story.append(PageBreak())

    # ==========================================================================
    # 3. RECOMMENDED MODEL: SQLCODER
    # ==========================================================================
    story.append(Paragraph("3. Recommended Model: SQLCoder", styles['SectionTitle']))

    story.append(Paragraph(
        "After evaluating multiple open-source models, we recommend SQLCoder as the primary base model "
        "for fine-tuning. SQLCoder is specifically designed for SQL generation tasks and outperforms "
        "general-purpose models on SQL benchmarks.",
        styles['BodyCustom']
    ))
    story.append(Spacer(1, 10))

    story.append(Paragraph("Why SQLCoder?", styles['SubSection']))

    sqlcoder_data = [
        ['Criteria', 'SQLCoder 15B', 'GPT-4', 'Claude 3'],
        ['SQL Accuracy (Spider)', '85%', '82%', '80%'],
        ['MSSQL Support', 'Excellent', 'Good', 'Good'],
        ['Fine-tuning Support', 'Full (open weights)', 'None', 'None'],
        ['Self-hosting', 'Yes', 'No', 'No'],
        ['Cost per Query', '~$0.001', '$0.03', '$0.015'],
        ['Latency', '300ms', '3s', '2s'],
        ['Offline Capable', 'Yes', 'No', 'No'],
    ]
    sqlcoder_table = Table(sqlcoder_data, colWidths=[110, 100, 80, 80])
    sqlcoder_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2d3748')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BACKGROUND', (1, 1), (1, -1), colors.HexColor('#c6f6d5')),
    ]))
    story.append(sqlcoder_table)
    story.append(Spacer(1, 15))

    story.append(Paragraph("SQLCoder Technical Specifications", styles['SubSection']))
    specs = [
        "Model Size: 15B parameters (also available in 7B)",
        "Architecture: Transformer (based on StarCoder)",
        "Context Window: 8,192 tokens",
        "Training: SQL-specific datasets including Spider, WikiSQL",
        "License: Apache 2.0 (commercial use allowed)",
        "Fine-tuning: Supports LoRA, QLoRA for efficient training",
        "Inference: Compatible with Ollama, vLLM, HuggingFace",
    ]
    for spec in specs:
        story.append(Paragraph(f"- {spec}", styles['BulletText']))
    story.append(PageBreak())

    # ==========================================================================
    # 4. ALTERNATIVE MODELS
    # ==========================================================================
    story.append(Paragraph("4. Alternative Models Comparison", styles['SectionTitle']))

    models_data = [
        ['Model', 'Size', 'Specialization', 'Pros', 'Cons'],
        ['SQLCoder', '7B-15B', 'SQL Generation', 'Best SQL accuracy', 'SQL-only'],
        ['CodeLlama', '7B-34B', 'General Code', 'Versatile, Meta backing', 'Less SQL-specific'],
        ['Llama 3', '8B-70B', 'General Purpose', 'Strong reasoning', 'Needs more fine-tuning'],
        ['Mistral', '7B', 'General Purpose', 'Fast, efficient', 'Smaller context'],
        ['DeepSeek-Coder', '6.7B-33B', 'Code', 'Good SQL support', 'Less community'],
        ['StarCoder2', '3B-15B', 'Code', 'GitHub trained', 'Less SQL focus'],
    ]
    models_table = Table(models_data, colWidths=[70, 50, 80, 100, 100])
    models_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2d3748')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 8),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f7fafc')]),
    ]))
    story.append(models_table)
    story.append(Spacer(1, 15))

    story.append(Paragraph("Recommendation by Use Case", styles['SubSection']))
    recommendations = [
        "Primary (SQL Generation): SQLCoder 15B - Best accuracy for SQL tasks",
        "Backup (Complex Logic): CodeLlama 34B - Better reasoning for stored procedures",
        "On-Premise (Resource Limited): SQLCoder 7B - Lower hardware requirements",
        "Future (Advanced): Llama 3 70B - When we have more training data",
    ]
    for rec in recommendations:
        story.append(Paragraph(f"- {rec}", styles['BulletText']))
    story.append(PageBreak())

    # ==========================================================================
    # 5. DATA COLLECTION STRATEGY
    # ==========================================================================
    story.append(Paragraph("5. Data Collection Strategy", styles['SectionTitle']))

    story.append(Paragraph(
        "The key to successful fine-tuning is high-quality training data. We will collect SQL migration "
        "patterns from customer engagements, with strict privacy controls.",
        styles['BodyCustom']
    ))
    story.append(Spacer(1, 10))

    story.append(Paragraph("What We Collect (Schema Only)", styles['SubSection']))
    collect_items = [
        "SQL stored procedure structures and logic patterns",
        "Table and column definitions (CREATE TABLE statements)",
        "JOIN patterns and relationships",
        "Data type mappings (MSSQL to dbt)",
        "Transformation logic (aggregations, filters, CTEs)",
        "Successfully generated dbt model outputs",
    ]
    for item in collect_items:
        story.append(Paragraph(f"[Collect] {item}", styles['BulletText']))
    story.append(Spacer(1, 10))

    story.append(Paragraph("What We NEVER Collect", styles['SubSection']))
    never_collect = [
        "Actual row-level data (customer records, transactions)",
        "Personal identifiable information (names, SSNs, emails)",
        "Financial data (account numbers, balances)",
        "Business-sensitive values (prices, salaries)",
        "Authentication credentials",
    ]
    for item in never_collect:
        story.append(Paragraph(f"[NEVER] {item}", styles['BulletText']))
    story.append(Spacer(1, 15))

    story.append(Paragraph("Training Data Quality Requirements", styles['SubSection']))
    quality_data = [
        ['Criteria', 'Requirement', 'Validation Method'],
        ['Success Score', '> 0.8 (80%)', 'Automated testing'],
        ['Completeness', 'Full input/output pair', 'Schema validation'],
        ['Diversity', 'Multiple SQL patterns', 'Pattern clustering'],
        ['Accuracy', 'Validated dbt output', 'dbt compile test'],
    ]
    quality_table = Table(quality_data, colWidths=[100, 150, 150])
    quality_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2d3748')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
    ]))
    story.append(quality_table)
    story.append(PageBreak())

    # ==========================================================================
    # 6. PRIVACY & ANONYMIZATION
    # ==========================================================================
    story.append(Paragraph("6. Privacy & Anonymization", styles['SectionTitle']))

    story.append(Paragraph(
        "All collected data undergoes mandatory anonymization before storage. This ensures customer "
        "privacy while retaining the SQL patterns needed for training.",
        styles['BodyCustom']
    ))
    story.append(Spacer(1, 10))

    story.append(Paragraph("Anonymization Process", styles['SubSection']))

    anon_example = """
BEFORE ANONYMIZATION (Customer SQL):
CREATE PROCEDURE dbo.GetCustomerCreditScore
    @CustomerSSN VARCHAR(11)
AS
SELECT customer_id, first_name, credit_score
FROM dbo.CustomerMaster
WHERE social_security_number = @CustomerSSN

AFTER ANONYMIZATION (Training Data):
CREATE PROCEDURE dbo.Proc_001
    @Param_A VARCHAR(11)
AS
SELECT col_001, col_002, col_003
FROM dbo.Table_A
WHERE col_004 = @Param_A
"""
    for line in anon_example.strip().split('\n'):
        story.append(Paragraph(line, styles['CodeStyle']))
    story.append(Spacer(1, 15))

    story.append(Paragraph("Three Levels of Data Collection", styles['SubSection']))
    levels_data = [
        ['Level', 'Description', 'Privacy', 'Model Quality'],
        ['Anonymized (Default)', 'All names replaced with generic tokens', 'Maximum', 'Good'],
        ['Named (Opt-in)', 'Table/column names preserved', 'High', 'Better'],
        ['No Collection', 'Customer opts out entirely', 'N/A', 'N/A'],
    ]
    levels_table = Table(levels_data, colWidths=[100, 180, 70, 70])
    levels_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2d3748')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
    ]))
    story.append(levels_table)
    story.append(PageBreak())

    # ==========================================================================
    # 7. TRAINING PIPELINE
    # ==========================================================================
    story.append(Paragraph("7. Training Pipeline", styles['SectionTitle']))

    story.append(Paragraph("End-to-End Process", styles['SubSection']))
    pipeline_steps = [
        "1. DATA COLLECTION: Capture successful migrations during customer engagements",
        "2. ANONYMIZATION: Remove all identifying information from SQL patterns",
        "3. QUALITY FILTER: Only keep examples with success score > 0.8",
        "4. FORMAT CONVERSION: Convert to JSONL training format",
        "5. BASE MODEL: Load SQLCoder 15B as starting point",
        "6. FINE-TUNING: Apply LoRA/QLoRA for efficient training",
        "7. EVALUATION: Test on held-out migration examples",
        "8. DEPLOYMENT: Deploy via Ollama or vLLM",
        "9. MONITORING: Track accuracy and collect feedback",
        "10. ITERATION: Retrain with new data quarterly",
    ]
    for step in pipeline_steps:
        story.append(Paragraph(step, styles['BulletText']))
    story.append(Spacer(1, 15))

    story.append(Paragraph("Training Data Format (JSONL)", styles['SubSection']))
    jsonl_example = """
{
  "messages": [
    {
      "role": "system",
      "content": "You are a SQL migration expert..."
    },
    {
      "role": "user",
      "content": "Convert MSSQL procedure: CREATE PROCEDURE..."
    },
    {
      "role": "assistant",
      "content": "SELECT col_001 FROM {{ ref('stg_table_a') }}..."
    }
  ]
}
"""
    for line in jsonl_example.strip().split('\n'):
        story.append(Paragraph(line, styles['CodeStyle']))
    story.append(PageBreak())

    # ==========================================================================
    # 8. BUSINESS BENEFITS & ROI
    # ==========================================================================
    story.append(Paragraph("8. Business Benefits & ROI", styles['SectionTitle']))

    story.append(Paragraph("Cost Comparison (Annual, 200 Migrations)", styles['SubSection']))
    cost_data = [
        ['Cost Category', 'API-Based', 'Fine-Tuned', 'Savings'],
        ['AI Inference', '36,000 DKK', '3,600 DKK', '32,400 DKK'],
        ['GPU Infrastructure', '0 DKK', '12,000 DKK', '-12,000 DKK'],
        ['ML Engineer (Part-time)', '0 DKK', '50,000 DKK', '-50,000 DKK'],
        ['Total Annual Cost', '36,000 DKK', '65,600 DKK', '-29,600 DKK'],
        ['Year 3 (500 migrations)', '90,000 DKK', '70,000 DKK', '20,000 DKK'],
    ]
    cost_table = Table(cost_data, colWidths=[130, 90, 90, 90])
    cost_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2d3748')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BACKGROUND', (0, 5), (-1, 5), colors.HexColor('#c6f6d5')),
    ]))
    story.append(cost_table)
    story.append(Spacer(1, 15))

    story.append(Paragraph("Strategic Benefits", styles['SubSection']))
    strategic = [
        "COMPETITIVE MOAT: Proprietary model trained on real migrations - cannot be replicated",
        "ON-PREMISE SALES: Enable enterprise deals requiring data isolation (500K-2M DKK)",
        "SPEED: 5-10x faster responses improve user experience",
        "SCALABILITY: Fixed infrastructure cost regardless of volume",
        "IP OWNERSHIP: Model becomes company asset, increases valuation",
    ]
    for item in strategic:
        story.append(Paragraph(f"* {item}", styles['BulletText']))
    story.append(PageBreak())

    # ==========================================================================
    # 9. IMPLEMENTATION ROADMAP
    # ==========================================================================
    story.append(Paragraph("9. Implementation Roadmap", styles['SectionTitle']))

    roadmap_data = [
        ['Phase', 'Timeline', 'Milestone', 'Investment'],
        ['Data Collection', 'Months 1-12', '500+ training examples', 'Built into product'],
        ['ML Hire', 'Month 6', 'Contract ML engineer', '50,000 DKK'],
        ['First Training', 'Month 9', 'v0.1 model on cloud GPU', '5,000 DKK'],
        ['Evaluation', 'Month 10', 'Benchmark vs Claude/GPT', 'Internal'],
        ['Production', 'Month 12', 'Deploy for simple tasks', '10,000 DKK'],
        ['Scale', 'Month 18', '80% tasks on fine-tuned model', 'GPU server'],
    ]
    roadmap_table = Table(roadmap_data, colWidths=[90, 80, 150, 90])
    roadmap_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2d3748')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
    ]))
    story.append(roadmap_table)
    story.append(PageBreak())

    # ==========================================================================
    # 10. INFRASTRUCTURE REQUIREMENTS
    # ==========================================================================
    story.append(Paragraph("10. Infrastructure Requirements", styles['SectionTitle']))

    story.append(Paragraph("Training Infrastructure (Cloud)", styles['SubSection']))
    training_infra = [
        "GPU: NVIDIA A100 80GB (cloud rental)",
        "Cost: ~$2-4/hour, training takes 4-8 hours",
        "Total per training run: ~$30-50",
        "Frequency: Quarterly retraining",
        "Provider: AWS, Lambda Labs, or RunPod",
    ]
    for item in training_infra:
        story.append(Paragraph(f"- {item}", styles['BulletText']))
    story.append(Spacer(1, 10))

    story.append(Paragraph("Inference Infrastructure (Production)", styles['SubSection']))
    inference_options = [
        ['Option', 'Hardware', 'Cost', 'Use Case'],
        ['Cloud GPU', 'A10G on AWS', '~500 DKK/month', 'Testing, low volume'],
        ['Dedicated Server', 'RTX 4090 (24GB)', '~35,000 DKK one-time', 'Production'],
        ['Enterprise', '2x A10 (48GB)', '~100,000 DKK', 'High availability'],
    ]
    inference_table = Table(inference_options, colWidths=[90, 120, 100, 110])
    inference_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2d3748')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
    ]))
    story.append(inference_table)
    story.append(PageBreak())

    # ==========================================================================
    # 11. RISK ANALYSIS
    # ==========================================================================
    story.append(Paragraph("11. Risk Analysis", styles['SectionTitle']))

    risks_data = [
        ['Risk', 'Likelihood', 'Impact', 'Mitigation'],
        ['Insufficient training data', 'Medium', 'High', 'Aggressive customer acquisition'],
        ['Model underperforms', 'Low', 'Medium', 'Keep Claude API as fallback'],
        ['Privacy breach', 'Low', 'Critical', 'Strict anonymization, audits'],
        ['GPU costs increase', 'Low', 'Low', 'Long-term cloud contracts'],
        ['Competitor catches up', 'Medium', 'Medium', 'Continuous data collection'],
    ]
    risks_table = Table(risks_data, colWidths=[120, 70, 60, 170])
    risks_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2d3748')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
    ]))
    story.append(risks_table)

    # Footer
    story.append(Spacer(1, 50))
    story.append(Paragraph(
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')} | OKO Investments | Confidential",
        ParagraphStyle(name='Footer', fontSize=9, textColor=colors.grey, alignment=TA_CENTER)
    ))

    doc.build(story)
    print(f"[OK] Created: {filename}")
    return filename


# =============================================================================
# ML STRATEGY DOCUMENTATION - WORD
# =============================================================================

def create_ml_strategy_word(output_dir: str = "docs"):
    """Generate comprehensive ML Strategy documentation Word"""

    os.makedirs(output_dir, exist_ok=True)
    filename = f"{output_dir}/DATAMIGRATE_AI_ML_STRATEGY.docx"

    doc = Document()

    # Title
    title = doc.add_heading('DataMigrate AI', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('Machine Learning Strategy')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('Fine-Tuning Open-Source Models for SQL Migration').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # Company info
    info_table = doc.add_table(rows=5, cols=2)
    info_table.cell(0, 0).text = 'Company:'
    info_table.cell(0, 1).text = 'OKO Investments'
    info_table.cell(1, 0).text = 'Author:'
    info_table.cell(1, 1).text = 'Alexander Garcia Angus'
    info_table.cell(2, 0).text = 'Date:'
    info_table.cell(2, 1).text = datetime.now().strftime('%B %Y')
    info_table.cell(3, 0).text = 'Version:'
    info_table.cell(3, 1).text = '1.0'
    info_table.cell(4, 0).text = 'Classification:'
    info_table.cell(4, 1).text = 'Confidential'

    doc.add_page_break()

    # Executive Summary
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        'This document outlines DataMigrate AI\'s strategy for developing proprietary AI models through '
        'fine-tuning open-source language models. By collecting anonymized SQL migration patterns from '
        'customer engagements, we can create specialized models that outperform generic AI APIs while '
        'reducing operational costs by 70% and creating a defensible competitive moat.'
    )

    summary_points = [
        'Primary Model: SQLCoder 15B - Best SQL generation accuracy',
        'Training Data: Anonymized schema patterns from customer migrations',
        'Cost Reduction: 70% savings vs API-based approach at scale',
        'Competitive Moat: Proprietary training data cannot be replicated',
        'Timeline: 12-18 months from data collection to production deployment',
    ]
    for point in summary_points:
        doc.add_paragraph(point, style='List Bullet')

    doc.add_page_break()

    # Recommended Model
    doc.add_heading('2. Recommended Model: SQLCoder', level=1)
    doc.add_paragraph(
        'After evaluating multiple open-source models, we recommend SQLCoder as the primary base model '
        'for fine-tuning. SQLCoder is specifically designed for SQL generation tasks.'
    )

    doc.add_heading('Why SQLCoder?', level=2)
    sqlcoder_benefits = [
        '85% accuracy on SQL benchmarks (vs 82% GPT-4)',
        'Open weights - full fine-tuning support',
        'Self-hosting capable - no API dependency',
        '~$0.001 per query vs $0.03 for GPT-4',
        '300ms latency vs 3s for cloud APIs',
        'Apache 2.0 license - commercial use allowed',
    ]
    for benefit in sqlcoder_benefits:
        doc.add_paragraph(benefit, style='List Bullet')

    doc.add_heading('Technical Specifications', level=2)
    specs = [
        'Model Size: 15B parameters (also 7B available)',
        'Architecture: Transformer (based on StarCoder)',
        'Context Window: 8,192 tokens',
        'Fine-tuning: Supports LoRA, QLoRA',
        'Inference: Ollama, vLLM, HuggingFace compatible',
    ]
    for spec in specs:
        doc.add_paragraph(spec, style='List Bullet')

    doc.add_page_break()

    # Data Collection
    doc.add_heading('3. Data Collection Strategy', level=1)

    doc.add_heading('What We Collect (Schema Only)', level=2)
    collect = [
        'SQL stored procedure structures and logic patterns',
        'Table and column definitions',
        'JOIN patterns and relationships',
        'Data type mappings (MSSQL to dbt)',
        'Successfully generated dbt model outputs',
    ]
    for item in collect:
        doc.add_paragraph(f'[COLLECT] {item}', style='List Bullet')

    doc.add_heading('What We NEVER Collect', level=2)
    never = [
        'Actual row-level data (customer records)',
        'Personal identifiable information (names, SSNs)',
        'Financial data (account numbers, balances)',
        'Authentication credentials',
    ]
    for item in never:
        doc.add_paragraph(f'[NEVER] {item}', style='List Bullet')

    doc.add_page_break()

    # Privacy
    doc.add_heading('4. Privacy & Anonymization', level=1)
    doc.add_paragraph(
        'All collected data undergoes mandatory anonymization before storage.'
    )

    doc.add_heading('Anonymization Example', level=2)
    doc.add_paragraph('BEFORE: CustomerMaster, credit_score, social_security_number')
    doc.add_paragraph('AFTER: Table_A, col_001, col_002')

    doc.add_heading('Three Collection Levels', level=2)
    levels = [
        'Anonymized (Default): All names replaced - Maximum privacy',
        'Named (Opt-in): Table names preserved - Better model quality, 15% discount',
        'No Collection: Customer opts out - Full privacy',
    ]
    for level in levels:
        doc.add_paragraph(level, style='List Bullet')

    doc.add_page_break()

    # Business Benefits
    doc.add_heading('5. Business Benefits & ROI', level=1)

    doc.add_heading('Strategic Benefits', level=2)
    benefits = [
        'COMPETITIVE MOAT: Proprietary model cannot be replicated by competitors',
        'ON-PREMISE SALES: Enable enterprise deals (500K-2M DKK) requiring data isolation',
        'SPEED: 5-10x faster responses improve user experience',
        'SCALABILITY: Fixed infrastructure cost regardless of volume',
        'IP OWNERSHIP: Model becomes company asset, increases valuation',
    ]
    for benefit in benefits:
        doc.add_paragraph(benefit, style='List Bullet')

    doc.add_heading('Cost Projection', level=2)
    doc.add_paragraph('Year 1-2: Investment phase (ML engineer, GPU infrastructure)')
    doc.add_paragraph('Year 3+: 70% cost reduction vs API-based approach')
    doc.add_paragraph('Break-even: ~500 migrations')

    doc.add_page_break()

    # Implementation Roadmap
    doc.add_heading('6. Implementation Roadmap', level=1)

    roadmap = [
        'Months 1-12: Data Collection (500+ training examples)',
        'Month 6: Hire contract ML engineer',
        'Month 9: First model training on cloud GPU',
        'Month 10: Benchmark evaluation vs Claude/GPT',
        'Month 12: Production deployment for simple tasks',
        'Month 18: 80% of tasks on fine-tuned model',
    ]
    for step in roadmap:
        doc.add_paragraph(step, style='List Bullet')

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d")} | OKO Investments | Confidential')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(filename)
    print(f"[OK] Created: {filename}")
    return filename


# =============================================================================
# DATA COLLECTION CONTRACT - PDF
# =============================================================================

def create_data_contract_pdf(output_dir: str = "docs"):
    """Generate Data Collection Agreement PDF"""

    os.makedirs(output_dir, exist_ok=True)
    filename = f"{output_dir}/DATA_COLLECTION_AGREEMENT_TEMPLATE.pdf"

    doc = SimpleDocTemplate(
        filename,
        pagesize=A4,
        rightMargin=50,
        leftMargin=50,
        topMargin=50,
        bottomMargin=50
    )

    styles = getSampleStyleSheet()

    styles.add(ParagraphStyle(
        name='ContractTitle',
        parent=styles['Heading1'],
        fontSize=20,
        spaceAfter=20,
        alignment=TA_CENTER,
        textColor=colors.HexColor('#1a365d')
    ))
    styles.add(ParagraphStyle(
        name='SectionTitle',
        parent=styles['Heading2'],
        fontSize=14,
        spaceBefore=20,
        spaceAfter=10,
        textColor=colors.HexColor('#2d3748')
    ))
    styles.add(ParagraphStyle(
        name='ContractBody',
        parent=styles['Normal'],
        fontSize=10,
        spaceAfter=8,
        leading=14,
        alignment=TA_JUSTIFY
    ))
    styles.add(ParagraphStyle(
        name='ContractBullet',
        parent=styles['Normal'],
        fontSize=10,
        leftIndent=20,
        spaceAfter=5,
        leading=13
    ))

    story = []

    # Header
    story.append(Paragraph("DATA COLLECTION AND AI IMPROVEMENT AGREEMENT", styles['ContractTitle']))
    story.append(Spacer(1, 20))

    story.append(Paragraph(
        "This Data Collection and AI Improvement Agreement (\"Agreement\") is entered into as of "
        "[DATE] by and between:",
        styles['ContractBody']
    ))
    story.append(Spacer(1, 10))

    story.append(Paragraph(
        "<b>OKO Investments</b> (\"DataMigrate AI\", \"Provider\", \"We\", \"Us\")",
        styles['ContractBody']
    ))
    story.append(Paragraph("Address: [Provider Address]", styles['ContractBody']))
    story.append(Paragraph("CVR: [Company Registration Number]", styles['ContractBody']))
    story.append(Spacer(1, 10))

    story.append(Paragraph("AND", styles['ContractBody']))
    story.append(Spacer(1, 10))

    story.append(Paragraph(
        "<b>[CUSTOMER NAME]</b> (\"Customer\", \"Client\", \"You\")",
        styles['ContractBody']
    ))
    story.append(Paragraph("Address: [Customer Address]", styles['ContractBody']))
    story.append(Paragraph("CVR: [Customer Registration Number]", styles['ContractBody']))
    story.append(PageBreak())

    # Section 1: Definitions
    story.append(Paragraph("1. DEFINITIONS", styles['SectionTitle']))

    definitions = [
        ("\"Schema Data\"", "Database structure information including table names, column names, data types, "
         "stored procedure logic, and SQL transformation patterns. Schema Data does NOT include actual "
         "row-level data or business content."),
        ("\"Anonymized Data\"", "Schema Data that has been processed to replace all customer-specific "
         "identifiers (table names, column names, procedure names) with generic tokens (Table_A, col_001, etc.)."),
        ("\"Named Data\"", "Schema Data where original table and column names are preserved, but no "
         "actual business data values are included."),
        ("\"Training Data\"", "Anonymized or Named Data used to improve DataMigrate AI's machine learning models."),
        ("\"Row-Level Data\"", "Actual data values stored in database tables. This is NEVER collected."),
    ]

    for term, definition in definitions:
        story.append(Paragraph(f"<b>{term}</b>: {definition}", styles['ContractBody']))
        story.append(Spacer(1, 5))
    story.append(PageBreak())

    # Section 2: Scope
    story.append(Paragraph("2. SCOPE OF DATA COLLECTION", styles['SectionTitle']))

    story.append(Paragraph("2.1 What We Collect", styles['ContractBody']))
    collect_items = [
        "SQL stored procedure structures and transformation logic",
        "CREATE TABLE statements and column definitions",
        "JOIN patterns and table relationships",
        "Data type mappings between MSSQL and dbt",
        "Successfully generated dbt model SQL code",
        "Migration success/failure metadata",
    ]
    for item in collect_items:
        story.append(Paragraph(f"- {item}", styles['ContractBullet']))
    story.append(Spacer(1, 10))

    story.append(Paragraph("2.2 What We NEVER Collect", styles['ContractBody']))
    never_items = [
        "Actual row-level data (customer records, transactions, etc.)",
        "Personal identifiable information (names, addresses, SSNs, emails)",
        "Financial data (account numbers, balances, credit card numbers)",
        "Health information (medical records, diagnoses)",
        "Authentication credentials (passwords, API keys)",
        "Any data classified as sensitive under GDPR Article 9",
    ]
    for item in never_items:
        story.append(Paragraph(f"- {item}", styles['ContractBullet']))
    story.append(PageBreak())

    # Section 3: Collection Options
    story.append(Paragraph("3. DATA COLLECTION OPTIONS", styles['SectionTitle']))

    story.append(Paragraph(
        "Customer may select one of the following data collection options:",
        styles['ContractBody']
    ))
    story.append(Spacer(1, 10))

    story.append(Paragraph("<b>OPTION A: ANONYMIZED COLLECTION (Default)</b>", styles['ContractBody']))
    story.append(Paragraph(
        "All Schema Data is automatically anonymized before storage. Table names, column names, and "
        "procedure names are replaced with generic identifiers. This option provides maximum privacy "
        "while still contributing to model improvement.",
        styles['ContractBody']
    ))
    story.append(Paragraph("Example: 'CustomerMaster.credit_score' becomes 'Table_A.col_001'", styles['ContractBullet']))
    story.append(Spacer(1, 10))

    story.append(Paragraph("<b>OPTION B: NAMED COLLECTION (Opt-in)</b>", styles['ContractBody']))
    story.append(Paragraph(
        "Table and column names are preserved in collected data. This improves model quality for "
        "industry-specific patterns. Customer receives a 15% discount on services in exchange for "
        "this contribution. No actual data values are ever collected.",
        styles['ContractBody']
    ))
    story.append(Paragraph("Benefit: Priority access to industry-specific model improvements", styles['ContractBullet']))
    story.append(Spacer(1, 10))

    story.append(Paragraph("<b>OPTION C: NO COLLECTION (Opt-out)</b>", styles['ContractBody']))
    story.append(Paragraph(
        "Customer may opt out of all data collection. No Schema Data will be retained after migration "
        "completion. Service pricing remains unchanged.",
        styles['ContractBody']
    ))
    story.append(PageBreak())

    # Section 4: Customer Selection
    story.append(Paragraph("4. CUSTOMER SELECTION", styles['SectionTitle']))
    story.append(Paragraph(
        "Customer hereby selects the following data collection option (check one):",
        styles['ContractBody']
    ))
    story.append(Spacer(1, 10))

    options_table = [
        ['[ ]', 'OPTION A: Anonymized Collection (Default)'],
        ['[ ]', 'OPTION B: Named Collection (15% discount)'],
        ['[ ]', 'OPTION C: No Collection (Opt-out)'],
    ]
    opt_table = Table(options_table, colWidths=[40, 350])
    opt_table.setStyle(TableStyle([
        ('FONTSIZE', (0, 0), (-1, -1), 11),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 10),
        ('TOPPADDING', (0, 0), (-1, -1), 10),
    ]))
    story.append(opt_table)
    story.append(PageBreak())

    # Section 5: Data Security
    story.append(Paragraph("5. DATA SECURITY AND COMPLIANCE", styles['SectionTitle']))

    security_items = [
        "5.1 All collected data is encrypted at rest (AES-256) and in transit (TLS 1.3).",
        "5.2 Data is stored in EU-based data centers compliant with GDPR requirements.",
        "5.3 Access to training data is restricted to authorized ML engineers only.",
        "5.4 Annual third-party security audits are conducted.",
        "5.5 Data retention: Training data is retained indefinitely for model improvement. "
        "Customer may request deletion at any time per Section 7.",
        "5.6 Provider maintains SOC 2 Type II compliance for data handling.",
    ]
    for item in security_items:
        story.append(Paragraph(item, styles['ContractBody']))
        story.append(Spacer(1, 5))
    story.append(PageBreak())

    # Section 6: Usage Rights
    story.append(Paragraph("6. USAGE RIGHTS", styles['SectionTitle']))

    story.append(Paragraph(
        "6.1 <b>Training Purpose:</b> Collected data may be used solely for training and improving "
        "DataMigrate AI's machine learning models for SQL migration tasks.",
        styles['ContractBody']
    ))
    story.append(Paragraph(
        "6.2 <b>No Sale:</b> Provider will not sell, license, or transfer raw training data to "
        "third parties.",
        styles['ContractBody']
    ))
    story.append(Paragraph(
        "6.3 <b>Model Ownership:</b> All trained models remain the intellectual property of "
        "OKO Investments.",
        styles['ContractBody']
    ))
    story.append(Paragraph(
        "6.4 <b>Aggregated Insights:</b> Provider may publish aggregated, anonymized statistics "
        "about migration patterns (e.g., '60% of migrations involve stored procedures').",
        styles['ContractBody']
    ))
    story.append(PageBreak())

    # Section 7: Customer Rights
    story.append(Paragraph("7. CUSTOMER RIGHTS", styles['SectionTitle']))

    rights = [
        "7.1 <b>Right to Withdraw:</b> Customer may withdraw consent and request deletion of their "
        "contributed training data at any time by written notice.",
        "7.2 <b>Right to Access:</b> Customer may request a summary of what data has been collected "
        "from their migrations.",
        "7.3 <b>Right to Change Option:</b> Customer may change their collection option at any time "
        "for future migrations.",
        "7.4 <b>Right to Audit:</b> Customer may request an audit of data handling practices with "
        "30 days notice.",
    ]
    for right in rights:
        story.append(Paragraph(right, styles['ContractBody']))
        story.append(Spacer(1, 5))
    story.append(PageBreak())

    # Section 8: GDPR
    story.append(Paragraph("8. GDPR COMPLIANCE", styles['SectionTitle']))

    story.append(Paragraph(
        "8.1 <b>Data Controller:</b> Customer is the Data Controller for their database content. "
        "Provider is the Data Processor for Schema Data collected during migrations.",
        styles['ContractBody']
    ))
    story.append(Paragraph(
        "8.2 <b>Legal Basis:</b> Data collection is based on legitimate interest (improving AI services) "
        "with customer consent as documented in Section 4.",
        styles['ContractBody']
    ))
    story.append(Paragraph(
        "8.3 <b>Data Protection Officer:</b> Provider's DPO can be contacted at dpo@okoinvestments.com.",
        styles['ContractBody']
    ))
    story.append(Paragraph(
        "8.4 <b>Cross-Border Transfer:</b> Training data is processed within the European Economic Area. "
        "No transfer to third countries without adequate safeguards.",
        styles['ContractBody']
    ))
    story.append(PageBreak())

    # Signatures
    story.append(Paragraph("9. SIGNATURES", styles['SectionTitle']))
    story.append(Spacer(1, 30))

    sig_data = [
        ['FOR PROVIDER:', 'FOR CUSTOMER:'],
        ['', ''],
        ['_______________________________', '_______________________________'],
        ['Name: Alexander Garcia Angus', 'Name: ___________________________'],
        ['Title: CEO, OKO Investments', 'Title: ___________________________'],
        ['Date: _________________________', 'Date: ___________________________'],
    ]
    sig_table = Table(sig_data, colWidths=[220, 220])
    sig_table.setStyle(TableStyle([
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 15),
        ('TOPPADDING', (0, 0), (-1, -1), 5),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
    ]))
    story.append(sig_table)

    # Footer
    story.append(Spacer(1, 50))
    story.append(Paragraph(
        "DataMigrate AI Data Collection Agreement v1.0 | OKO Investments",
        ParagraphStyle(name='Footer', fontSize=9, textColor=colors.grey, alignment=TA_CENTER)
    ))

    doc.build(story)
    print(f"[OK] Created: {filename}")
    return filename


# =============================================================================
# DATA COLLECTION CONTRACT - WORD
# =============================================================================

def create_data_contract_word(output_dir: str = "docs"):
    """Generate Data Collection Agreement Word"""

    os.makedirs(output_dir, exist_ok=True)
    filename = f"{output_dir}/DATA_COLLECTION_AGREEMENT_TEMPLATE.docx"

    doc = Document()

    # Title
    title = doc.add_heading('DATA COLLECTION AND AI IMPROVEMENT AGREEMENT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph(
        'This Data Collection and AI Improvement Agreement ("Agreement") is entered into as of '
        '[DATE] by and between:'
    )

    doc.add_paragraph('OKO Investments ("DataMigrate AI", "Provider")')
    doc.add_paragraph('Address: [Provider Address]')
    doc.add_paragraph('CVR: [Company Registration Number]')

    doc.add_paragraph()
    doc.add_paragraph('AND')
    doc.add_paragraph()

    doc.add_paragraph('[CUSTOMER NAME] ("Customer", "Client")')
    doc.add_paragraph('Address: [Customer Address]')
    doc.add_paragraph('CVR: [Customer Registration Number]')

    doc.add_page_break()

    # Definitions
    doc.add_heading('1. DEFINITIONS', level=1)

    definitions = [
        ('"Schema Data"', 'Database structure information including table names, column names, '
         'data types, stored procedure logic. Does NOT include actual row-level data.'),
        ('"Anonymized Data"', 'Schema Data with all identifiers replaced with generic tokens.'),
        ('"Named Data"', 'Schema Data where original names are preserved, but no data values.'),
        ('"Row-Level Data"', 'Actual data values in tables. This is NEVER collected.'),
    ]

    for term, definition in definitions:
        p = doc.add_paragraph()
        p.add_run(term).bold = True
        p.add_run(f': {definition}')

    doc.add_page_break()

    # Scope
    doc.add_heading('2. SCOPE OF DATA COLLECTION', level=1)

    doc.add_heading('2.1 What We Collect', level=2)
    collect = [
        'SQL stored procedure structures',
        'CREATE TABLE statements',
        'JOIN patterns and relationships',
        'Data type mappings',
        'Generated dbt model code',
    ]
    for item in collect:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('2.2 What We NEVER Collect', level=2)
    never = [
        'Actual row-level data',
        'Personal identifiable information',
        'Financial data',
        'Health information',
        'Authentication credentials',
    ]
    for item in never:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # Options
    doc.add_heading('3. DATA COLLECTION OPTIONS', level=1)

    doc.add_heading('OPTION A: Anonymized Collection (Default)', level=2)
    doc.add_paragraph(
        'All Schema Data is automatically anonymized. Table names become Table_A, '
        'column names become col_001, etc. Maximum privacy.'
    )

    doc.add_heading('OPTION B: Named Collection (15% discount)', level=2)
    doc.add_paragraph(
        'Table and column names preserved. Improves model quality for your industry. '
        'No actual data values collected.'
    )

    doc.add_heading('OPTION C: No Collection (Opt-out)', level=2)
    doc.add_paragraph(
        'No Schema Data retained after migration. Full privacy.'
    )

    doc.add_page_break()

    # Customer Selection
    doc.add_heading('4. CUSTOMER SELECTION', level=1)
    doc.add_paragraph('Customer selects the following option (check one):')
    doc.add_paragraph()
    doc.add_paragraph('[ ] OPTION A: Anonymized Collection (Default)')
    doc.add_paragraph('[ ] OPTION B: Named Collection (15% discount)')
    doc.add_paragraph('[ ] OPTION C: No Collection (Opt-out)')

    doc.add_page_break()

    # Security
    doc.add_heading('5. DATA SECURITY', level=1)
    security = [
        'All data encrypted at rest (AES-256) and in transit (TLS 1.3)',
        'EU-based data centers, GDPR compliant',
        'Access restricted to authorized ML engineers',
        'Annual security audits',
        'SOC 2 Type II compliance',
    ]
    for item in security:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # Customer Rights
    doc.add_heading('6. CUSTOMER RIGHTS', level=1)
    rights = [
        'Right to Withdraw: Request deletion of contributed data anytime',
        'Right to Access: Request summary of collected data',
        'Right to Change: Modify collection option for future migrations',
        'Right to Audit: Request audit of data handling practices',
    ]
    for right in rights:
        doc.add_paragraph(right, style='List Bullet')

    doc.add_page_break()

    # Signatures
    doc.add_heading('7. SIGNATURES', level=1)
    doc.add_paragraph()
    doc.add_paragraph()

    sig_table = doc.add_table(rows=6, cols=2)
    sig_table.cell(0, 0).text = 'FOR PROVIDER:'
    sig_table.cell(0, 1).text = 'FOR CUSTOMER:'
    sig_table.cell(2, 0).text = '_____________________________'
    sig_table.cell(2, 1).text = '_____________________________'
    sig_table.cell(3, 0).text = 'Name: Alexander Garcia Angus'
    sig_table.cell(3, 1).text = 'Name: _____________________'
    sig_table.cell(4, 0).text = 'Title: CEO, OKO Investments'
    sig_table.cell(4, 1).text = 'Title: _____________________'
    sig_table.cell(5, 0).text = 'Date: _____________________'
    sig_table.cell(5, 1).text = 'Date: _____________________'

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph('DataMigrate AI Data Collection Agreement v1.0 | OKO Investments')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(filename)
    print(f"[OK] Created: {filename}")
    return filename


# =============================================================================
# MAIN
# =============================================================================

if __name__ == "__main__":
    print("\n" + "=" * 60)
    print("DATAMIGRATE AI - ML STRATEGY & CONTRACT GENERATOR")
    print("=" * 60 + "\n")

    try:
        # ML Strategy Documents
        ml_pdf = create_ml_strategy_pdf()
        ml_word = create_ml_strategy_word()

        # Contract Documents
        contract_pdf = create_data_contract_pdf()
        contract_word = create_data_contract_word()

        print("\n" + "=" * 60)
        print("GENERATION COMPLETE!")
        print("=" * 60)
        print(f"\nFiles created:")
        print(f"  ML Strategy:")
        print(f"    - {ml_pdf}")
        print(f"    - {ml_word}")
        print(f"  Data Collection Contract:")
        print(f"    - {contract_pdf}")
        print(f"    - {contract_word}")

    except Exception as e:
        print(f"\n[ERROR] {e}")
        import traceback
        traceback.print_exc()
