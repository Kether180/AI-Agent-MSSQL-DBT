"""
Script to generate Attack Surface Analysis in Word (.docx) and PDF formats
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from datetime import datetime

def set_cell_shading(cell, fill_color):
    """Set cell background color"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)

def add_table_with_header(doc, headers, rows, header_color='8B0000'):
    """Add a formatted table with colored header"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'

    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        header_cells[idx].text = header
        set_cell_shading(header_cells[idx], header_color)
        for paragraph in header_cells[idx].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(10)

    for row_data in rows:
        row = table.add_row()
        for idx, cell_text in enumerate(row_data):
            row.cells[idx].text = str(cell_text)
            for paragraph in row.cells[idx].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    return table

def add_risk_cell(cell, risk_level):
    """Color code risk cells"""
    colors = {
        'CRITICAL': 'FF0000',
        'HIGH': 'FF6600',
        'MEDIUM': 'FFCC00',
        'LOW': '00CC00'
    }
    if risk_level.upper() in colors:
        set_cell_shading(cell, colors[risk_level.upper()])

def create_attack_surface_document():
    doc = Document()

    # Title
    title = doc.add_heading('DataMigrate AI', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading('Attack Surface Analysis & Security Assessment', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Classification banner
    banner = doc.add_paragraph()
    banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = banner.add_run('CONFIDENTIAL - Internal Use Only')
    run.bold = True
    run.font.color.rgb = RGBColor(139, 0, 0)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run(f'Version: 1.0 | Date: {datetime.now().strftime("%B %Y")}').italic = True

    doc.add_page_break()

    # Table of Contents
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        '1. Executive Summary',
        '2. Platform Architecture Overview',
        '3. Attack Surface Mapping',
        '4. Detailed Vulnerability Analysis',
        '5. Security Controls Matrix',
        '6. Attack Priority Matrix',
        '7. Penetration Testing Guidelines',
        '8. Recommended Security Enhancements',
        '9. Incident Response Procedures',
        '10. Appendix: Security Checklist'
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Number')

    doc.add_page_break()

    # 1. Executive Summary
    doc.add_heading('1. Executive Summary', level=1)

    doc.add_heading('Platform Overview', level=2)
    doc.add_paragraph(
        'DataMigrate AI is an enterprise platform for migrating Microsoft SQL Server databases '
        'to modern cloud data warehouses (Snowflake, BigQuery, Databricks, etc.) using AI-powered automation.'
    )

    doc.add_heading('Critical Assets Protected', level=2)
    assets = [
        ('Customer Database Credentials', 'CRITICAL', 'AES-256-GCM Encryption'),
        ('User Account Passwords', 'HIGH', 'Bcrypt Hashing (cost 10)'),
        ('JWT Authentication Tokens', 'HIGH', 'HMAC-SHA256, Short Expiry'),
        ('Migration Data/Schema', 'MEDIUM', 'User Isolation, Access Controls'),
        ('AI/RAG Knowledge Base', 'MEDIUM', 'Organization Isolation'),
        ('API Keys', 'HIGH', 'Secure Generation, Single Display')
    ]
    add_table_with_header(doc, ['Asset', 'Sensitivity', 'Protection Level'], assets)

    doc.add_paragraph()
    doc.add_heading('Security Posture Summary', level=2)
    posture = [
        'Authentication: JWT-based with bcrypt password hashing',
        'Authorization: Role-based access control (RBAC) with organization isolation',
        'Encryption: AES-256-GCM at rest, TLS in transit',
        'Input Validation: Multi-layer pattern detection (SQL, XSS, Command, Prompt Injection)',
        'Rate Limiting: Sliding window algorithm (300/min, 3000/hr)',
        'Audit Logging: Comprehensive security event logging'
    ]
    for p in posture:
        doc.add_paragraph(p, style='List Bullet')

    doc.add_page_break()

    # 2. Platform Architecture
    doc.add_heading('2. Platform Architecture Overview', level=1)

    doc.add_heading('Network Boundaries', level=2)
    boundaries = [
        ('Public', 'Frontend, API Gateway', 'Untrusted'),
        ('DMZ', 'Go Backend, Python AI Service', 'Semi-Trusted'),
        ('Private', 'PostgreSQL, Internal Services', 'Trusted'),
        ('External', 'Customer Databases', 'Untrusted')
    ]
    add_table_with_header(doc, ['Zone', 'Components', 'Trust Level'], boundaries)

    doc.add_page_break()

    # 3. Attack Surface Mapping
    doc.add_heading('3. Attack Surface Mapping', level=1)

    doc.add_heading('3.1 API Endpoints (Go Backend)', level=2)
    endpoints = [
        ('/api/v1/auth/*', 'Public', 'HIGH', 'Brute force, Credential stuffing'),
        ('/api/v1/connections/*', 'JWT Required', 'CRITICAL', 'Credential theft, SSRF, Injection'),
        ('/api/v1/migrations/*', 'JWT Required', 'HIGH', 'Path traversal, Data exfiltration'),
        ('/api/v1/chat', 'JWT Required', 'HIGH', 'Prompt injection, Data leakage'),
        ('/api/v1/api-keys/*', 'JWT Required', 'HIGH', 'Key theft, Privilege escalation'),
        ('/api/v1/security/*', 'JWT + Admin', 'MEDIUM', 'Information disclosure'),
        ('/health', 'Public', 'LOW', 'Information disclosure'),
        ('/metrics', 'Public', 'MEDIUM', 'Information disclosure')
    ]
    add_table_with_header(doc, ['Endpoint', 'Auth', 'Risk', 'Attack Vectors'], endpoints)

    doc.add_page_break()

    # 4. Detailed Vulnerability Analysis
    doc.add_heading('4. Detailed Vulnerability Analysis', level=1)

    doc.add_heading('4.1 Authentication Attacks', level=2)

    doc.add_heading('A. Brute Force Login Attack', level=3)
    doc.add_paragraph('Target: POST /api/v1/auth/login')
    doc.add_paragraph('Attack Description: Attacker attempts multiple password combinations to guess valid credentials.')

    doc.add_paragraph('Current Protections:', style='Heading 4')
    protections = [
        'Rate limiting: 300 requests/minute per IP',
        'Bcrypt with cost factor 10 (~100ms per hash)',
        'Generic error messages ("Invalid credentials")'
    ]
    for p in protections:
        doc.add_paragraph(p, style='List Bullet')

    doc.add_paragraph('Recommended Enhancements:', style='Heading 4')
    enhancements = [
        'Account lockout after 5 failed attempts (15-minute lockout)',
        'CAPTCHA after 3 failed attempts',
        'Geographic anomaly detection',
        'Email notification on suspicious login patterns'
    ]
    for e in enhancements:
        doc.add_paragraph(e, style='List Bullet')

    doc.add_heading('4.2 Database Credential Attacks', level=2)

    doc.add_heading('A. Credential Theft from Database', level=3)
    doc.add_paragraph('Target: database_connections table, password column')

    doc.add_paragraph('Current Protections:', style='Heading 4')
    cred_protections = [
        'AES-256-GCM encryption',
        'Encryption key stored separately (environment variable)',
        'Key never logged'
    ]
    for p in cred_protections:
        doc.add_paragraph(p, style='List Bullet')

    doc.add_heading('B. SSRF via Database Connections', level=3)
    doc.add_paragraph('Target: POST /api/v1/connections, POST /api/v1/connections/:id/test')
    doc.add_paragraph('Vulnerability: Currently allows localhost, 127.0.0.1, and private IP ranges.')

    doc.add_paragraph('Recommended Blocks (Production):', style='Heading 4')
    ip_blocks = [
        '127.0.0.0/8 (Loopback)',
        '10.0.0.0/8 (Private)',
        '172.16.0.0/12 (Private)',
        '192.168.0.0/16 (Private)',
        '169.254.0.0/16 (Link-local)',
        '::1 (IPv6 Loopback)'
    ]
    for ip in ip_blocks:
        doc.add_paragraph(ip, style='List Bullet')

    doc.add_heading('4.3 Injection Attacks', level=2)

    doc.add_heading('SQL Injection Patterns Detected', level=3)
    sql_patterns = [
        'SELECT...FROM, INSERT INTO, DROP TABLE',
        'UNION SELECT',
        'OR 1=1, AND 1=1',
        '--, /* */',
        'xp_cmdshell, WAITFOR DELAY',
        'BENCHMARK(), SLEEP()',
        'INFORMATION_SCHEMA'
    ]
    for p in sql_patterns:
        doc.add_paragraph(p, style='List Bullet')

    doc.add_heading('XSS Patterns Detected', level=3)
    xss_patterns = [
        '<script>, javascript:',
        'onclick=, onload=, onerror=',
        '<iframe>, <embed>, <object>',
        'data:text/html',
        '<svg onload='
    ]
    for p in xss_patterns:
        doc.add_paragraph(p, style='List Bullet')

    doc.add_heading('Prompt Injection Patterns Detected', level=3)
    prompt_patterns = [
        'ignore previous instructions',
        'disregard all, forget previous',
        'system: you are, new instructions',
        '</system>, <|im_start|>',
        'HUMAN:, ASSISTANT:',
        'reveal your system prompt'
    ]
    for p in prompt_patterns:
        doc.add_paragraph(p, style='List Bullet')

    doc.add_page_break()

    # 5. Security Controls Matrix
    doc.add_heading('5. Security Controls Matrix', level=1)

    controls = [
        ('Password Hashing', 'Yes', 'auth.go', 'Strong (bcrypt)'),
        ('JWT Tokens', 'Yes', 'middleware/auth.go', 'Strong'),
        ('RBAC', 'Yes', 'middleware/auth.go', 'Strong'),
        ('SQL Injection Prevention', 'Yes', 'Guardian, validation', 'Strong'),
        ('XSS Prevention', 'Yes', 'Guardian', 'Strong'),
        ('Prompt Injection Prevention', 'Yes', 'Guardian', 'Strong'),
        ('Credentials Encryption', 'Yes', 'AES-256-GCM', 'Strong'),
        ('Rate Limiting', 'Yes', 'rate_limiter.go', 'Strong'),
        ('Audit Logging', 'Yes', 'audit_logger.go', 'Strong'),
        ('Security Headers', 'Partial', '-', 'Needs Enhancement'),
        ('MFA', 'No', '-', 'Not Implemented')
    ]
    add_table_with_header(doc, ['Control', 'Implemented', 'Location', 'Effectiveness'], controls)

    doc.add_page_break()

    # 6. Attack Priority Matrix
    doc.add_heading('6. Attack Priority Matrix', level=1)

    doc.add_paragraph('Risk Score = Likelihood x Impact')

    priorities = [
        ('Database Credential Theft', '2', '5', '10', 'CRITICAL'),
        ('SQL Injection', '2', '5', '10', 'CRITICAL'),
        ('Prompt Injection (AI)', '4', '3', '12', 'CRITICAL'),
        ('JWT Token Theft', '3', '4', '12', 'CRITICAL'),
        ('Brute Force Login', '3', '3', '9', 'HIGH'),
        ('XSS Attack', '2', '4', '8', 'HIGH'),
        ('SSRF via Connections', '2', '4', '8', 'HIGH'),
        ('API Key Abuse', '2', '4', '8', 'HIGH'),
        ('DDoS Attack', '3', '3', '9', 'HIGH'),
        ('Path Traversal', '2', '3', '6', 'MEDIUM'),
        ('File Upload Attack', '2', '3', '6', 'MEDIUM'),
        ('CSRF', '2', '2', '4', 'LOW'),
        ('Clickjacking', '2', '2', '4', 'LOW')
    ]
    add_table_with_header(doc, ['Attack', 'Likelihood', 'Impact', 'Risk Score', 'Priority'], priorities)

    doc.add_page_break()

    # 7. Penetration Testing Guidelines
    doc.add_heading('7. Penetration Testing Guidelines', level=1)

    doc.add_heading('7.1 Scope', level=2)

    doc.add_paragraph('In Scope:', style='Heading 4')
    in_scope = [
        'All API endpoints (/api/v1/*)',
        'Authentication mechanisms',
        'Authorization controls',
        'Input validation',
        'File handling',
        'AI/Chat functionality',
        'Rate limiting effectiveness'
    ]
    for s in in_scope:
        doc.add_paragraph(s, style='List Bullet')

    doc.add_paragraph('Out of Scope:', style='Heading 4')
    out_scope = [
        'Third-party services (Anthropic API, SMTP)',
        'Customer database systems',
        'Physical security',
        'Social engineering'
    ]
    for s in out_scope:
        doc.add_paragraph(s, style='List Bullet')

    doc.add_heading('7.2 Test Cases', level=2)
    test_cases = [
        ('TC-AUTH-001', 'Brute force login with rate limit bypass attempts'),
        ('TC-AUTH-002', 'Password reset token prediction/brute force'),
        ('TC-AUTH-003', 'JWT token manipulation (algorithm confusion)'),
        ('TC-AUTHZ-001', 'Horizontal privilege escalation (access other user data)'),
        ('TC-AUTHZ-002', 'Vertical privilege escalation (user to admin)'),
        ('TC-INJ-001', 'SQL injection in all input fields'),
        ('TC-INJ-004', 'XSS (reflected, stored, DOM-based)'),
        ('TC-INJ-005', 'Prompt injection variations'),
        ('TC-BL-001', 'Rate limit bypass techniques')
    ]
    add_table_with_header(doc, ['Test ID', 'Description'], test_cases)

    doc.add_heading('7.3 Recommended Tools', level=2)
    tools = [
        ('Web Scanning', 'Burp Suite Pro, OWASP ZAP'),
        ('API Testing', 'Postman, Insomnia'),
        ('SQL Injection', 'sqlmap'),
        ('Fuzzing', 'ffuf, wfuzz'),
        ('JWT Testing', 'jwt_tool')
    ]
    add_table_with_header(doc, ['Category', 'Tools'], tools)

    doc.add_page_break()

    # 8. Recommended Security Enhancements
    doc.add_heading('8. Recommended Security Enhancements', level=1)

    doc.add_heading('Immediate (P0 - Before Production)', level=2)
    p0 = [
        ('Account Lockout', 'Lock after 5 failed attempts', '4 hours'),
        ('Security Headers', 'Add CSP, X-Frame-Options, etc.', '2 hours'),
        ('Block Internal IPs', 'Prevent SSRF to private networks', '4 hours'),
        ('HTTPS Enforcement', 'Require TLS in production', '2 hours')
    ]
    add_table_with_header(doc, ['Enhancement', 'Description', 'Effort'], p0)

    doc.add_paragraph()
    doc.add_heading('Short-Term (P1 - Within 2 Weeks)', level=2)
    p1 = [
        ('MFA Implementation', 'TOTP-based two-factor auth', '16 hours'),
        ('API Key Expiration', 'Auto-expire keys after 90 days', '8 hours'),
        ('Refresh Token Rotation', 'Implement secure token refresh', '8 hours'),
        ('AI Output Filtering', 'Filter sensitive data from AI responses', '8 hours')
    ]
    add_table_with_header(doc, ['Enhancement', 'Description', 'Effort'], p1)

    doc.add_page_break()

    # 9. Incident Response
    doc.add_heading('9. Incident Response Procedures', level=1)

    doc.add_heading('9.1 Severity Levels', level=2)
    severities = [
        ('SEV-1', 'Active exploitation, data breach', 'Immediate', 'Credential theft, SQL injection success'),
        ('SEV-2', 'Attempted attack, partial success', '1 hour', 'Multiple blocked attacks'),
        ('SEV-3', 'Security anomaly detected', '4 hours', 'Unusual patterns, rate limit triggers'),
        ('SEV-4', 'Minor security event', '24 hours', 'Failed logins, blocked requests')
    ]
    add_table_with_header(doc, ['Level', 'Description', 'Response Time', 'Examples'], severities)

    doc.add_heading('9.2 Credential Theft Response', level=2)
    response_steps = [
        'Immediate: Rotate encryption key',
        'Immediate: Force password reset for all affected users',
        '1 hour: Revoke all active sessions',
        '4 hours: Notify affected customers',
        '24 hours: Full forensic analysis'
    ]
    for idx, step in enumerate(response_steps, 1):
        doc.add_paragraph(f'{idx}. {step}')

    doc.add_page_break()

    # 10. Security Checklist
    doc.add_heading('10. Appendix: Security Checklist', level=1)

    doc.add_heading('Pre-Deployment Checklist', level=2)

    checklist_items = [
        'JWT_SECRET is strong and unique (32+ characters)',
        'ENCRYPTION_KEY set for credential encryption',
        'TLS 1.2+ enforced',
        'HTTPS required for all endpoints',
        'Account lockout configured',
        'CORS configured for specific domains',
        'Internal IPs blocked for connections',
        'Rate limiting configured',
        'Security headers configured',
        'Audit logging enabled',
        'Firewall rules configured',
        'Database access restricted',
        'Backup encryption enabled'
    ]
    for item in checklist_items:
        p = doc.add_paragraph()
        p.add_run('\u2610 ').font.size = Pt(12)
        p.add_run(item)

    doc.add_page_break()

    # Document Control
    doc.add_heading('Document Control', level=1)
    doc_control = [
        ('1.0', datetime.now().strftime('%B %Y'), 'DataMigrate AI Security Team', 'Initial release')
    ]
    add_table_with_header(doc, ['Version', 'Date', 'Author', 'Changes'], doc_control)

    doc.add_paragraph()
    doc.add_paragraph()

    # Footer
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run('CONFIDENTIAL - This document contains sensitive security information.')
    run.italic = True
    run.font.color.rgb = RGBColor(139, 0, 0)

    return doc

def main():
    print("Generating Attack Surface Analysis Document...")
    doc = create_attack_surface_document()

    base_path = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    docs_path = os.path.join(base_path, 'docs')
    os.makedirs(docs_path, exist_ok=True)

    docx_path = os.path.join(docs_path, 'ATTACK_SURFACE_ANALYSIS.docx')
    doc.save(docx_path)
    print(f"Word document saved: {docx_path}")

    # Try to convert to PDF
    try:
        from docx2pdf import convert
        pdf_path = os.path.join(docs_path, 'ATTACK_SURFACE_ANALYSIS.pdf')
        convert(docx_path, pdf_path)
        print(f"PDF document saved: {pdf_path}")
    except Exception as e:
        print(f"PDF conversion note: {e}")
        print("Open the .docx file in Word and export as PDF manually.")

    print("\nDone!")

if __name__ == '__main__':
    main()
