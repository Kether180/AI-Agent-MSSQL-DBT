"""
Convert Markdown documentation files to PDF via Word
Uses markdown2 for parsing, python-docx for Word, and docx2pdf for PDF
"""

import markdown2
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx2pdf import convert
from pathlib import Path
import sys
import re
from html.parser import HTMLParser


class MarkdownToDocx:
    """Convert markdown to Word document with proper formatting"""

    def __init__(self):
        self.doc = Document()
        self._setup_styles()

    def _setup_styles(self):
        """Setup document styles"""
        # Modify existing heading styles
        styles = self.doc.styles

        # Title style
        title_style = styles['Title']
        title_style.font.size = Pt(28)
        title_style.font.color.rgb = RGBColor(79, 70, 229)  # Indigo
        title_style.font.bold = True

        # Heading 1
        h1_style = styles['Heading 1']
        h1_style.font.size = Pt(22)
        h1_style.font.color.rgb = RGBColor(79, 70, 229)
        h1_style.font.bold = True

        # Heading 2
        h2_style = styles['Heading 2']
        h2_style.font.size = Pt(16)
        h2_style.font.color.rgb = RGBColor(99, 102, 241)
        h2_style.font.bold = True

        # Heading 3
        h3_style = styles['Heading 3']
        h3_style.font.size = Pt(14)
        h3_style.font.color.rgb = RGBColor(124, 58, 237)
        h3_style.font.bold = True

    def _clean_html(self, text):
        """Remove HTML tags from text"""
        clean = re.compile('<.*?>')
        text = re.sub(clean, '', text)
        # Decode HTML entities
        text = text.replace('&lt;', '<').replace('&gt;', '>')
        text = text.replace('&amp;', '&').replace('&quot;', '"')
        text = text.replace('&#39;', "'")
        return text.strip()

    def _parse_table_html(self, html):
        """Parse HTML table and return rows"""
        rows = []
        # Find all rows
        row_pattern = re.compile(r'<tr>(.*?)</tr>', re.DOTALL)
        cell_pattern = re.compile(r'<t[hd]>(.*?)</t[hd]>', re.DOTALL)

        for row_match in row_pattern.finditer(html):
            row_html = row_match.group(1)
            cells = [self._clean_html(c.group(1)) for c in cell_pattern.finditer(row_html)]
            if cells:
                rows.append(cells)
        return rows

    def add_markdown(self, md_content):
        """Parse and add markdown content to document"""
        # Convert markdown to HTML
        html = markdown2.markdown(
            md_content,
            extras=['tables', 'fenced-code-blocks', 'header-ids', 'code-friendly']
        )

        # Process line by line
        lines = html.split('\n')
        in_code_block = False
        code_lines = []
        in_table = False
        table_html = []

        i = 0
        while i < len(lines):
            line = lines[i].strip()

            # Handle code blocks
            if '<pre>' in line or in_code_block:
                if '<pre>' in line:
                    in_code_block = True
                    code_lines = []
                    # Get content after <pre>
                    start = line.find('<pre>')
                    content = line[start + 5:]
                    if '<code>' in content:
                        content = content.replace('<code>', '')
                    code_lines.append(content)
                elif '</pre>' in line:
                    in_code_block = False
                    content = line.replace('</pre>', '').replace('</code>', '')
                    code_lines.append(content)
                    # Add code block to document
                    code_text = '\n'.join(code_lines)
                    code_text = self._clean_html(code_text)
                    if code_text.strip():
                        para = self.doc.add_paragraph()
                        run = para.add_run(code_text)
                        run.font.name = 'Consolas'
                        run.font.size = Pt(9)
                        para.paragraph_format.left_indent = Inches(0.25)
                else:
                    code_lines.append(line)
                i += 1
                continue

            # Handle tables
            if '<table>' in line:
                in_table = True
                table_html = [line]
                i += 1
                continue
            elif in_table:
                table_html.append(line)
                if '</table>' in line:
                    in_table = False
                    # Parse and add table
                    full_table = '\n'.join(table_html)
                    rows = self._parse_table_html(full_table)
                    if rows:
                        self._add_table(rows)
                i += 1
                continue

            # Handle headings
            if line.startswith('<h1'):
                text = self._clean_html(line)
                if text:
                    self.doc.add_heading(text, level=1)
            elif line.startswith('<h2'):
                text = self._clean_html(line)
                if text:
                    self.doc.add_heading(text, level=2)
            elif line.startswith('<h3'):
                text = self._clean_html(line)
                if text:
                    self.doc.add_heading(text, level=3)
            elif line.startswith('<h4'):
                text = self._clean_html(line)
                if text:
                    self.doc.add_heading(text, level=4)

            # Handle lists
            elif line.startswith('<li>'):
                text = self._clean_html(line)
                if text:
                    self.doc.add_paragraph(text, style='List Bullet')

            # Handle blockquotes
            elif line.startswith('<blockquote>'):
                text = self._clean_html(line)
                if text:
                    para = self.doc.add_paragraph()
                    para.paragraph_format.left_indent = Inches(0.5)
                    run = para.add_run(text)
                    run.font.italic = True
                    run.font.color.rgb = RGBColor(107, 114, 128)

            # Handle horizontal rules
            elif '<hr' in line:
                para = self.doc.add_paragraph()
                para.paragraph_format.space_before = Pt(12)
                para.paragraph_format.space_after = Pt(12)
                run = para.add_run('─' * 80)
                run.font.color.rgb = RGBColor(209, 213, 219)

            # Handle paragraphs
            elif line.startswith('<p>'):
                text = self._clean_html(line)
                if text:
                    para = self.doc.add_paragraph(text)
                    para.paragraph_format.space_after = Pt(8)

            # Any other content
            elif line and not line.startswith('<'):
                text = self._clean_html(line)
                if text:
                    self.doc.add_paragraph(text)

            i += 1

    def _add_table(self, rows):
        """Add a table to the document"""
        if not rows:
            return

        num_cols = max(len(row) for row in rows)
        table = self.doc.add_table(rows=len(rows), cols=num_cols)
        table.style = 'Table Grid'

        for i, row_data in enumerate(rows):
            for j, cell_text in enumerate(row_data):
                if j < num_cols:
                    cell = table.rows[i].cells[j]
                    cell.text = cell_text
                    # Style header row
                    if i == 0:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True

        self.doc.add_paragraph()  # Space after table

    def add_footer(self, title):
        """Add document footer"""
        self.doc.add_paragraph()
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run('─' * 60)
        run.font.color.rgb = RGBColor(209, 213, 219)

        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run('© 2025 DataMigrate AI. All rights reserved.')
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(156, 163, 175)

        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run('Confidential - Internal Use Only')
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(156, 163, 175)

    def save(self, filepath):
        """Save the document"""
        self.doc.save(filepath)


def convert_md_to_pdf(md_path: Path, output_dir: Path = None):
    """Convert a markdown file to PDF via Word"""

    if not md_path.exists():
        print(f"Error: File not found: {md_path}")
        return False

    # Read markdown content
    with open(md_path, 'r', encoding='utf-8') as f:
        md_content = f.read()

    # Get title from filename
    title = md_path.stem.replace('_', ' ').title()

    # Determine output paths
    if output_dir is None:
        output_dir = md_path.parent

    docx_path = output_dir / f"{md_path.stem}.docx"
    pdf_path = output_dir / f"{md_path.stem}.pdf"

    try:
        # Create Word document
        converter = MarkdownToDocx()
        converter.add_markdown(md_content)
        converter.add_footer(title)
        converter.save(docx_path)
        print(f"  [OK] Word: {docx_path.name}")

        # Convert to PDF
        convert(str(docx_path), str(pdf_path))
        print(f"  [OK] PDF:  {pdf_path.name}")

        return True
    except Exception as e:
        print(f"  [ERROR] {e}")
        return False


def main():
    """Convert specific documentation files to PDF"""

    docs_dir = Path(__file__).parent.parent / 'docs'

    # Files to convert
    files_to_convert = [
        'API_DOCUMENTATION.md',
        'DOMAIN_STRATEGY.md',
        'PHASE2_INTEGRATIONS.md',
        'RAG_AI_ML_DOCUMENTATION.md',
        'TARGET_MARKET_STRATEGY.md'
    ]

    print("=" * 60)
    print("DataMigrate AI - Markdown to PDF Converter")
    print("=" * 60)
    print()

    success_count = 0
    total_count = len(files_to_convert)

    for filename in files_to_convert:
        md_path = docs_dir / filename
        print(f"Converting: {filename}")

        if convert_md_to_pdf(md_path):
            success_count += 1
        print()

    print("=" * 60)
    print(f"Conversion complete: {success_count}/{total_count} files converted")
    print("=" * 60)

    return success_count == total_count


if __name__ == '__main__':
    success = main()
    sys.exit(0 if success else 1)
