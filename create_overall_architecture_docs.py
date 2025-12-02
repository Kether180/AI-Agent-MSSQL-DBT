"""
DataMigrate AI - Complete Architecture Documentation Generator
Creates comprehensive PDF and Word documents covering the entire system

Author: Alexander Garcia Angus
Company: OKO Investments
"""

import os
from datetime import datetime
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
)
from reportlab.lib.units import inch

# Word document imports
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def create_architecture_pdf(output_dir: str = "docs"):
    """Generate comprehensive architecture documentation PDF"""

    os.makedirs(output_dir, exist_ok=True)
    filename = f"{output_dir}/DATAMIGRATE_AI_COMPLETE_ARCHITECTURE.pdf"

    doc = SimpleDocTemplate(
        filename,
        pagesize=A4,
        rightMargin=50,
        leftMargin=50,
        topMargin=50,
        bottomMargin=50
    )

    styles = getSampleStyleSheet()

    # Custom styles
    styles.add(ParagraphStyle(
        name='MainTitle',
        parent=styles['Heading1'],
        fontSize=28,
        spaceAfter=20,
        alignment=TA_CENTER,
        textColor=colors.HexColor('#1a365d')
    ))
    styles.add(ParagraphStyle(
        name='Subtitle',
        parent=styles['Heading2'],
        fontSize=16,
        spaceAfter=30,
        alignment=TA_CENTER,
        textColor=colors.HexColor('#4a5568')
    ))
    styles.add(ParagraphStyle(
        name='SectionTitle',
        parent=styles['Heading1'],
        fontSize=18,
        spaceBefore=25,
        spaceAfter=15,
        textColor=colors.HexColor('#2d3748')
    ))
    styles.add(ParagraphStyle(
        name='SubSection',
        parent=styles['Heading3'],
        fontSize=13,
        spaceBefore=15,
        spaceAfter=8,
        textColor=colors.HexColor('#4a5568')
    ))
    styles.add(ParagraphStyle(
        name='BodyCustom',
        parent=styles['Normal'],
        fontSize=10,
        spaceAfter=8,
        leading=14,
        alignment=TA_JUSTIFY
    ))
    styles.add(ParagraphStyle(
        name='BulletText',
        parent=styles['Normal'],
        fontSize=10,
        leftIndent=20,
        spaceAfter=5,
        leading=13
    ))
    styles.add(ParagraphStyle(
        name='CodeStyle',
        fontName='Courier',
        fontSize=8,
        backColor=colors.HexColor('#f7fafc'),
        leftIndent=10,
        spaceAfter=8
    ))

    story = []

    # ==========================================================================
    # COVER PAGE
    # ==========================================================================
    story.append(Spacer(1, 80))
    story.append(Paragraph("DataMigrate AI", styles['MainTitle']))
    story.append(Spacer(1, 10))
    story.append(Paragraph("Complete System Architecture", styles['Subtitle']))
    story.append(Spacer(1, 30))
    story.append(Paragraph("Technical Documentation", styles['SubSection']))
    story.append(Paragraph("Version 2.0 - With Model Router & Fine-Tuning", styles['SubSection']))
    story.append(Spacer(1, 50))

    # Company info box
    company_info = [
        ['Company:', 'OKO Investments'],
        ['Author:', 'Alexander Garcia Angus'],
        ['Date:', datetime.now().strftime('%B %Y')],
        ['Version:', '2.0'],
    ]
    company_table = Table(company_info, colWidths=[100, 200])
    company_table.setStyle(TableStyle([
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('TEXTCOLOR', (0, 0), (0, -1), colors.HexColor('#4a5568')),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
    ]))
    story.append(company_table)
    story.append(PageBreak())

    # ==========================================================================
    # TABLE OF CONTENTS
    # ==========================================================================
    story.append(Paragraph("Table of Contents", styles['SectionTitle']))
    toc_items = [
        "1. System Overview",
        "2. High-Level Architecture",
        "3. Frontend Layer (Vue.js 3)",
        "4. Backend API Layer (Go)",
        "5. AI Agents Layer (Python + LangGraph)",
        "6. Model Router & AI Provider System",
        "7. Fine-Tuning Data Pipeline",
        "8. Guardian Agent Security",
        "9. Database Layer (PostgreSQL)",
        "10. Deployment Architecture",
        "11. Data Flow Diagrams",
        "12. Technology Stack Summary",
    ]
    for item in toc_items:
        story.append(Paragraph(item, styles['BulletText']))
    story.append(PageBreak())

    # ==========================================================================
    # 1. SYSTEM OVERVIEW
    # ==========================================================================
    story.append(Paragraph("1. System Overview", styles['SectionTitle']))
    story.append(Paragraph(
        "DataMigrate AI is an enterprise-grade platform that automates the migration of Microsoft SQL Server "
        "databases to dbt (data build tool). The system uses a multi-agent AI architecture orchestrated by "
        "LangGraph, with a Go backend for high-performance API handling and a Vue.js 3 frontend for user interaction.",
        styles['BodyCustom']
    ))
    story.append(Spacer(1, 10))

    story.append(Paragraph("Core Components:", styles['SubSection']))
    components = [
        "Frontend: Vue.js 3 + TypeScript + Pinia (Port 5173)",
        "Backend API: Go + Gin Framework (Port 8000)",
        "AI Service: Python + LangGraph + Model Router (Port 8001)",
        "Database: PostgreSQL 16",
        "Security: Guardian Agent (Python + Go implementations)",
        "AI Providers: Claude, GPT, Ollama, Custom Fine-tuned Models",
    ]
    for comp in components:
        story.append(Paragraph(f"- {comp}", styles['BulletText']))
    story.append(PageBreak())

    # ==========================================================================
    # 2. HIGH-LEVEL ARCHITECTURE
    # ==========================================================================
    story.append(Paragraph("2. High-Level Architecture", styles['SectionTitle']))
    story.append(Paragraph(
        "The system follows a microservices architecture with clear separation of concerns:",
        styles['BodyCustom']
    ))
    story.append(Spacer(1, 10))

    arch_text = """
    +-------------------+     +-------------------+     +-------------------+
    |   Vue.js 3        |---->|   Go API          |---->|   Python AI       |
    |   Frontend        |     |   (Gin)           |     |   (LangGraph)     |
    |   Port 5173       |     |   Port 8000       |     |   Port 8001       |
    +-------------------+     +-------------------+     +-------------------+
                                      |                         |
                                      v                         v
                              +-------------------+     +-------------------+
                              |   PostgreSQL      |     |   Model Router    |
                              |   Database        |     |   (Multi-Provider)|
                              +-------------------+     +-------------------+
                                                               |
                              +----------------+----------------+----------------+
                              |                |                |                |
                              v                v                v                v
                        +---------+      +---------+      +---------+      +---------+
                        | Claude  |      | GPT-4   |      | Ollama  |      | Custom  |
                        | API     |      | API     |      | (Local) |      | Model   |
                        +---------+      +---------+      +---------+      +---------+
    """
    for line in arch_text.strip().split('\n'):
        story.append(Paragraph(line, styles['CodeStyle']))
    story.append(PageBreak())

    # ==========================================================================
    # 3. FRONTEND LAYER
    # ==========================================================================
    story.append(Paragraph("3. Frontend Layer (Vue.js 3)", styles['SectionTitle']))
    story.append(Paragraph(
        "The frontend provides a modern, responsive user interface built with Vue.js 3 and TypeScript.",
        styles['BodyCustom']
    ))

    frontend_data = [
        ['Component', 'Technology', 'Purpose'],
        ['Framework', 'Vue.js 3', 'Reactive UI with Composition API'],
        ['Language', 'TypeScript', 'Type-safe development'],
        ['State Management', 'Pinia', 'Centralized state stores'],
        ['Routing', 'Vue Router 4', 'Client-side navigation'],
        ['HTTP Client', 'Axios', 'API communication'],
        ['Styling', 'Tailwind CSS', 'Utility-first CSS framework'],
        ['Build Tool', 'Vite', 'Fast development server and bundler'],
    ]
    frontend_table = Table(frontend_data, colWidths=[100, 100, 250])
    frontend_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2d3748')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
    ]))
    story.append(frontend_table)
    story.append(Spacer(1, 15))

    story.append(Paragraph("Directory Structure:", styles['SubSection']))
    frontend_dirs = [
        "frontend/src/api/ - Axios API client and endpoints",
        "frontend/src/components/ - Reusable Vue components",
        "frontend/src/views/ - Page-level components",
        "frontend/src/stores/ - Pinia state management",
        "frontend/src/router/ - Vue Router configuration",
        "frontend/src/types/ - TypeScript interfaces",
    ]
    for dir in frontend_dirs:
        story.append(Paragraph(f"- {dir}", styles['BulletText']))
    story.append(PageBreak())

    # ==========================================================================
    # 4. BACKEND API LAYER
    # ==========================================================================
    story.append(Paragraph("4. Backend API Layer (Go)", styles['SectionTitle']))
    story.append(Paragraph(
        "The Go backend provides a high-performance REST API using the Gin framework. "
        "It handles 95% of all requests (authentication, CRUD operations, data retrieval) "
        "while delegating AI-intensive tasks to the Python service.",
        styles['BodyCustom']
    ))

    backend_data = [
        ['Component', 'Technology', 'Purpose'],
        ['Framework', 'Gin', 'High-performance HTTP router'],
        ['ORM', 'GORM', 'PostgreSQL database access'],
        ['Auth', 'JWT-Go', 'Token-based authentication'],
        ['Security', 'Guardian Agent', 'Request validation and rate limiting'],
        ['Config', 'Viper', 'Configuration management'],
        ['Logging', 'Zap', 'Structured logging'],
    ]
    backend_table = Table(backend_data, colWidths=[100, 100, 250])
    backend_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2d3748')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
    ]))
    story.append(backend_table)
    story.append(Spacer(1, 15))

    story.append(Paragraph("API Endpoints:", styles['SubSection']))
    endpoints = [
        "POST /api/v1/auth/register - User registration",
        "POST /api/v1/auth/login - User login (returns JWT)",
        "GET /api/v1/migrations - List user's migrations",
        "POST /api/v1/migrations - Create new migration (triggers AI)",
        "GET /api/v1/migrations/:id - Get migration details",
        "GET /api/v1/security/dashboard - Security metrics (admin)",
        "GET /api/v1/security/audit-logs - Audit trail (admin)",
    ]
    for ep in endpoints:
        story.append(Paragraph(f"- {ep}", styles['BulletText']))
    story.append(PageBreak())

    # ==========================================================================
    # 5. AI AGENTS LAYER
    # ==========================================================================
    story.append(Paragraph("5. AI Agents Layer (Python + LangGraph)", styles['SectionTitle']))
    story.append(Paragraph(
        "The AI layer uses LangGraph to orchestrate a multi-agent workflow for database migration. "
        "Each agent specializes in a specific aspect of the migration process.",
        styles['BodyCustom']
    ))

    agents_data = [
        ['Agent', 'Responsibility', 'Output'],
        ['Assessment Agent', 'Analyzes MSSQL schema complexity', 'Complexity score, risk factors'],
        ['Planner Agent', 'Creates migration strategy', 'Ordered list of models to create'],
        ['Executor Agent', 'Generates dbt model SQL', 'dbt .sql files'],
        ['Tester Agent', 'Validates generated models', 'Test results, errors found'],
        ['Rebuilder Agent', 'Fixes failed models', 'Corrected SQL'],
        ['Optimizer Agent', 'Improves model performance', 'Optimized SQL, indexes'],
        ['Guardian Agent', 'Security validation', 'Threat detection, audit logs'],
    ]
    agents_table = Table(agents_data, colWidths=[100, 180, 170])
    agents_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2d3748')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f7fafc')]),
    ]))
    story.append(agents_table)
    story.append(Spacer(1, 15))

    story.append(Paragraph("LangGraph Workflow:", styles['SubSection']))
    workflow = [
        "1. START -> Assessment Agent (analyze schema)",
        "2. Assessment -> Planner Agent (create migration plan)",
        "3. Planner -> Executor Agent (generate dbt models)",
        "4. Executor -> Tester Agent (validate models)",
        "5. Tester -> Rebuilder Agent (if errors) OR Optimizer Agent (if success)",
        "6. Rebuilder -> Tester Agent (retry validation)",
        "7. Optimizer -> END (migration complete)",
    ]
    for step in workflow:
        story.append(Paragraph(step, styles['BulletText']))
    story.append(PageBreak())

    # ==========================================================================
    # 6. MODEL ROUTER & AI PROVIDER SYSTEM
    # ==========================================================================
    story.append(Paragraph("6. Model Router & AI Provider System", styles['SectionTitle']))
    story.append(Paragraph(
        "The Model Router provides a unified interface for multiple AI providers, enabling cost optimization, "
        "fallback chains, and future fine-tuning support. This is a key competitive advantage.",
        styles['BodyCustom']
    ))
    story.append(Spacer(1, 10))

    story.append(Paragraph("Supported Providers:", styles['SubSection']))
    providers_data = [
        ['Provider', 'Models', 'Use Case', 'Cost'],
        ['Anthropic', 'Claude Opus, Sonnet, Haiku', 'Complex reasoning, critical tasks', '$$$'],
        ['OpenAI', 'GPT-4o, GPT-4o-mini', 'General purpose, JSON generation', '$$'],
        ['Ollama', 'Llama 3, Mistral, CodeLlama', 'Simple tasks, on-premise', 'Free'],
        ['Custom/vLLM', 'Fine-tuned models', 'Domain-specific SQL generation', 'Free'],
    ]
    providers_table = Table(providers_data, colWidths=[80, 150, 150, 60])
    providers_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2d3748')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
    ]))
    story.append(providers_table)
    story.append(Spacer(1, 15))

    story.append(Paragraph("Smart Routing by Task Complexity:", styles['SubSection']))
    routing_data = [
        ['Complexity', 'Default Model', 'Example Tasks'],
        ['SIMPLE', 'Claude Haiku (cheap)', 'Basic SELECT queries, simple transformations'],
        ['MEDIUM', 'GPT-4o-mini (balanced)', 'Standard migrations, moderate logic'],
        ['COMPLEX', 'Claude Sonnet (quality)', 'Stored procedures, business logic'],
        ['CRITICAL', 'Claude Opus (best)', 'Production migrations, complex validation'],
    ]
    routing_table = Table(routing_data, colWidths=[80, 130, 240])
    routing_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2d3748')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
    ]))
    story.append(routing_table)
    story.append(Spacer(1, 15))

    story.append(Paragraph("Fallback Chains:", styles['SubSection']))
    story.append(Paragraph(
        "If a model fails or is unavailable, the router automatically tries the next model in the chain:",
        styles['BodyCustom']
    ))
    fallbacks = [
        "Claude Opus -> Claude Sonnet -> GPT-4o",
        "Claude Sonnet -> GPT-4o -> Claude Haiku",
        "Claude Haiku -> GPT-4o-mini -> Llama-3-8b",
        "Custom Model -> Claude Sonnet -> GPT-4o",
    ]
    for fb in fallbacks:
        story.append(Paragraph(f"- {fb}", styles['BulletText']))
    story.append(PageBreak())

    # ==========================================================================
    # 7. FINE-TUNING DATA PIPELINE
    # ==========================================================================
    story.append(Paragraph("7. Fine-Tuning Data Pipeline", styles['SectionTitle']))
    story.append(Paragraph(
        "The platform automatically collects high-quality migration examples for future model fine-tuning. "
        "This creates a competitive moat that grows stronger with each customer.",
        styles['BodyCustom']
    ))
    story.append(Spacer(1, 10))

    story.append(Paragraph("Data Collection Process:", styles['SubSection']))
    collection_steps = [
        "1. Migration Execution - Agents process customer schema",
        "2. Quality Scoring - Each output is scored (0-1) based on success",
        "3. Filtering - Only high-quality examples (score > 0.8) are saved",
        "4. Storage - Input/output pairs stored in JSONL format",
        "5. Anonymization - Customer-specific data is removed/masked",
    ]
    for step in collection_steps:
        story.append(Paragraph(step, styles['BulletText']))
    story.append(Spacer(1, 15))

    story.append(Paragraph("Training Data Structure:", styles['SubSection']))
    training_data_example = """
    {
      "messages": [
        {"role": "system", "content": "You are a SQL migration expert..."},
        {"role": "user", "content": "Convert this MSSQL: CREATE PROCEDURE..."},
        {"role": "assistant", "content": "SELECT customer_id FROM..."}
      ]
    }
    """
    for line in training_data_example.strip().split('\n'):
        story.append(Paragraph(line, styles['CodeStyle']))
    story.append(Spacer(1, 15))

    story.append(Paragraph("ML Engineer Workflow:", styles['SubSection']))
    ml_workflow = [
        "1. Export Data - collector.export_for_training('jsonl')",
        "2. Choose Base Model - Llama 3, Mistral, or SQLCoder",
        "3. Fine-tune - Using LoRA or full fine-tuning",
        "4. Evaluate - Test on held-out migration examples",
        "5. Deploy - Via Ollama or vLLM",
        "6. Register - router.register_custom_model('datamigrate-v1', config)",
    ]
    for step in ml_workflow:
        story.append(Paragraph(step, styles['BulletText']))
    story.append(PageBreak())

    # ==========================================================================
    # 8. GUARDIAN AGENT SECURITY
    # ==========================================================================
    story.append(Paragraph("8. Guardian Agent Security", styles['SectionTitle']))
    story.append(Paragraph(
        "Enterprise-grade security layer that protects all AI operations from malicious input and abuse.",
        styles['BodyCustom']
    ))

    security_features = [
        ['Feature', 'Implementation', 'Protection'],
        ['Prompt Injection', '25+ regex patterns', 'Blocks AI manipulation attempts'],
        ['SQL Injection', 'OWASP patterns', 'Prevents malicious SQL generation'],
        ['XSS Protection', 'HTML/JS detection', 'Blocks script injection'],
        ['Rate Limiting', 'Sliding window', 'Prevents abuse (100 req/min)'],
        ['Audit Logging', 'Full event capture', 'SOC 2/GDPR compliance'],
        ['Multi-Tenant', 'Per-org policies', 'Isolated security rules'],
    ]
    security_table = Table(security_features, colWidths=[100, 130, 220])
    security_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2d3748')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
    ]))
    story.append(security_table)
    story.append(Spacer(1, 15))

    story.append(Paragraph("Dual Implementation:", styles['SubSection']))
    story.append(Paragraph(
        "Guardian Agent is implemented in both Python (for AI agents) and Go (for API middleware):",
        styles['BodyCustom']
    ))
    implementations = [
        "Python: @protected_agent decorator wraps agent functions",
        "Go: Gin middleware validates all incoming API requests",
        "Both: Share same detection patterns and rate limit rules",
        "Both: Write to same PostgreSQL audit_logs table",
    ]
    for impl in implementations:
        story.append(Paragraph(f"- {impl}", styles['BulletText']))
    story.append(PageBreak())

    # ==========================================================================
    # 9. DATABASE LAYER
    # ==========================================================================
    story.append(Paragraph("9. Database Layer (PostgreSQL)", styles['SectionTitle']))
    story.append(Paragraph(
        "PostgreSQL 16 serves as the central data store for users, migrations, and security audit logs.",
        styles['BodyCustom']
    ))

    db_tables = [
        ['Table', 'Purpose', 'Key Fields'],
        ['users', 'User accounts', 'id, email, password_hash, org_id'],
        ['organizations', 'Multi-tenant orgs', 'id, name, plan_type'],
        ['api_keys', 'API authentication', 'id, key_hash, user_id, scopes'],
        ['migrations', 'Migration projects', 'id, user_id, status, config'],
        ['migration_models', 'Generated dbt models', 'id, migration_id, sql, status'],
        ['audit_logs', 'Security events', 'id, event_type, user_id, details'],
        ['fine_tuning_data', 'Training examples', 'id, task_type, input, output, score'],
    ]
    db_table = Table(db_tables, colWidths=[100, 150, 200])
    db_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2d3748')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
    ]))
    story.append(db_table)
    story.append(PageBreak())

    # ==========================================================================
    # 10. DEPLOYMENT ARCHITECTURE
    # ==========================================================================
    story.append(Paragraph("10. Deployment Architecture", styles['SectionTitle']))

    story.append(Paragraph("Development Environment:", styles['SubSection']))
    dev_stack = [
        "Frontend: npm run dev (Vite, Port 5173)",
        "Backend: go run cmd/server/main.go (Port 8000)",
        "AI Service: uvicorn ai_service:app (Port 8001)",
        "Database: PostgreSQL local or Docker (Port 5432)",
    ]
    for item in dev_stack:
        story.append(Paragraph(f"- {item}", styles['BulletText']))
    story.append(Spacer(1, 15))

    story.append(Paragraph("Production Environment (AWS):", styles['SubSection']))
    aws_stack = [
        "Frontend: S3 + CloudFront (static hosting)",
        "Backend: ECS Fargate (Go container)",
        "AI Service: ECS Fargate (Python container)",
        "Database: RDS PostgreSQL (Multi-AZ)",
        "Load Balancer: Application Load Balancer",
        "CDN: CloudFront for global distribution",
        "Secrets: AWS Secrets Manager",
        "Monitoring: CloudWatch + X-Ray",
    ]
    for item in aws_stack:
        story.append(Paragraph(f"- {item}", styles['BulletText']))
    story.append(Spacer(1, 15))

    story.append(Paragraph("On-Premise Option:", styles['SubSection']))
    onprem_stack = [
        "Docker Compose for all services",
        "Ollama for local AI inference (no API calls)",
        "PostgreSQL container",
        "Nginx reverse proxy",
        "No internet dependency for migrations",
    ]
    for item in onprem_stack:
        story.append(Paragraph(f"- {item}", styles['BulletText']))
    story.append(PageBreak())

    # ==========================================================================
    # 11. DATA FLOW DIAGRAMS
    # ==========================================================================
    story.append(Paragraph("11. Data Flow Diagrams", styles['SectionTitle']))

    story.append(Paragraph("User Login Flow:", styles['SubSection']))
    login_flow = [
        "1. User -> Frontend: Enter credentials",
        "2. Frontend -> Go API: POST /auth/login",
        "3. Go API -> PostgreSQL: Verify credentials",
        "4. Go API -> Frontend: Return JWT token",
        "5. Frontend: Store token in localStorage",
    ]
    for step in login_flow:
        story.append(Paragraph(step, styles['BulletText']))
    story.append(Spacer(1, 15))

    story.append(Paragraph("Migration Creation Flow:", styles['SubSection']))
    migration_flow = [
        "1. User -> Frontend: Submit migration config",
        "2. Frontend -> Go API: POST /migrations (with JWT)",
        "3. Go API -> PostgreSQL: Create migration record",
        "4. Go API -> User: Return migration ID (30ms)",
        "5. Go API -> Python AI: Async call to start agents",
        "6. Python AI -> Model Router: Select appropriate model",
        "7. Model Router -> AI Provider: Generate SQL",
        "8. Python AI -> PostgreSQL: Save generated models",
        "9. User -> Frontend: Poll for status updates",
    ]
    for step in migration_flow:
        story.append(Paragraph(step, styles['BulletText']))
    story.append(PageBreak())

    # ==========================================================================
    # 12. TECHNOLOGY STACK SUMMARY
    # ==========================================================================
    story.append(Paragraph("12. Technology Stack Summary", styles['SectionTitle']))

    tech_stack = [
        ['Layer', 'Technology', 'Version'],
        ['Frontend', 'Vue.js + TypeScript', '3.4+'],
        ['State Management', 'Pinia', '2.0+'],
        ['Styling', 'Tailwind CSS', '3.0+'],
        ['Backend API', 'Go + Gin', '1.21+'],
        ['ORM', 'GORM', '1.25+'],
        ['AI Orchestration', 'LangGraph', '0.0.40+'],
        ['AI Framework', 'LangChain', '0.1+'],
        ['AI Providers', 'Claude, GPT, Ollama', 'Latest'],
        ['Database', 'PostgreSQL', '16'],
        ['Security', 'Guardian Agent', 'Custom'],
        ['Infrastructure', 'AWS CDK', '2.0+'],
        ['Container', 'Docker', '24+'],
    ]
    tech_table = Table(tech_stack, colWidths=[130, 200, 100])
    tech_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2d3748')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f7fafc')]),
    ]))
    story.append(tech_table)

    # Footer
    story.append(Spacer(1, 50))
    story.append(Paragraph(
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')} | OKO Investments | DataMigrate AI v2.0",
        ParagraphStyle(name='Footer', fontSize=9, textColor=colors.grey, alignment=TA_CENTER)
    ))

    doc.build(story)
    print(f"[OK] Created: {filename}")
    return filename


def create_architecture_word(output_dir: str = "docs"):
    """Generate comprehensive architecture documentation Word document"""

    os.makedirs(output_dir, exist_ok=True)
    filename = f"{output_dir}/DATAMIGRATE_AI_COMPLETE_ARCHITECTURE.docx"

    doc = Document()

    # Title
    title = doc.add_heading('DataMigrate AI', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('Complete System Architecture')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('Technical Documentation').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Version 2.0 - With Model Router & Fine-Tuning').alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Company info
    info_table = doc.add_table(rows=4, cols=2)
    info_table.cell(0, 0).text = 'Company:'
    info_table.cell(0, 1).text = 'OKO Investments'
    info_table.cell(1, 0).text = 'Author:'
    info_table.cell(1, 1).text = 'Alexander Garcia Angus'
    info_table.cell(2, 0).text = 'Date:'
    info_table.cell(2, 1).text = datetime.now().strftime('%B %Y')
    info_table.cell(3, 0).text = 'Version:'
    info_table.cell(3, 1).text = '2.0'

    doc.add_page_break()

    # 1. System Overview
    doc.add_heading('1. System Overview', level=1)
    doc.add_paragraph(
        'DataMigrate AI is an enterprise-grade platform that automates the migration of Microsoft SQL Server '
        'databases to dbt (data build tool). The system uses a multi-agent AI architecture orchestrated by '
        'LangGraph, with a Go backend for high-performance API handling and a Vue.js 3 frontend.'
    )

    doc.add_heading('Core Components:', level=2)
    components = [
        'Frontend: Vue.js 3 + TypeScript + Pinia (Port 5173)',
        'Backend API: Go + Gin Framework (Port 8000)',
        'AI Service: Python + LangGraph + Model Router (Port 8001)',
        'Database: PostgreSQL 16',
        'Security: Guardian Agent (Python + Go)',
        'AI Providers: Claude, GPT, Ollama, Custom Models',
    ]
    for comp in components:
        doc.add_paragraph(comp, style='List Bullet')

    doc.add_page_break()

    # 2. High-Level Architecture
    doc.add_heading('2. High-Level Architecture', level=1)
    doc.add_paragraph('The system follows a microservices architecture:')

    arch_table = doc.add_table(rows=5, cols=3)
    arch_table.style = 'Table Grid'
    headers = ['Service', 'Technology', 'Port']
    for i, header in enumerate(headers):
        arch_table.cell(0, i).text = header

    services = [
        ('Frontend', 'Vue.js 3 + TypeScript', '5173'),
        ('Backend API', 'Go + Gin', '8000'),
        ('AI Service', 'Python + LangGraph', '8001'),
        ('Database', 'PostgreSQL', '5432'),
    ]
    for i, (svc, tech, port) in enumerate(services, 1):
        arch_table.cell(i, 0).text = svc
        arch_table.cell(i, 1).text = tech
        arch_table.cell(i, 2).text = port

    doc.add_page_break()

    # 3. AI Agents
    doc.add_heading('3. AI Agents Layer', level=1)

    agents_table = doc.add_table(rows=8, cols=3)
    agents_table.style = 'Table Grid'
    agent_headers = ['Agent', 'Responsibility', 'Output']
    for i, header in enumerate(agent_headers):
        agents_table.cell(0, i).text = header

    agents = [
        ('Assessment', 'Analyzes schema complexity', 'Risk score'),
        ('Planner', 'Creates migration strategy', 'Model order'),
        ('Executor', 'Generates dbt SQL', 'SQL files'),
        ('Tester', 'Validates models', 'Test results'),
        ('Rebuilder', 'Fixes errors', 'Corrected SQL'),
        ('Optimizer', 'Improves performance', 'Optimized SQL'),
        ('Guardian', 'Security validation', 'Audit logs'),
    ]
    for i, (agent, resp, output) in enumerate(agents, 1):
        agents_table.cell(i, 0).text = agent
        agents_table.cell(i, 1).text = resp
        agents_table.cell(i, 2).text = output

    doc.add_page_break()

    # 4. Model Router
    doc.add_heading('4. Model Router System', level=1)
    doc.add_paragraph(
        'The Model Router provides a unified interface for multiple AI providers, '
        'enabling cost optimization and future fine-tuning support.'
    )

    doc.add_heading('Supported Providers:', level=2)
    providers = [
        'Anthropic: Claude Opus, Sonnet, Haiku - Complex reasoning',
        'OpenAI: GPT-4o, GPT-4o-mini - General purpose',
        'Ollama: Llama 3, Mistral - Local, free',
        'Custom/vLLM: Fine-tuned models - Domain-specific',
    ]
    for p in providers:
        doc.add_paragraph(p, style='List Bullet')

    doc.add_heading('Smart Routing:', level=2)
    routing = [
        'SIMPLE tasks -> Claude Haiku (cheapest)',
        'MEDIUM tasks -> GPT-4o-mini (balanced)',
        'COMPLEX tasks -> Claude Sonnet (quality)',
        'CRITICAL tasks -> Claude Opus (best)',
    ]
    for r in routing:
        doc.add_paragraph(r, style='List Bullet')

    doc.add_page_break()

    # 5. Fine-Tuning Pipeline
    doc.add_heading('5. Fine-Tuning Data Pipeline', level=1)
    doc.add_paragraph(
        'The platform automatically collects high-quality migration examples '
        'for future model fine-tuning, creating a competitive moat.'
    )

    doc.add_heading('Data Collection Process:', level=2)
    collection = [
        '1. Migration generates input/output pairs',
        '2. Quality scoring (0-1) based on success',
        '3. Only high-quality (>0.8) saved',
        '4. Stored in JSONL format',
        '5. Anonymized for privacy',
    ]
    for step in collection:
        doc.add_paragraph(step)

    doc.add_heading('ML Engineer Workflow:', level=2)
    ml_steps = [
        '1. Export training data',
        '2. Choose base model (Llama, Mistral)',
        '3. Fine-tune with LoRA',
        '4. Deploy via Ollama/vLLM',
        '5. Register in Model Router',
    ]
    for step in ml_steps:
        doc.add_paragraph(step)

    doc.add_page_break()

    # 6. Guardian Agent Security
    doc.add_heading('6. Guardian Agent Security', level=1)

    security_table = doc.add_table(rows=7, cols=2)
    security_table.style = 'Table Grid'
    sec_headers = ['Feature', 'Protection']
    for i, header in enumerate(sec_headers):
        security_table.cell(0, i).text = header

    security = [
        ('Prompt Injection', '25+ detection patterns'),
        ('SQL Injection', 'OWASP pattern matching'),
        ('XSS Protection', 'Script tag detection'),
        ('Rate Limiting', '100 requests/minute'),
        ('Audit Logging', 'Full event capture'),
        ('Multi-Tenant', 'Per-org isolation'),
    ]
    for i, (feature, protection) in enumerate(security, 1):
        security_table.cell(i, 0).text = feature
        security_table.cell(i, 1).text = protection

    doc.add_page_break()

    # 7. Technology Stack
    doc.add_heading('7. Technology Stack Summary', level=1)

    tech_table = doc.add_table(rows=14, cols=3)
    tech_table.style = 'Table Grid'
    tech_headers = ['Layer', 'Technology', 'Version']
    for i, header in enumerate(tech_headers):
        tech_table.cell(0, i).text = header

    tech = [
        ('Frontend', 'Vue.js + TypeScript', '3.4+'),
        ('State Management', 'Pinia', '2.0+'),
        ('Styling', 'Tailwind CSS', '3.0+'),
        ('Backend API', 'Go + Gin', '1.21+'),
        ('ORM', 'GORM', '1.25+'),
        ('AI Orchestration', 'LangGraph', '0.0.40+'),
        ('AI Framework', 'LangChain', '0.1+'),
        ('AI Providers', 'Claude, GPT, Ollama', 'Latest'),
        ('Database', 'PostgreSQL', '16'),
        ('Security', 'Guardian Agent', 'Custom'),
        ('Infrastructure', 'AWS CDK', '2.0+'),
        ('Container', 'Docker', '24+'),
        ('Model Router', 'Custom Python', '1.0'),
    ]
    for i, (layer, technology, version) in enumerate(tech, 1):
        tech_table.cell(i, 0).text = layer
        tech_table.cell(i, 1).text = technology
        tech_table.cell(i, 2).text = version

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d")} | OKO Investments | v2.0')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(filename)
    print(f"[OK] Created: {filename}")
    return filename


if __name__ == "__main__":
    print("\n" + "=" * 60)
    print("DATAMIGRATE AI - COMPLETE ARCHITECTURE DOCUMENTATION")
    print("=" * 60 + "\n")

    try:
        pdf_file = create_architecture_pdf()
        word_file = create_architecture_word()

        print("\n" + "=" * 60)
        print("GENERATION COMPLETE!")
        print("=" * 60)
        print(f"\nFiles created:")
        print(f"  - {pdf_file}")
        print(f"  - {word_file}")

    except Exception as e:
        print(f"\n[ERROR] {e}")
        import traceback
        traceback.print_exc()
