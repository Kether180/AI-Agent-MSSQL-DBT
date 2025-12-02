"""
DataMigrate AI - Danish Market Sales Deck Generator
Creates comprehensive PDF and Word sales presentations

Author: Alexander Garcia Angus
Company: OKO Investments
"""

import os
from datetime import datetime
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
)
from reportlab.lib.units import inch

# Word document imports
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def create_sales_deck_pdf(output_dir: str = "docs"):
    """Generate comprehensive Danish market sales deck PDF"""

    os.makedirs(output_dir, exist_ok=True)
    filename = f"{output_dir}/DATAMIGRATE_AI_SALES_DECK_DENMARK.pdf"

    doc = SimpleDocTemplate(
        filename,
        pagesize=A4,
        rightMargin=50,
        leftMargin=50,
        topMargin=50,
        bottomMargin=50
    )

    styles = getSampleStyleSheet()

    # Custom styles
    styles.add(ParagraphStyle(
        name='MainTitle',
        parent=styles['Heading1'],
        fontSize=28,
        spaceAfter=20,
        alignment=TA_CENTER,
        textColor=colors.HexColor('#1a365d')
    ))
    styles.add(ParagraphStyle(
        name='Subtitle',
        parent=styles['Heading2'],
        fontSize=16,
        spaceAfter=30,
        alignment=TA_CENTER,
        textColor=colors.HexColor('#4a5568')
    ))
    styles.add(ParagraphStyle(
        name='SectionTitle',
        parent=styles['Heading1'],
        fontSize=18,
        spaceBefore=25,
        spaceAfter=15,
        textColor=colors.HexColor('#2d3748')
    ))
    styles.add(ParagraphStyle(
        name='SubSection',
        parent=styles['Heading3'],
        fontSize=13,
        spaceBefore=15,
        spaceAfter=8,
        textColor=colors.HexColor('#4a5568')
    ))
    styles.add(ParagraphStyle(
        name='BodyCustom',
        parent=styles['Normal'],
        fontSize=10,
        spaceAfter=8,
        leading=14,
        alignment=TA_JUSTIFY
    ))
    styles.add(ParagraphStyle(
        name='BulletText',
        parent=styles['Normal'],
        fontSize=10,
        leftIndent=20,
        spaceAfter=5,
        leading=13
    ))
    styles.add(ParagraphStyle(
        name='Highlight',
        parent=styles['Normal'],
        fontSize=11,
        spaceAfter=10,
        backColor=colors.HexColor('#e6fffa'),
        borderPadding=8,
        leading=14
    ))

    story = []

    # ==========================================================================
    # COVER PAGE
    # ==========================================================================
    story.append(Spacer(1, 80))
    story.append(Paragraph("DataMigrate AI", styles['MainTitle']))
    story.append(Spacer(1, 10))
    story.append(Paragraph("AI-Powered MSSQL to dbt Migration Platform", styles['Subtitle']))
    story.append(Spacer(1, 30))
    story.append(Paragraph("Sales Deck & Market Analysis", styles['SubSection']))
    story.append(Paragraph("Denmark Market Entry Strategy", styles['SubSection']))
    story.append(Spacer(1, 50))

    # Company info box
    company_info = [
        ['Company:', 'OKO Investments'],
        ['Author:', 'Alexander Garcia Angus'],
        ['Date:', datetime.now().strftime('%B %Y')],
        ['Version:', '1.0'],
    ]
    company_table = Table(company_info, colWidths=[100, 200])
    company_table.setStyle(TableStyle([
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('TEXTCOLOR', (0, 0), (0, -1), colors.HexColor('#4a5568')),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
    ]))
    story.append(company_table)
    story.append(PageBreak())

    # ==========================================================================
    # EXECUTIVE SUMMARY
    # ==========================================================================
    story.append(Paragraph("Executive Summary", styles['SectionTitle']))
    story.append(Paragraph(
        "DataMigrate AI is an enterprise-grade platform that automates the migration of Microsoft SQL Server "
        "databases to dbt (data build tool), reducing migration time by 80% and costs by 60% compared to "
        "traditional manual approaches.",
        styles['BodyCustom']
    ))
    story.append(Spacer(1, 10))

    story.append(Paragraph("Key Value Propositions:", styles['SubSection']))
    value_props = [
        "80% faster migrations - 2-4 weeks instead of 3-6 months",
        "60% cost reduction - AI automation replaces manual SQL conversion",
        "Zero data loss guarantee - Comprehensive validation at every step",
        "Enterprise security - SOC 2, GDPR, HIPAA compliant",
        "Future-proof - Continuous learning improves with each migration"
    ]
    for vp in value_props:
        story.append(Paragraph(f"- {vp}", styles['BulletText']))
    story.append(PageBreak())

    # ==========================================================================
    # DENMARK MARKET ANALYSIS
    # ==========================================================================
    story.append(Paragraph("Denmark Market Analysis", styles['SectionTitle']))

    story.append(Paragraph("Target Industries", styles['SubSection']))
    story.append(Paragraph(
        "Denmark has a strong digital economy with many companies running legacy MSSQL databases "
        "that need modernization. Our primary target industries are:",
        styles['BodyCustom']
    ))

    industries_data = [
        ['Industry', 'Example Companies', 'Typical DB Size', 'Migration Need'],
        ['Retail/E-commerce', 'Salling Group, Coop, JYSK', '500+ tables', 'High - Analytics modernization'],
        ['Financial Services', 'Danske Bank, Nordea, Saxo', '1000+ tables', 'Critical - Regulatory compliance'],
        ['Manufacturing', 'Vestas, Grundfos, Danfoss', '300+ tables', 'Medium - IoT data integration'],
        ['Logistics', 'Maersk, DSV, DFDS', '400+ tables', 'High - Real-time analytics'],
        ['Healthcare/Pharma', 'Novo Nordisk, Lundbeck', '600+ tables', 'Critical - Research data'],
    ]
    industries_table = Table(industries_data, colWidths=[90, 130, 80, 130])
    industries_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2d3748')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 8),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f7fafc')]),
    ]))
    story.append(industries_table)
    story.append(Spacer(1, 15))

    story.append(Paragraph("Company Segmentation", styles['SubSection']))
    segments_data = [
        ['Segment', 'Employees', 'Annual Revenue', 'Decision Maker', 'Sales Cycle'],
        ['SMB', '50-200', '50-200M DKK', 'CTO/IT Manager', '1-2 months'],
        ['Mid-Market', '200-1000', '200M-1B DKK', 'VP Engineering', '2-4 months'],
        ['Enterprise', '1000+', '1B+ DKK', 'CIO/CDO', '4-8 months'],
    ]
    segments_table = Table(segments_data, colWidths=[80, 80, 100, 100, 80])
    segments_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2d3748')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
    ]))
    story.append(segments_table)
    story.append(PageBreak())

    # ==========================================================================
    # PRICING STRATEGY
    # ==========================================================================
    story.append(Paragraph("Pricing Strategy - Denmark", styles['SectionTitle']))

    story.append(Paragraph("Pricing Tiers (in DKK)", styles['SubSection']))
    pricing_data = [
        ['Tier', 'Price Range', 'Includes', 'Target'],
        ['Starter', '50,000 - 150,000 DKK', 'Up to 100 tables, Basic support', 'SMB'],
        ['Professional', '150,000 - 500,000 DKK', 'Up to 500 tables, Priority support, Custom models', 'Mid-Market'],
        ['Enterprise', '500,000 - 2,000,000 DKK', 'Unlimited tables, Dedicated support, On-premise option', 'Enterprise'],
        ['Custom', 'Contact Sales', 'Full customization, SLA guarantees, Training', 'Large Enterprise'],
    ]
    pricing_table = Table(pricing_data, colWidths=[80, 120, 180, 70])
    pricing_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2d3748')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ('TOPPADDING', (0, 0), (-1, -1), 8),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f7fafc')]),
    ]))
    story.append(pricing_table)
    story.append(Spacer(1, 15))

    story.append(Paragraph("Revenue Model", styles['SubSection']))
    revenue_items = [
        "One-time migration fee (70% of revenue) - Project-based pricing",
        "Annual maintenance (20% of revenue) - Updates, support, monitoring",
        "Custom model training (10% of revenue) - Fine-tuned AI for specific schemas",
    ]
    for item in revenue_items:
        story.append(Paragraph(f"- {item}", styles['BulletText']))
    story.append(Spacer(1, 10))

    story.append(Paragraph("Competitive Pricing Comparison", styles['SubSection']))
    comparison_data = [
        ['Solution', 'Cost for 500 Tables', 'Timeline', 'Risk'],
        ['Manual Migration', '1,500,000 - 3,000,000 DKK', '4-8 months', 'High - Human error'],
        ['Traditional Tools', '800,000 - 1,500,000 DKK', '3-5 months', 'Medium - Limited automation'],
        ['DataMigrate AI', '200,000 - 400,000 DKK', '2-4 weeks', 'Low - AI validation'],
    ]
    comparison_table = Table(comparison_data, colWidths=[110, 140, 90, 110])
    comparison_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2d3748')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ('TOPPADDING', (0, 0), (-1, -1), 8),
        ('BACKGROUND', (0, 3), (-1, 3), colors.HexColor('#c6f6d5')),
    ]))
    story.append(comparison_table)
    story.append(PageBreak())

    # ==========================================================================
    # ML/FINE-TUNING BUSINESS BENEFITS
    # ==========================================================================
    story.append(Paragraph("AI/ML Competitive Advantage", styles['SectionTitle']))

    story.append(Paragraph(
        "Our proprietary Model Router and Fine-Tuning infrastructure creates a sustainable competitive moat "
        "that grows stronger with each customer engagement.",
        styles['BodyCustom']
    ))
    story.append(Spacer(1, 10))

    story.append(Paragraph("Business Benefits of ML Infrastructure", styles['SubSection']))
    ml_benefits = [
        ("Cost Reduction", "Start with premium AI (Claude/GPT), migrate to fine-tuned models = 70% cost savings"),
        ("Competitive Moat", "Each migration improves our models - competitors can't replicate our training data"),
        ("Speed Improvement", "Fine-tuned models are 3-5x faster than generic models for SQL tasks"),
        ("Domain Expertise", "Models learn industry-specific patterns (retail, finance, healthcare)"),
        ("Offline Capability", "Local models enable on-premise deployments for security-conscious clients"),
        ("Quality Improvement", "Continuous learning from validated migrations increases accuracy"),
    ]
    for title, desc in ml_benefits:
        story.append(Paragraph(f"<b>{title}:</b> {desc}", styles['BulletText']))
    story.append(Spacer(1, 15))

    story.append(Paragraph("Monetization Strategies for ML", styles['SubSection']))
    monetization = [
        ("Premium Fine-Tuning Service", "150,000-500,000 DKK", "Custom model trained on client's specific schema patterns"),
        ("Industry-Specific Models", "50,000-150,000 DKK/year", "Pre-trained models for retail, finance, healthcare"),
        ("On-Premise License", "1,000,000+ DKK", "Full platform with local AI models, no cloud dependency"),
        ("Model Marketplace", "Revenue share", "Sell industry models to other customers (anonymized)"),
        ("Consulting Services", "2,000 DKK/hour", "ML engineering for custom integrations"),
    ]
    monetization_data = [['Service', 'Price', 'Description']]
    for name, price, desc in monetization:
        monetization_data.append([name, price, desc])

    monetization_table = Table(monetization_data, colWidths=[130, 100, 220])
    monetization_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2d3748')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ('TOPPADDING', (0, 0), (-1, -1), 8),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f7fafc')]),
    ]))
    story.append(monetization_table)
    story.append(PageBreak())

    # ==========================================================================
    # ML FLYWHEEL EFFECT
    # ==========================================================================
    story.append(Paragraph("The ML Flywheel Effect", styles['SectionTitle']))

    story.append(Paragraph(
        "Our business model creates a virtuous cycle where more customers lead to better AI, "
        "which attracts more customers:",
        styles['BodyCustom']
    ))
    story.append(Spacer(1, 10))

    flywheel_steps = [
        "1. Customer Migration - Successful migration generates training data",
        "2. Data Collection - High-quality input/output pairs stored securely",
        "3. Model Training - ML engineer fine-tunes open-source models",
        "4. Better Performance - Faster, more accurate migrations",
        "5. Cost Reduction - Less reliance on expensive API calls",
        "6. Price Advantage - Lower costs enable competitive pricing",
        "7. More Customers - Better value attracts new business",
        "8. REPEAT - Each cycle strengthens the competitive moat",
    ]
    for step in flywheel_steps:
        story.append(Paragraph(step, styles['BulletText']))
    story.append(Spacer(1, 15))

    story.append(Paragraph("Projected Cost Savings (Per Migration)", styles['SubSection']))
    cost_savings_data = [
        ['Phase', 'API Costs', 'Fine-Tuned Costs', 'Savings'],
        ['Year 1 (0-50 customers)', '15,000 DKK', '15,000 DKK', '0%'],
        ['Year 2 (50-200 customers)', '15,000 DKK', '8,000 DKK', '47%'],
        ['Year 3 (200+ customers)', '15,000 DKK', '4,500 DKK', '70%'],
    ]
    cost_table = Table(cost_savings_data, colWidths=[130, 100, 100, 80])
    cost_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2d3748')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ('TOPPADDING', (0, 0), (-1, -1), 8),
    ]))
    story.append(cost_table)
    story.append(PageBreak())

    # ==========================================================================
    # MVP STATUS & ROADMAP
    # ==========================================================================
    story.append(Paragraph("MVP Status & Roadmap", styles['SectionTitle']))

    story.append(Paragraph("Current Capabilities (Ready Now)", styles['SubSection']))
    ready_features = [
        "7 AI Agents - Assessment, Planning, Execution, Testing, Rebuilding, Optimization, Guardian",
        "Vue.js 3 Frontend - Modern, responsive user interface",
        "Go Backend API - High-performance REST API with JWT auth",
        "Guardian Agent - Enterprise security (prompt injection, rate limiting, audit)",
        "Model Router - Multi-provider AI support (Claude, GPT, Ollama, custom)",
        "Fine-Tuning Pipeline - Data collection for future model training",
        "Mock Mode - Full demo capability without database connection",
    ]
    for feature in ready_features:
        story.append(Paragraph(f"[OK] {feature}", styles['BulletText']))
    story.append(Spacer(1, 15))

    story.append(Paragraph("MVP Completion Needed", styles['SubSection']))
    needed_features = [
        ("Live MSSQL Connection", "2 weeks", "Connect to real customer databases"),
        ("Real dbt Compile/Run", "1 week", "Execute actual dbt commands"),
        ("Payment Integration", "2 weeks", "Stripe/invoice billing system"),
        ("Data Validation", "1 week", "Row count and checksum verification"),
        ("Progress Dashboard", "1 week", "Real-time migration monitoring"),
    ]
    needed_data = [['Feature', 'Effort', 'Description']]
    for name, effort, desc in needed_features:
        needed_data.append([name, effort, desc])

    needed_table = Table(needed_data, colWidths=[130, 80, 240])
    needed_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#e53e3e')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ('TOPPADDING', (0, 0), (-1, -1), 8),
    ]))
    story.append(needed_table)
    story.append(Spacer(1, 10))
    story.append(Paragraph("<b>Total MVP Completion: ~7 weeks of development</b>", styles['BodyCustom']))
    story.append(PageBreak())

    # ==========================================================================
    # GO-TO-MARKET STRATEGY
    # ==========================================================================
    story.append(Paragraph("Go-To-Market Strategy - Denmark", styles['SectionTitle']))

    story.append(Paragraph("Phase 1: Foundation (Months 1-3)", styles['SubSection']))
    phase1 = [
        "Complete MVP features (7 weeks)",
        "Establish Danish entity (ApS or IVS)",
        "Build initial pipeline through LinkedIn outreach",
        "Partner with 2-3 Danish consultancies (Netcompany, Avanade, Tata)",
        "Attend IT Branchen and Danish Tech events",
    ]
    for item in phase1:
        story.append(Paragraph(f"- {item}", styles['BulletText']))

    story.append(Paragraph("Phase 2: Early Traction (Months 4-6)", styles['SubSection']))
    phase2 = [
        "Land 2-3 pilot customers (free or discounted)",
        "Generate case studies and testimonials",
        "Collect training data for fine-tuning",
        "Hire Danish sales representative",
        "Apply for Danish Innovation Fund grant",
    ]
    for item in phase2:
        story.append(Paragraph(f"- {item}", styles['BulletText']))

    story.append(Paragraph("Phase 3: Scale (Months 7-12)", styles['SubSection']))
    phase3 = [
        "Deploy first fine-tuned models",
        "Expand to 10+ paying customers",
        "Launch industry-specific packages",
        "Partner with Microsoft Denmark",
        "Target: 2-5M DKK ARR by end of Year 1",
    ]
    for item in phase3:
        story.append(Paragraph(f"- {item}", styles['BulletText']))
    story.append(PageBreak())

    # ==========================================================================
    # ENTERPRISE SECURITY SELLING POINTS
    # ==========================================================================
    story.append(Paragraph("Enterprise Security - Key Selling Points", styles['SectionTitle']))

    story.append(Paragraph(
        "Danish enterprises, especially in financial services and healthcare, have strict security requirements. "
        "Our Guardian Agent addresses these concerns:",
        styles['BodyCustom']
    ))
    story.append(Spacer(1, 10))

    security_points = [
        ("GDPR Compliance", "All data processing logged for audit, data minimization enforced"),
        ("Prompt Injection Prevention", "25+ patterns block AI manipulation attempts"),
        ("SQL Injection Protection", "All generated SQL validated before execution"),
        ("Rate Limiting", "Prevents abuse and DoS attacks"),
        ("Audit Trail", "Complete logging for SOC 2, HIPAA compliance"),
        ("Multi-Tenant Isolation", "Per-organization security policies"),
        ("On-Premise Option", "Local deployment with no cloud data transfer"),
    ]

    security_data = [['Feature', 'Description']]
    for name, desc in security_points:
        security_data.append([name, desc])

    security_table = Table(security_data, colWidths=[150, 300])
    security_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2d3748')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ('TOPPADDING', (0, 0), (-1, -1), 8),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f7fafc')]),
    ]))
    story.append(security_table)
    story.append(PageBreak())

    # ==========================================================================
    # FINANCIAL PROJECTIONS
    # ==========================================================================
    story.append(Paragraph("Financial Projections (Denmark)", styles['SectionTitle']))

    story.append(Paragraph("Year 1 Revenue Targets", styles['SubSection']))
    year1_data = [
        ['Quarter', 'Customers', 'Avg Deal Size', 'Revenue', 'Cumulative'],
        ['Q1', '1 (pilot)', '0 DKK', '0 DKK', '0 DKK'],
        ['Q2', '3', '150,000 DKK', '450,000 DKK', '450,000 DKK'],
        ['Q3', '5', '200,000 DKK', '1,000,000 DKK', '1,450,000 DKK'],
        ['Q4', '6', '250,000 DKK', '1,500,000 DKK', '2,950,000 DKK'],
    ]
    year1_table = Table(year1_data, colWidths=[60, 80, 100, 100, 100])
    year1_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2d3748')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ('TOPPADDING', (0, 0), (-1, -1), 8),
        ('BACKGROUND', (0, -1), (-1, -1), colors.HexColor('#c6f6d5')),
    ]))
    story.append(year1_table)
    story.append(Spacer(1, 15))

    story.append(Paragraph("3-Year Projection", styles['SubSection']))
    projection_data = [
        ['Metric', 'Year 1', 'Year 2', 'Year 3'],
        ['Customers', '15', '50', '150'],
        ['Revenue', '3M DKK', '12M DKK', '40M DKK'],
        ['Gross Margin', '70%', '80%', '85%'],
        ['Team Size', '3', '8', '20'],
    ]
    projection_table = Table(projection_data, colWidths=[100, 100, 100, 100])
    projection_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2d3748')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 10),
        ('TOPPADDING', (0, 0), (-1, -1), 10),
    ]))
    story.append(projection_table)
    story.append(PageBreak())

    # ==========================================================================
    # CONTACT & NEXT STEPS
    # ==========================================================================
    story.append(Paragraph("Next Steps", styles['SectionTitle']))

    story.append(Paragraph("For Interested Companies:", styles['SubSection']))
    next_steps = [
        "1. Schedule a 30-minute demo call",
        "2. Provide sample schema (anonymized) for assessment",
        "3. Receive free migration complexity report",
        "4. Discuss pilot program terms",
        "5. Start migration within 2 weeks",
    ]
    for step in next_steps:
        story.append(Paragraph(step, styles['BulletText']))
    story.append(Spacer(1, 30))

    story.append(Paragraph("Contact Information", styles['SubSection']))
    contact_info = [
        "Company: OKO Investments",
        "Contact: Alexander Garcia Angus",
        "Email: alexander@okoinvestments.com",
        "Website: www.datamigrate.ai",
        "LinkedIn: linkedin.com/in/alexandergarciaangus",
    ]
    for info in contact_info:
        story.append(Paragraph(info, styles['BodyCustom']))

    # Footer
    story.append(Spacer(1, 50))
    story.append(Paragraph(
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')} | OKO Investments | Confidential",
        ParagraphStyle(name='Footer', fontSize=9, textColor=colors.grey, alignment=TA_CENTER)
    ))

    doc.build(story)
    print(f"[OK] Created: {filename}")
    return filename


def create_sales_deck_word(output_dir: str = "docs"):
    """Generate comprehensive Danish market sales deck Word document"""

    os.makedirs(output_dir, exist_ok=True)
    filename = f"{output_dir}/DATAMIGRATE_AI_SALES_DECK_DENMARK.docx"

    doc = Document()

    # Title
    title = doc.add_heading('DataMigrate AI', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('AI-Powered MSSQL to dbt Migration Platform')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('Sales Deck & Market Analysis').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Denmark Market Entry Strategy').alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Company info
    info_table = doc.add_table(rows=4, cols=2)
    info_table.cell(0, 0).text = 'Company:'
    info_table.cell(0, 1).text = 'OKO Investments'
    info_table.cell(1, 0).text = 'Author:'
    info_table.cell(1, 1).text = 'Alexander Garcia Angus'
    info_table.cell(2, 0).text = 'Date:'
    info_table.cell(2, 1).text = datetime.now().strftime('%B %Y')
    info_table.cell(3, 0).text = 'Version:'
    info_table.cell(3, 1).text = '1.0'

    doc.add_page_break()

    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(
        'DataMigrate AI is an enterprise-grade platform that automates the migration of Microsoft SQL Server '
        'databases to dbt (data build tool), reducing migration time by 80% and costs by 60% compared to '
        'traditional manual approaches.'
    )

    doc.add_heading('Key Value Propositions:', level=2)
    value_props = [
        '80% faster migrations - 2-4 weeks instead of 3-6 months',
        '60% cost reduction - AI automation replaces manual SQL conversion',
        'Zero data loss guarantee - Comprehensive validation at every step',
        'Enterprise security - SOC 2, GDPR, HIPAA compliant',
        'Future-proof - Continuous learning improves with each migration',
    ]
    for vp in value_props:
        doc.add_paragraph(vp, style='List Bullet')

    doc.add_page_break()

    # Denmark Market Analysis
    doc.add_heading('Denmark Market Analysis', level=1)

    doc.add_heading('Target Industries', level=2)
    doc.add_paragraph(
        'Denmark has a strong digital economy with many companies running legacy MSSQL databases '
        'that need modernization.'
    )

    industries_table = doc.add_table(rows=6, cols=4)
    industries_table.style = 'Table Grid'
    headers = ['Industry', 'Example Companies', 'Typical DB Size', 'Migration Need']
    for i, header in enumerate(headers):
        industries_table.cell(0, i).text = header

    industries = [
        ('Retail/E-commerce', 'Salling Group, Coop, JYSK', '500+ tables', 'High'),
        ('Financial Services', 'Danske Bank, Nordea, Saxo', '1000+ tables', 'Critical'),
        ('Manufacturing', 'Vestas, Grundfos, Danfoss', '300+ tables', 'Medium'),
        ('Logistics', 'Maersk, DSV, DFDS', '400+ tables', 'High'),
        ('Healthcare/Pharma', 'Novo Nordisk, Lundbeck', '600+ tables', 'Critical'),
    ]
    for i, (ind, companies, size, need) in enumerate(industries, 1):
        industries_table.cell(i, 0).text = ind
        industries_table.cell(i, 1).text = companies
        industries_table.cell(i, 2).text = size
        industries_table.cell(i, 3).text = need

    doc.add_page_break()

    # Pricing Strategy
    doc.add_heading('Pricing Strategy - Denmark', level=1)

    doc.add_heading('Pricing Tiers (in DKK)', level=2)
    pricing_table = doc.add_table(rows=5, cols=4)
    pricing_table.style = 'Table Grid'
    pricing_headers = ['Tier', 'Price Range', 'Includes', 'Target']
    for i, header in enumerate(pricing_headers):
        pricing_table.cell(0, i).text = header

    pricing = [
        ('Starter', '50,000 - 150,000 DKK', 'Up to 100 tables', 'SMB'),
        ('Professional', '150,000 - 500,000 DKK', 'Up to 500 tables', 'Mid-Market'),
        ('Enterprise', '500,000 - 2,000,000 DKK', 'Unlimited tables', 'Enterprise'),
        ('Custom', 'Contact Sales', 'Full customization', 'Large Enterprise'),
    ]
    for i, (tier, price, includes, target) in enumerate(pricing, 1):
        pricing_table.cell(i, 0).text = tier
        pricing_table.cell(i, 1).text = price
        pricing_table.cell(i, 2).text = includes
        pricing_table.cell(i, 3).text = target

    doc.add_paragraph()

    doc.add_heading('Competitive Comparison', level=2)
    comparison_table = doc.add_table(rows=4, cols=4)
    comparison_table.style = 'Table Grid'
    comp_headers = ['Solution', 'Cost (500 tables)', 'Timeline', 'Risk']
    for i, header in enumerate(comp_headers):
        comparison_table.cell(0, i).text = header

    comparisons = [
        ('Manual Migration', '1.5-3M DKK', '4-8 months', 'High'),
        ('Traditional Tools', '800K-1.5M DKK', '3-5 months', 'Medium'),
        ('DataMigrate AI', '200-400K DKK', '2-4 weeks', 'Low'),
    ]
    for i, (sol, cost, time, risk) in enumerate(comparisons, 1):
        comparison_table.cell(i, 0).text = sol
        comparison_table.cell(i, 1).text = cost
        comparison_table.cell(i, 2).text = time
        comparison_table.cell(i, 3).text = risk

    doc.add_page_break()

    # AI/ML Business Benefits
    doc.add_heading('AI/ML Competitive Advantage', level=1)

    doc.add_paragraph(
        'Our proprietary Model Router and Fine-Tuning infrastructure creates a sustainable competitive moat '
        'that grows stronger with each customer engagement.'
    )

    doc.add_heading('Business Benefits of ML Infrastructure', level=2)
    ml_benefits = [
        'Cost Reduction: Start with premium AI, migrate to fine-tuned models = 70% cost savings',
        'Competitive Moat: Each migration improves our models - competitors can\'t replicate our data',
        'Speed Improvement: Fine-tuned models are 3-5x faster than generic models',
        'Domain Expertise: Models learn industry-specific patterns',
        'Offline Capability: Local models enable on-premise deployments',
        'Quality Improvement: Continuous learning increases accuracy',
    ]
    for benefit in ml_benefits:
        doc.add_paragraph(benefit, style='List Bullet')

    doc.add_heading('Monetization Strategies', level=2)
    monetization = [
        'Premium Fine-Tuning Service: 150,000-500,000 DKK - Custom model for client schema',
        'Industry-Specific Models: 50,000-150,000 DKK/year - Pre-trained for retail, finance, healthcare',
        'On-Premise License: 1,000,000+ DKK - Full platform with local AI models',
        'Model Marketplace: Revenue share - Sell industry models to other customers',
        'Consulting Services: 2,000 DKK/hour - ML engineering for custom integrations',
    ]
    for item in monetization:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # ML Flywheel
    doc.add_heading('The ML Flywheel Effect', level=1)

    doc.add_paragraph(
        'Our business model creates a virtuous cycle where more customers lead to better AI:'
    )

    flywheel = [
        '1. Customer Migration - Successful migration generates training data',
        '2. Data Collection - High-quality input/output pairs stored securely',
        '3. Model Training - ML engineer fine-tunes open-source models',
        '4. Better Performance - Faster, more accurate migrations',
        '5. Cost Reduction - Less reliance on expensive API calls',
        '6. Price Advantage - Lower costs enable competitive pricing',
        '7. More Customers - Better value attracts new business',
        '8. REPEAT - Each cycle strengthens the competitive moat',
    ]
    for step in flywheel:
        doc.add_paragraph(step)

    doc.add_heading('Projected Cost Savings', level=2)
    savings_table = doc.add_table(rows=4, cols=4)
    savings_table.style = 'Table Grid'
    savings_headers = ['Phase', 'API Costs', 'Fine-Tuned Costs', 'Savings']
    for i, header in enumerate(savings_headers):
        savings_table.cell(0, i).text = header

    savings = [
        ('Year 1 (0-50 customers)', '15,000 DKK', '15,000 DKK', '0%'),
        ('Year 2 (50-200 customers)', '15,000 DKK', '8,000 DKK', '47%'),
        ('Year 3 (200+ customers)', '15,000 DKK', '4,500 DKK', '70%'),
    ]
    for i, (phase, api, ft, save) in enumerate(savings, 1):
        savings_table.cell(i, 0).text = phase
        savings_table.cell(i, 1).text = api
        savings_table.cell(i, 2).text = ft
        savings_table.cell(i, 3).text = save

    doc.add_page_break()

    # MVP Status
    doc.add_heading('MVP Status & Roadmap', level=1)

    doc.add_heading('Current Capabilities (Ready Now)', level=2)
    ready = [
        '7 AI Agents - Assessment, Planning, Execution, Testing, Rebuilding, Optimization, Guardian',
        'Vue.js 3 Frontend - Modern, responsive user interface',
        'Go Backend API - High-performance REST API with JWT auth',
        'Guardian Agent - Enterprise security (prompt injection, rate limiting, audit)',
        'Model Router - Multi-provider AI support (Claude, GPT, Ollama, custom)',
        'Fine-Tuning Pipeline - Data collection for future model training',
    ]
    for item in ready:
        doc.add_paragraph(f'[OK] {item}', style='List Bullet')

    doc.add_heading('MVP Completion Needed (~7 weeks)', level=2)
    needed = [
        'Live MSSQL Connection (2 weeks)',
        'Real dbt Compile/Run (1 week)',
        'Payment Integration (2 weeks)',
        'Data Validation (1 week)',
        'Progress Dashboard (1 week)',
    ]
    for item in needed:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # Financial Projections
    doc.add_heading('Financial Projections (Denmark)', level=1)

    doc.add_heading('Year 1 Revenue Targets', level=2)
    revenue_table = doc.add_table(rows=5, cols=5)
    revenue_table.style = 'Table Grid'
    rev_headers = ['Quarter', 'Customers', 'Avg Deal', 'Revenue', 'Cumulative']
    for i, header in enumerate(rev_headers):
        revenue_table.cell(0, i).text = header

    revenue = [
        ('Q1', '1 (pilot)', '0', '0', '0'),
        ('Q2', '3', '150K', '450K', '450K'),
        ('Q3', '5', '200K', '1M', '1.45M'),
        ('Q4', '6', '250K', '1.5M', '2.95M'),
    ]
    for i, (q, cust, avg, rev, cum) in enumerate(revenue, 1):
        revenue_table.cell(i, 0).text = q
        revenue_table.cell(i, 1).text = cust
        revenue_table.cell(i, 2).text = avg
        revenue_table.cell(i, 3).text = rev
        revenue_table.cell(i, 4).text = cum

    doc.add_paragraph()

    doc.add_heading('3-Year Projection', level=2)
    projection_table = doc.add_table(rows=5, cols=4)
    projection_table.style = 'Table Grid'
    proj_headers = ['Metric', 'Year 1', 'Year 2', 'Year 3']
    for i, header in enumerate(proj_headers):
        projection_table.cell(0, i).text = header

    projections = [
        ('Customers', '15', '50', '150'),
        ('Revenue', '3M DKK', '12M DKK', '40M DKK'),
        ('Gross Margin', '70%', '80%', '85%'),
        ('Team Size', '3', '8', '20'),
    ]
    for i, (metric, y1, y2, y3) in enumerate(projections, 1):
        projection_table.cell(i, 0).text = metric
        projection_table.cell(i, 1).text = y1
        projection_table.cell(i, 2).text = y2
        projection_table.cell(i, 3).text = y3

    doc.add_page_break()

    # Contact
    doc.add_heading('Next Steps & Contact', level=1)

    doc.add_heading('For Interested Companies:', level=2)
    steps = [
        '1. Schedule a 30-minute demo call',
        '2. Provide sample schema (anonymized) for assessment',
        '3. Receive free migration complexity report',
        '4. Discuss pilot program terms',
        '5. Start migration within 2 weeks',
    ]
    for step in steps:
        doc.add_paragraph(step)

    doc.add_paragraph()
    doc.add_heading('Contact Information', level=2)
    contact = [
        'Company: OKO Investments',
        'Contact: Alexander Garcia Angus',
        'Email: alexander@okoinvestments.com',
        'Website: www.datamigrate.ai',
    ]
    for info in contact:
        doc.add_paragraph(info)

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d")} | OKO Investments | Confidential')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(filename)
    print(f"[OK] Created: {filename}")
    return filename


if __name__ == "__main__":
    print("\n" + "=" * 60)
    print("DATAMIGRATE AI - SALES DECK GENERATOR")
    print("=" * 60 + "\n")

    try:
        pdf_file = create_sales_deck_pdf()
        word_file = create_sales_deck_word()

        print("\n" + "=" * 60)
        print("GENERATION COMPLETE!")
        print("=" * 60)
        print(f"\nFiles created:")
        print(f"  - {pdf_file}")
        print(f"  - {word_file}")

    except Exception as e:
        print(f"\n[ERROR] {e}")
        import traceback
        traceback.print_exc()
