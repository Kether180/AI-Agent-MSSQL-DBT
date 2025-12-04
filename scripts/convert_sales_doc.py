#!/usr/bin/env python3
"""
Convert DataMigrate AI Sales Document from Markdown to Word (.docx)
"""

import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from pathlib import Path


def create_word_document():
    """Create a professionally formatted Word document from the sales content."""

    doc = Document()

    # Set up styles
    styles = doc.styles

    # Title style
    title_style = styles['Title']
    title_style.font.size = Pt(28)
    title_style.font.color.rgb = RGBColor(79, 70, 229)  # Indigo
    title_style.font.bold = True

    # Heading 1 style
    h1_style = styles['Heading 1']
    h1_style.font.size = Pt(20)
    h1_style.font.color.rgb = RGBColor(79, 70, 229)
    h1_style.font.bold = True

    # Heading 2 style
    h2_style = styles['Heading 2']
    h2_style.font.size = Pt(16)
    h2_style.font.color.rgb = RGBColor(99, 102, 241)
    h2_style.font.bold = True

    # Heading 3 style
    h3_style = styles['Heading 3']
    h3_style.font.size = Pt(14)
    h3_style.font.color.rgb = RGBColor(67, 56, 202)
    h3_style.font.bold = True

    # ===== COVER PAGE =====
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('DataMigrate AI')
    run.font.size = Pt(48)
    run.font.color.rgb = RGBColor(79, 70, 229)
    run.font.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Enterprise Sales Document')
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(107, 114, 128)

    doc.add_paragraph()
    doc.add_paragraph()

    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tagline.add_run('Intelligent Data Migration for the Modern Enterprise')
    run.font.size = Pt(16)
    run.font.italic = True
    run.font.color.rgb = RGBColor(107, 114, 128)

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    # Document info
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run('Version 2.0 | December 2025 | Confidential')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(156, 163, 175)

    doc.add_page_break()

    # ===== EXECUTIVE SUMMARY =====
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(
        'DataMigrate AI is a revolutionary AI-powered platform that automates the migration of legacy '
        'MSSQL databases to modern dbt (data build tool) projects. Our multi-agent AI architecture '
        'delivers unprecedented speed, accuracy, and cost savings for enterprises looking to modernize '
        'their data infrastructure.'
    )
    doc.add_paragraph()

    # ===== MISSION =====
    doc.add_heading('Our Mission', level=1)
    mission_para = doc.add_paragraph()
    run = mission_para.add_run(
        'To democratize data modernization by making legacy database migration accessible, affordable, '
        'and intelligent for organizations of all sizes.'
    )
    run.font.italic = True
    run.font.size = Pt(12)
    doc.add_paragraph(
        'We believe every company deserves access to modern data analytics infrastructure without the '
        'prohibitive costs and technical barriers of traditional migration approaches.'
    )
    doc.add_paragraph()

    # ===== VISION =====
    doc.add_heading('Our Vision', level=1)
    vision_para = doc.add_paragraph()
    run = vision_para.add_run(
        'To become the global leader in AI-powered data migration, enabling 10,000+ enterprises '
        'worldwide to transition from legacy systems to modern data stacks by 2028.'
    )
    run.font.italic = True
    run.font.size = Pt(12)

    doc.add_paragraph('We envision a future where:')
    bullets = [
        'Data migrations are completed in days, not months',
        'Zero-code solutions empower non-technical teams',
        'AI agents handle 95% of migration complexity automatically',
        'Every business decision is powered by accessible, clean data'
    ]
    for bullet in bullets:
        doc.add_paragraph(bullet, style='List Bullet')
    doc.add_paragraph()

    # ===== STRATEGIC GOALS =====
    doc.add_heading('Strategic Goals', level=1)

    doc.add_heading('Short-Term Goals (Year 1)', level=2)
    short_goals = [
        'Market Launch - Successfully deploy to 100+ enterprise customers',
        'Platform Stability - Achieve 99.9% uptime and reliability',
        'Customer Success - Maintain 95%+ customer satisfaction rate',
        'Feature Completion - Launch all three AI agents (Migration, Support, Business Intelligence)'
    ]
    for goal in short_goals:
        doc.add_paragraph(goal, style='List Bullet')

    doc.add_heading('Medium-Term Goals (Years 2-3)', level=2)
    medium_goals = [
        'Market Expansion - Expand to European and Asian markets',
        'Revenue Growth - Achieve $10M ARR milestone',
        'Technology Leadership - Develop proprietary ML models for migration optimization',
        'Partnership Ecosystem - Establish partnerships with major cloud providers (AWS, Azure, GCP)'
    ]
    for goal in medium_goals:
        doc.add_paragraph(goal, style='List Bullet')

    doc.add_heading('Long-Term Goals (Years 4-5)', level=2)
    long_goals = [
        'Market Dominance - Capture 25% of the enterprise data migration market',
        'Product Suite - Expand to support all major database systems',
        'AI Innovation - Lead the industry in intelligent data transformation',
        'Global Presence - Operations in 20+ countries'
    ]
    for goal in long_goals:
        doc.add_paragraph(goal, style='List Bullet')

    doc.add_page_break()

    # ===== COMPETITIVE ADVANTAGES =====
    doc.add_heading('Competitive Advantages', level=1)

    doc.add_heading('1. Eight-Agent AI Architecture', level=2)
    doc.add_paragraph(
        'Unlike competitors who offer single-purpose tools, DataMigrate AI features a comprehensive '
        'eight-agent AI system with RAG, Security, and DataPrep capabilities:'
    )

    # Agent table
    table = doc.add_table(rows=10, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['Agent', 'Function', 'Competitive Edge']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        from docx.oxml.ns import nsdecls
        from docx.oxml import parse_xml
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="4F46E5"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    agent_data = [
        ['Migration Agent', 'MSSQL to dbt conversion', '90% faster migrations'],
        ['Customer Support Agent', '24/7 AI assistance', '60% lower support costs'],
        ['Business Intelligence Agent', 'Data insights & analytics', 'Unlock hidden value'],
        ['Data Quality Agent', 'Validation & reconciliation', '99.9% data accuracy'],
        ['Documentation Agent (RAG)', 'RAG-powered auto docs', 'Zero manual documentation'],
        ['Security Agent', 'Threat detection & compliance', 'Enterprise-grade security'],
        ['ML Fine-Tuning Agent', 'Custom model training', 'Industry-specific AI'],
        ['DataPrep AI Agent', 'Data preparation & cleaning', 'Analytics & ML-ready data']
    ]
    for row_idx, row_data in enumerate(agent_data, 1):
        for col_idx, cell_data in enumerate(row_data):
            table.rows[row_idx].cells[col_idx].text = cell_data

    doc.add_paragraph()

    doc.add_heading('2. MSSQL Specialization', level=2)
    doc.add_paragraph('While competitors offer generic migration tools, we specialize in:')
    specializations = [
        'Deep MSSQL Integration - Native understanding of SQL Server schemas, stored procedures, and data types',
        'Windows Authentication Support - Seamless enterprise security integration',
        'Complex Query Translation - Intelligent conversion of T-SQL to dbt models'
    ]
    for spec in specializations:
        doc.add_paragraph(spec, style='List Bullet')

    doc.add_heading('3. Cost Efficiency', level=2)
    cost_benefits = [
        '70% lower cost than traditional consulting-based migrations',
        'No per-row pricing - Unlimited data migration',
        'Predictable pricing - Flat monthly subscription'
    ]
    for benefit in cost_benefits:
        doc.add_paragraph(benefit, style='List Bullet')

    doc.add_heading('4. Speed to Value', level=2)

    speed_table = doc.add_table(rows=4, cols=3)
    speed_table.style = 'Table Grid'
    speed_headers = ['Migration Size', 'Traditional Approach', 'DataMigrate AI']
    for i, header in enumerate(speed_headers):
        cell = speed_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True

    speed_data = [
        ['Small (10 tables)', '2-4 weeks', '1-2 days'],
        ['Medium (50 tables)', '2-3 months', '1-2 weeks'],
        ['Large (200+ tables)', '6-12 months', '4-6 weeks']
    ]
    for row_idx, row_data in enumerate(speed_data, 1):
        for col_idx, cell_data in enumerate(row_data):
            speed_table.rows[row_idx].cells[col_idx].text = cell_data

    doc.add_page_break()

    # ===== MARKET COMPETITIVE ANALYSIS =====
    doc.add_heading('Market Competitive Analysis', level=1)
    doc.add_paragraph(
        'The data migration and transformation market is growing rapidly, but few solutions address '
        'the specific challenge of MSSQL to dbt migration with AI-powered automation.'
    )

    doc.add_heading('Key Competitors', level=2)

    # Datafold
    doc.add_heading('Datafold', level=3)
    doc.add_paragraph('Focus: Data diff, migration testing, and CI/CD for data')
    doc.add_paragraph('Strengths:')
    datafold_strengths = [
        'Cross-database data diffing (primary differentiator)',
        'Strong focus on data validation and testing',
        'Integration with dbt Cloud',
        'Column-level lineage tracking'
    ]
    for s in datafold_strengths:
        doc.add_paragraph(s, style='List Bullet')
    doc.add_paragraph('Our Advantage: DataMigrate AI provides end-to-end automation from MSSQL extraction to dbt model generation, while Datafold only validates migrations you\'ve already built manually.')

    # Mage AI
    doc.add_heading('Mage AI', level=3)
    doc.add_paragraph('Focus: Open-source data pipeline orchestration')
    doc.add_paragraph('Our Advantage: Purpose-built for MSSQL to dbt with intelligent automation vs. manual pipeline building.')

    # dbt Labs + Consultants
    doc.add_heading('dbt Labs + Consulting Partners', level=3)
    doc.add_paragraph('Focus: dbt Cloud professional services')
    doc.add_paragraph('Our Advantage: AI-powered automation at 1/10th the cost with consistent, repeatable results.')

    doc.add_heading('Competitive Feature Matrix', level=2)
    comp_table = doc.add_table(rows=6, cols=4)
    comp_table.style = 'Table Grid'
    comp_headers = ['Feature', 'DataMigrate AI', 'Datafold', 'Consultants']
    for i, header in enumerate(comp_headers):
        cell = comp_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True

    comp_data = [
        ['MSSQL Specialization', '✓ Deep', '✗ Generic', '~ Varies'],
        ['Automated dbt Generation', '✓ Full', '✗ None', '~ Manual'],
        ['Data Validation', '✓ Built-in', '✓ Core Focus', '~ Manual'],
        ['Stored Proc Conversion', '✓ Automated', '✗ None', '~ Manual'],
        ['Cost Efficiency', '$$', '$$$', '$$$$$$']
    ]
    for row_idx, row_data in enumerate(comp_data, 1):
        for col_idx, cell_data in enumerate(row_data):
            comp_table.rows[row_idx].cells[col_idx].text = cell_data

    doc.add_paragraph()
    unique_para = doc.add_paragraph()
    run = unique_para.add_run('Our Unique Position: ')
    run.font.bold = True
    unique_para.add_run('DataMigrate AI is the only solution that combines MSSQL expertise, AI automation, dbt native output, and an eight-agent architecture with RAG, Security, and DataPrep capabilities.')

    doc.add_page_break()

    # ===== AI AGENT SUITE =====
    doc.add_heading('AI Agent Suite (8 Agents)', level=1)

    doc.add_heading('Agent 1: Migration Agent (Core Product)', level=2)
    doc.add_heading('Capabilities', level=3)
    migration_caps = [
        'Automated Schema Analysis - Scans and maps database structure',
        'Intelligent Data Type Mapping - Converts MSSQL types to dbt-compatible formats',
        'Relationship Preservation - Maintains foreign keys and dependencies',
        'dbt Model Generation - Creates production-ready dbt models',
        'Incremental Migration - Supports phased migration approaches',
        'Rollback Support - Safe migration with recovery options'
    ]
    for cap in migration_caps:
        doc.add_paragraph(cap, style='List Bullet')

    doc.add_heading('Agent 2: Customer Support Agent', level=2)
    doc.add_paragraph('AI-powered 24/7 customer support that reduces ticket resolution time and support costs.')
    doc.add_heading('Capabilities', level=3)
    support_caps = [
        'Intelligent Ticket Routing - Automatically categorizes and prioritizes issues',
        'Instant Response - Answers common questions in seconds',
        'Migration Troubleshooting - Diagnoses and resolves migration errors',
        'Guided Walkthroughs - Step-by-step assistance for complex tasks',
        'Escalation Management - Seamlessly transfers to human agents when needed',
        'Knowledge Base Integration - Learns from resolved tickets'
    ]
    for cap in support_caps:
        doc.add_paragraph(cap, style='List Bullet')

    # Support benefits table
    doc.add_paragraph()
    support_table = doc.add_table(rows=5, cols=3)
    support_table.style = 'Table Grid'
    support_headers = ['Metric', 'Without AI Agent', 'With AI Agent']
    for i, header in enumerate(support_headers):
        cell = support_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True

    support_data = [
        ['Average Response Time', '4-8 hours', '< 30 seconds'],
        ['Ticket Resolution', '24-48 hours', '2-4 hours'],
        ['Support Staff Needed', '5-10 agents', '1-2 agents'],
        ['Support Cost/Month', '$15,000+', '$3,000-5,000']
    ]
    for row_idx, row_data in enumerate(support_data, 1):
        for col_idx, cell_data in enumerate(row_data):
            support_table.rows[row_idx].cells[col_idx].text = cell_data

    doc.add_paragraph()

    doc.add_heading('Agent 3: Business Intelligence Agent', level=2)
    doc.add_paragraph('Transform your migrated data into strategic business insights and competitive advantages.')
    doc.add_heading('Capabilities', level=3)
    bi_caps = [
        'Automated Data Analysis - Discovers patterns and trends automatically',
        'Anomaly Detection - Identifies unusual data patterns and potential issues',
        'Predictive Analytics - Forecasts trends based on historical data',
        'Natural Language Queries - Ask questions in plain English',
        'Executive Dashboards - Auto-generated KPI dashboards',
        'Competitive Intelligence - Benchmark analysis and market insights'
    ]
    for cap in bi_caps:
        doc.add_paragraph(cap, style='List Bullet')

    doc.add_paragraph()

    # Agent 4: Data Quality Agent
    doc.add_heading('Agent 4: Data Quality Agent (NEW)', level=2)
    doc.add_paragraph('Ensure data integrity and accuracy throughout the migration process with comprehensive validation and reconciliation.')
    doc.add_heading('Capabilities', level=3)
    dq_caps = [
        'Cross-Database Data Diffing - Compare source MSSQL with target dbt models row-by-row',
        'Schema Validation - Verify all columns, types, and constraints are preserved',
        'Referential Integrity Checks - Ensure all relationships remain intact',
        'Statistical Validation - Compare row counts, sums, averages, and distributions',
        'Anomaly Detection - Identify data quality issues before and after migration',
        'Automated Reconciliation Reports - Generate compliance-ready documentation'
    ]
    for cap in dq_caps:
        doc.add_paragraph(cap, style='List Bullet')

    doc.add_paragraph()

    # Agent 5: Documentation Agent
    doc.add_heading('Agent 5: Documentation Agent with RAG (NEW)', level=2)
    doc.add_paragraph('Automatically generate comprehensive documentation for your migrated dbt project using RAG (Retrieval-Augmented Generation) technology.')
    doc.add_heading('Capabilities', level=3)
    doc_caps = [
        'Auto-Generated Model Documentation - Creates detailed descriptions for every dbt model',
        'Column-Level Documentation - Documents every field with business context',
        'Lineage Documentation - Visual and textual data flow documentation',
        'Business Glossary Generation - Creates consistent terminology definitions',
        'dbt Docs Integration - Seamlessly integrates with dbt\'s documentation system'
    ]
    for cap in doc_caps:
        doc.add_paragraph(cap, style='List Bullet')

    doc.add_paragraph()
    doc.add_heading('RAG Architecture', level=3)
    doc.add_paragraph('Our Documentation Agent is powered by a sophisticated RAG (Retrieval-Augmented Generation) architecture:')

    rag_components = [
        'Vector Database - ChromaDB/Pinecone/Weaviate for semantic document search',
        'Embedding Models - OpenAI Ada-002 or Cohere for document vectorization',
        'LLM Integration - GPT-4/Claude for intelligent response generation',
        'Document Chunking - Smart text splitting with overlap for context preservation',
        'Orchestration - LangChain/LlamaIndex for RAG pipeline management'
    ]
    for comp in rag_components:
        doc.add_paragraph(comp, style='List Bullet')

    doc.add_paragraph()
    doc.add_heading('RAG Technology Stack', level=3)
    rag_table = doc.add_table(rows=5, cols=2)
    rag_table.style = 'Table Grid'
    rag_headers = ['Component', 'Technology']
    for i, header in enumerate(rag_headers):
        cell = rag_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True

    rag_data = [
        ['Vector Store', 'ChromaDB / Pinecone / Weaviate'],
        ['Embeddings', 'OpenAI Ada-002 / Cohere Embed'],
        ['LLM', 'GPT-4 / Claude 3 / LLaMA'],
        ['Orchestration', 'LangChain / LlamaIndex']
    ]
    for row_idx, row_data in enumerate(rag_data, 1):
        for col_idx, cell_data in enumerate(row_data):
            rag_table.rows[row_idx].cells[col_idx].text = cell_data

    doc.add_paragraph()

    # Agent 6: ML Fine-Tuning Agent
    doc.add_heading('Agent 6: ML Fine-Tuning Agent (NEW)', level=2)
    doc.add_paragraph('Enable customers to fine-tune open-source ML models on their own data, adding custom machine learning capabilities to their business.')
    doc.add_heading('Capabilities', level=3)
    ml_caps = [
        'Model Selection - Choose from curated open-source models (LLaMA, Mistral, Falcon, etc.)',
        'Data Preparation - Automated data preprocessing and formatting for fine-tuning',
        'Fine-Tuning Pipeline - End-to-end training with LoRA, QLoRA, and full fine-tuning options',
        'Model Evaluation - Comprehensive benchmarking against your use cases',
        'Deployment - One-click deployment to production environments',
        'Industry Templates - Pre-configured fine-tuning recipes for Financial, Healthcare, Retail, Manufacturing'
    ]
    for cap in ml_caps:
        doc.add_paragraph(cap, style='List Bullet')

    # ML Benefits table
    doc.add_paragraph()
    ml_table = doc.add_table(rows=5, cols=3)
    ml_table.style = 'Table Grid'
    ml_headers = ['Benefit', 'Traditional Approach', 'With ML Agent']
    for i, header in enumerate(ml_headers):
        cell = ml_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True

    ml_data = [
        ['Time to Deploy ML', '3-6 months', '2-4 weeks'],
        ['ML Team Required', '3-5 specialists', '0-1 specialist'],
        ['Infrastructure Cost', '$50K-200K/year', 'Included'],
        ['Model Maintenance', 'Continuous effort', 'Automated']
    ]
    for row_idx, row_data in enumerate(ml_data, 1):
        for col_idx, cell_data in enumerate(row_data):
            ml_table.rows[row_idx].cells[col_idx].text = cell_data

    doc.add_paragraph()

    # Agent 7: Security Agent
    doc.add_heading('Agent 7: Security Agent (NEW)', level=2)
    doc.add_paragraph('Enterprise-grade AI security ensuring data protection, compliance, and threat detection throughout the migration process.')
    doc.add_heading('Capabilities', level=3)
    security_caps = [
        'Data Classification - Automatic PII, PHI, PCI data detection and tagging',
        'SQL Injection Prevention - AI-powered query analysis and sanitization',
        'Compliance Monitoring - GDPR, HIPAA, SOX, PCI-DSS, SOC2, CCPA compliance',
        'Threat Detection - Real-time anomaly detection and security alerts',
        'Access Control - Role-based permissions and audit logging',
        'SIEM Integration - Splunk, DataDog, ELK Stack integration'
    ]
    for cap in security_caps:
        doc.add_paragraph(cap, style='List Bullet')

    doc.add_paragraph()
    doc.add_heading('Compliance Frameworks', level=3)
    compliance_table = doc.add_table(rows=7, cols=2)
    compliance_table.style = 'Table Grid'
    compliance_headers = ['Framework', 'Coverage']
    for i, header in enumerate(compliance_headers):
        cell = compliance_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True

    compliance_data = [
        ['GDPR', 'Data privacy, right to be forgotten, consent management'],
        ['HIPAA', 'PHI protection, audit trails, encryption at rest/transit'],
        ['SOX', 'Financial data integrity, change management, audit logs'],
        ['PCI-DSS', 'Payment card data protection, tokenization, encryption'],
        ['SOC 2', 'Security, availability, processing integrity, confidentiality'],
        ['CCPA', 'California consumer privacy rights, data deletion']
    ]
    for row_idx, row_data in enumerate(compliance_data, 1):
        for col_idx, cell_data in enumerate(row_data):
            compliance_table.rows[row_idx].cells[col_idx].text = cell_data

    doc.add_paragraph()
    doc.add_heading('Security Architecture', level=3)
    security_arch = [
        'Zero Trust Model - Verify every request, assume breach',
        'End-to-End Encryption - TLS 1.3, AES-256 encryption',
        'Secret Management - HashiCorp Vault, AWS Secrets Manager integration',
        'Network Security - VPC isolation, private endpoints, WAF protection'
    ]
    for arch in security_arch:
        doc.add_paragraph(arch, style='List Bullet')

    doc.add_paragraph()

    # Agent 8: DataPrep AI Agent
    doc.add_heading('Agent 8: DataPrep AI Agent (NEW)', level=2)
    doc.add_paragraph('Intelligent data preparation and cleaning for analytics and ML workloads. Available as migration add-on or standalone product.')
    doc.add_heading('Capabilities', level=3)
    dataprep_caps = [
        'Automated Data Profiling - Comprehensive column analysis with statistics and patterns',
        'Intelligent Null Handling - Smart imputation strategies (mean, median, mode, predictive)',
        'Duplicate Detection - Fuzzy matching and deduplication across large datasets',
        'Type Inference - Automatic data type detection and conversion recommendations',
        'Outlier Detection - Multiple methods (IQR, Z-score, Isolation Forest, LOF)',
        'Feature Engineering - Auto-generate derived columns and aggregations',
        'Date/Text Standardization - Normalize formats across data sources',
        'ML-Ready Output - Prepare data for Snowflake, Databricks, BigQuery, Redshift'
    ]
    for cap in dataprep_caps:
        doc.add_paragraph(cap, style='List Bullet')

    # DataPrep Pricing
    doc.add_paragraph()
    doc.add_heading('DataPrep AI Pricing (Hybrid Model)', level=3)
    doc.add_paragraph('Available as migration add-on OR standalone product:')

    dataprep_table = doc.add_table(rows=3, cols=3)
    dataprep_table.style = 'Table Grid'
    dataprep_headers = ['Option', 'Price', 'Best For']
    for i, header in enumerate(dataprep_headers):
        cell = dataprep_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True

    dataprep_data = [
        ['Migration Add-on', '$999/month', 'Customers migrating MSSQL to dbt'],
        ['Standalone Product', '$1,999/month', 'Data teams needing prep without migration']
    ]
    for row_idx, row_data in enumerate(dataprep_data, 1):
        for col_idx, cell_data in enumerate(row_data):
            dataprep_table.rows[row_idx].cells[col_idx].text = cell_data

    doc.add_paragraph()

    doc.add_page_break()

    # ===== TARGET MARKETS =====
    doc.add_heading('Target Market & Customer Niches', level=1)

    markets = [
        {
            'name': 'Enterprise Financial Services',
            'profile': 'Banks, insurance companies, investment firms',
            'pain_points': ['Legacy mainframe databases', 'Regulatory compliance requirements',
                          'Data governance challenges', 'High migration risk tolerance'],
            'value_prop': 'Secure, compliant migration with audit trails',
            'size': '500-10,000+ employees',
            'budget': '$50K-500K annually'
        },
        {
            'name': 'Healthcare & Life Sciences',
            'profile': 'Hospitals, pharmaceutical companies, health insurers',
            'pain_points': ['HIPAA compliance requirements', 'Complex patient data schemas',
                          'Integration with EHR systems', 'Data quality concerns'],
            'value_prop': 'HIPAA-compliant migration with data integrity guarantees',
            'size': '200-5,000+ employees',
            'budget': '$30K-300K annually'
        },
        {
            'name': 'Retail & E-Commerce',
            'profile': 'Retailers, e-commerce platforms, consumer goods',
            'pain_points': ['High transaction volumes', 'Real-time analytics needs',
                          'Multi-channel data integration', 'Seasonal scaling requirements'],
            'value_prop': 'Fast migration with real-time analytics capabilities',
            'size': '100-5,000 employees',
            'budget': '$20K-150K annually'
        },
        {
            'name': 'Manufacturing & Industrial',
            'profile': 'Manufacturers, supply chain companies, industrial firms',
            'pain_points': ['IoT data integration', 'Supply chain visibility',
                          'Legacy ERP systems', 'Operational efficiency'],
            'value_prop': 'Comprehensive ERP and operational data migration',
            'size': '200-10,000 employees',
            'budget': '$30K-200K annually'
        },
        {
            'name': 'Technology & SaaS',
            'profile': 'Software companies, tech startups, SaaS providers',
            'pain_points': ['Rapid scaling requirements', 'Modern architecture needs',
                          'DevOps integration', 'Data-driven decision making'],
            'value_prop': 'Developer-friendly migration with CI/CD integration',
            'size': '50-1,000 employees',
            'budget': '$15K-100K annually'
        },
        {
            'name': 'Government & Public Sector',
            'profile': 'Government agencies, municipalities, public institutions',
            'pain_points': ['Strict compliance requirements', 'Budget constraints',
                          'Legacy system dependencies', 'Security concerns'],
            'value_prop': 'Secure, compliant, cost-effective migration',
            'size': '100-10,000+ employees',
            'budget': '$25K-250K annually'
        }
    ]

    for market in markets:
        doc.add_heading(market['name'], level=2)
        doc.add_paragraph(f"Profile: {market['profile']}")
        doc.add_paragraph('Pain Points:')
        for point in market['pain_points']:
            doc.add_paragraph(point, style='List Bullet')
        doc.add_paragraph(f"Value Proposition: {market['value_prop']}")
        doc.add_paragraph(f"Company Size: {market['size']}")
        doc.add_paragraph(f"Budget Range: {market['budget']}")
        doc.add_paragraph()

    doc.add_page_break()

    # ===== PRICING =====
    doc.add_heading('Pricing Structure', level=1)

    plans = [
        {
            'name': 'Starter Plan - $499/month',
            'best_for': 'Small businesses, startups, pilot projects',
            'features': [
                'Up to 25 database tables',
                '1 MSSQL connection',
                'Migration Agent (basic)',
                'Email support',
                '100GB data transfer/month',
                'Basic documentation generation'
            ]
        },
        {
            'name': 'Professional Plan - $1,499/month',
            'best_for': 'Mid-size companies, growing teams',
            'features': [
                'Up to 100 database tables',
                '5 MSSQL connections',
                'Migration Agent (full)',
                'Customer Support Agent',
                'Priority email & chat support',
                '500GB data transfer/month',
                'Advanced documentation',
                'API access',
                'Slack integration'
            ]
        },
        {
            'name': 'Enterprise Plan - $4,999/month',
            'best_for': 'Large enterprises, complex migrations',
            'features': [
                'Unlimited database tables',
                'Unlimited MSSQL connections',
                'All seven AI Agents (Migration, Support, BI, Data Quality, Documentation, Security, ML)',
                '24/7 dedicated support',
                'Unlimited data transfer',
                'Custom integrations',
                'SSO/SAML authentication',
                'Dedicated success manager',
                'SLA guarantees (99.9% uptime)',
                'On-premise deployment option',
                'Custom training sessions',
                'ML Fine-Tuning credits included'
            ]
        }
    ]

    for plan in plans:
        doc.add_heading(plan['name'], level=2)
        doc.add_paragraph(f"Best for: {plan['best_for']}")
        doc.add_paragraph('Includes:')
        for feature in plan['features']:
            doc.add_paragraph(feature, style='List Bullet')
        doc.add_paragraph()

    doc.add_heading('Volume Discounts', level=2)
    discount_table = doc.add_table(rows=4, cols=2)
    discount_table.style = 'Table Grid'
    discount_headers = ['Annual Commitment', 'Discount']
    for i, header in enumerate(discount_headers):
        cell = discount_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True

    discount_data = [
        ['1 year prepaid', '10% off'],
        ['2 year prepaid', '20% off'],
        ['3 year prepaid', '30% off']
    ]
    for row_idx, row_data in enumerate(discount_data, 1):
        for col_idx, cell_data in enumerate(row_data):
            discount_table.rows[row_idx].cells[col_idx].text = cell_data

    doc.add_page_break()

    # ===== COST REDUCTION BENEFITS =====
    doc.add_heading('Cost Reduction Benefits', level=1)

    doc.add_heading('Direct Cost Savings', level=2)

    doc.add_heading('1. Reduced Migration Consulting Costs', level=3)
    cost_table = doc.add_table(rows=5, cols=2)
    cost_table.style = 'Table Grid'
    cost_headers = ['Approach', 'Cost for 100-Table Migration']
    for i, header in enumerate(cost_headers):
        cell = cost_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True

    cost_data = [
        ['Traditional Consulting', '$150,000 - $300,000'],
        ['In-house Development', '$80,000 - $150,000'],
        ['DataMigrate AI', '$18,000 - $36,000'],
        ['Your Savings', '$62,000 - $264,000']
    ]
    for row_idx, row_data in enumerate(cost_data, 1):
        for col_idx, cell_data in enumerate(row_data):
            cell = cost_table.rows[row_idx].cells[col_idx]
            cell.text = cell_data
            if row_idx == 4:  # Highlight savings row
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(34, 197, 94)  # Green

    doc.add_paragraph()

    doc.add_heading('2. Reduced Labor Costs', level=3)
    doc.add_paragraph(
        'Traditional migrations require 400-800 hours of engineering time. '
        'DataMigrate AI reduces this to 22-44 hours.'
    )
    doc.add_paragraph(
        'At $100/hour average, this represents $37,800 - $75,600 in labor savings.',
        style='List Bullet'
    )

    doc.add_heading('3. Total Cost of Ownership (3-Year)', level=3)
    tco_table = doc.add_table(rows=6, cols=3)
    tco_table.style = 'Table Grid'
    tco_headers = ['Cost Category', 'Traditional', 'DataMigrate AI']
    for i, header in enumerate(tco_headers):
        cell = tco_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True

    tco_data = [
        ['Initial Migration', '$200,000', '$36,000'],
        ['Annual Maintenance', '$60,000/yr', '$18,000/yr'],
        ['Support & Training', '$30,000', '$5,000'],
        ['Error Remediation', '$50,000', '$5,000'],
        ['3-Year Total', '$460,000', '$100,000']
    ]
    for row_idx, row_data in enumerate(tco_data, 1):
        for col_idx, cell_data in enumerate(row_data):
            cell = tco_table.rows[row_idx].cells[col_idx]
            cell.text = cell_data
            if row_idx == 5:  # Highlight total row
                cell.paragraphs[0].runs[0].font.bold = True

    doc.add_paragraph()
    savings_para = doc.add_paragraph()
    run = savings_para.add_run('Total Savings: $360,000 (78%)')
    run.font.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(34, 197, 94)

    doc.add_page_break()

    # ===== VALUE PROPOSITION =====
    doc.add_heading('Sales Benefits & Value Proposition', level=1)

    doc.add_heading('For C-Suite Executives', level=2)

    doc.add_heading('CFO Value Proposition', level=3)
    cfo_points = [
        '70-80% cost reduction vs. traditional migration approaches',
        'Predictable monthly costs with subscription model',
        'Faster ROI - value delivered in weeks, not months',
        'Reduced risk - AI validation catches errors before production'
    ]
    for point in cfo_points:
        doc.add_paragraph(point, style='List Bullet')

    doc.add_heading('CTO Value Proposition', level=3)
    cto_points = [
        'Modern architecture - dbt is the industry standard',
        'Scalable solution - handles enterprise workloads',
        'Integration ready - connects to existing CI/CD pipelines',
        'Future-proof - continuous AI improvements included'
    ]
    for point in cto_points:
        doc.add_paragraph(point, style='List Bullet')

    doc.add_heading('CEO Value Proposition', level=3)
    ceo_points = [
        'Competitive advantage - faster data-driven decisions',
        'Innovation enabler - unlocks modern analytics capabilities',
        'Risk mitigation - reduce dependency on legacy systems',
        'Growth foundation - scalable data infrastructure'
    ]
    for point in ceo_points:
        doc.add_paragraph(point, style='List Bullet')

    doc.add_heading('Key Differentiators', level=2)
    differentiators = [
        ('Speed', 'What takes months with traditional tools takes weeks with DataMigrate AI'),
        ('Accuracy', 'AI-powered validation catches 95% of migration errors before they reach production'),
        ('Cost', '70% lower total cost of ownership compared to consulting-based approaches'),
        ('Support', '24/7 AI-powered support ensures you\'re never blocked'),
        ('Intelligence', 'Not just migration - unlock business insights from your data')
    ]
    for name, desc in differentiators:
        para = doc.add_paragraph()
        run = para.add_run(f'{name}: ')
        run.font.bold = True
        para.add_run(desc)

    doc.add_heading('ROI Summary', level=2)
    roi_table = doc.add_table(rows=3, cols=3)
    roi_table.style = 'Table Grid'
    roi_headers = ['Investment', 'Year 1 Value', '3-Year Value']
    for i, header in enumerate(roi_headers):
        cell = roi_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True

    roi_data = [
        ['Professional Plan ($18K/yr)', '$100K+', '$350K+'],
        ['Enterprise Plan ($60K/yr)', '$300K+', '$1M+']
    ]
    for row_idx, row_data in enumerate(roi_data, 1):
        for col_idx, cell_data in enumerate(row_data):
            roi_table.rows[row_idx].cells[col_idx].text = cell_data

    doc.add_paragraph()
    roi_para = doc.add_paragraph()
    run = roi_para.add_run('Average ROI: 400-600% in Year 1')
    run.font.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(79, 70, 229)

    doc.add_page_break()

    # ===== MVP DEVELOPMENT ROADMAP =====
    doc.add_heading('MVP Development Roadmap', level=1)

    doc.add_heading('Phase 1: Core Platform (Complete)', level=2)
    phase1_items = [
        'MSSQL Connector - Secure database connection with Windows/SQL auth ✓',
        'Schema Extractor - Automated table, column, and relationship discovery ✓',
        'dbt Model Generator - Core conversion of tables to dbt models ✓',
        'Web UI - Migration wizard and dashboard ✓',
        'API Backend - REST API for all operations ✓'
    ]
    for item in phase1_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Phase 2: Competitive Parity Features', level=2)
    doc.add_paragraph('Priority: P0 - Required to Compete with Datafold')
    phase2_items = [
        'Data Validation Engine - Cross-database row-level comparison',
        'Stored Procedure Converter - T-SQL to dbt macro transformation',
        'Reconciliation Reports - Automated pre/post migration comparison',
        'Schema Diff Tool - Visual comparison of source vs target'
    ]
    for item in phase2_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Phase 3: AI Agent Expansion', level=2)
    doc.add_paragraph('Priority: P1 - Competitive Differentiation')
    phase3_items = [
        'Data Quality Agent - Anomaly detection, quality scoring, continuous monitoring',
        'Documentation Agent - RAG-powered docs, business glossary, lineage visualization',
        'Customer Support Agent - Ticket routing, instant responses, troubleshooting'
    ]
    for item in phase3_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Phase 4: ML & Advanced Features', level=2)
    doc.add_paragraph('Priority: P2 - Market Leadership')
    phase4_items = [
        'ML Fine-Tuning Agent - Custom model training on customer data',
        'SSIS Package Importer - Convert SSIS packages to dbt pipelines',
        'Informatica Connector - Import Informatica mappings',
        'Cost Optimization Agent - Query cost analysis, optimization recommendations',
        'Self-Service Portal - White-label for consulting partners'
    ]
    for item in phase4_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Development Milestones', level=2)

    # Milestone table
    milestone_table = doc.add_table(rows=4, cols=3)
    milestone_table.style = 'Table Grid'
    milestone_headers = ['Milestone', 'Goal', 'Key Deliverables']
    for i, header in enumerate(milestone_headers):
        cell = milestone_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True

    milestone_data = [
        ['MVP Launch', 'First paying customers', 'Core migration, validation, reconciliation'],
        ['Enterprise Ready', 'Land enterprise customers', 'All agents, SOC2, SSO/SAML'],
        ['AI Platform', 'Market differentiation', 'ML Fine-Tuning, industry templates']
    ]
    for row_idx, row_data in enumerate(milestone_data, 1):
        for col_idx, cell_data in enumerate(row_data):
            milestone_table.rows[row_idx].cells[col_idx].text = cell_data

    doc.add_page_break()

    # ===== NEXT STEPS =====
    doc.add_heading('Next Steps', level=1)

    doc.add_heading('Ready to Transform Your Data Infrastructure?', level=2)

    doc.add_heading('Option 1: Free Assessment', level=3)
    doc.add_paragraph('Schedule a complimentary migration assessment to understand:')
    assessment_points = [
        'Current database complexity',
        'Estimated migration timeline',
        'Projected cost savings',
        'Recommended plan'
    ]
    for point in assessment_points:
        doc.add_paragraph(point, style='List Bullet')

    doc.add_heading('Option 2: Live Demo', level=3)
    doc.add_paragraph('See DataMigrate AI in action with a personalized demo:')
    demo_points = [
        '30-minute overview',
        'Your use case discussion',
        'Q&A with product experts'
    ]
    for point in demo_points:
        doc.add_paragraph(point, style='List Bullet')

    doc.add_heading('Option 3: Pilot Program', level=3)
    doc.add_paragraph('Start with a low-risk pilot:')
    pilot_points = [
        'Migrate 5-10 tables',
        'Full platform access',
        'Dedicated support',
        'No long-term commitment'
    ]
    for point in pilot_points:
        doc.add_paragraph(point, style='List Bullet')

    doc.add_paragraph()

    # Contact info
    doc.add_heading('Contact Information', level=2)
    contact_info = [
        ('Sales Inquiries:', 'sales@datamigrate.ai'),
        ('Technical Questions:', 'support@datamigrate.ai'),
        ('Partnership Opportunities:', 'partners@datamigrate.ai')
    ]
    for label, email in contact_info:
        para = doc.add_paragraph()
        run = para.add_run(f'{label} ')
        run.font.bold = True
        para.add_run(email)

    doc.add_paragraph()
    doc.add_paragraph()

    # Footer
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run('DataMigrate AI - Intelligent Data Migration for the Modern Enterprise')
    run.font.italic = True
    run.font.color.rgb = RGBColor(107, 114, 128)

    copyright_para = doc.add_paragraph()
    copyright_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = copyright_para.add_run('© 2025 DataMigrate AI. All rights reserved.')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(156, 163, 175)

    # Save the document
    output_path = Path(__file__).parent.parent / 'docs' / 'DATAMIGRATE_AI_SALES_DOCUMENT_v2.docx'
    doc.save(output_path)
    print(f'Word document created: {output_path}')
    return output_path


if __name__ == '__main__':
    create_word_document()
