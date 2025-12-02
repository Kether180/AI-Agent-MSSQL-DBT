"""
Agent Documentation Generator

Generates comprehensive documentation for all agents in PDF and Word formats.
Includes: Guardian Agent, Assessment Agent, Planner Agent, Executor Agent,
Tester Agent, Rebuilder Agent, and Evaluator Agent.

Run: python create_agent_documentation.py
"""

import os
from datetime import datetime

# Check for required packages
try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, ListFlowable, ListItem
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
except ImportError:
    print("Installing reportlab...")
    os.system("pip install reportlab")
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, ListFlowable, ListItem
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("Installing python-docx...")
    os.system("pip install python-docx")
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE


# =============================================================================
# AGENT DOCUMENTATION DATA
# =============================================================================

AGENTS_DATA = {
    "guardian_agent": {
        "name": "Guardian Agent",
        "subtitle": "Security & Protection Layer",
        "description": """The Guardian Agent is a comprehensive security layer designed to protect the
DataMigrate AI platform from various threats including prompt injection attacks, SQL injection,
data exfiltration, and unauthorized access. It operates at two levels: Go Backend (API protection)
and Python (AI agent protection).""",
        "location": {
            "go": "backend/internal/security/",
            "python": "agents/guardian_agent.py"
        },
        "responsibilities": [
            "Monitor and validate all API requests",
            "Detect and block prompt injection attacks",
            "Enforce rate limiting per user/organization",
            "Detect SQL injection and XSS attempts",
            "Audit log all security events",
            "Protect AI agents from malicious inputs",
            "Validate AI agent outputs for data leakage",
            "Enforce per-organization security policies"
        ],
        "benefits": [
            "Zero additional API costs (no AI calls for security)",
            "Deterministic, predictable security checks",
            "Works completely offline",
            "Multi-tenant security isolation",
            "Comprehensive audit trail for compliance",
            "Protects against OWASP Top 10 vulnerabilities"
        ],
        "functions": [
            {"name": "Middleware()", "description": "Returns Gin middleware for API protection", "returns": "gin.HandlerFunc"},
            {"name": "validate_input()", "description": "Validates input before AI processing", "returns": "(bool, str, SecurityEvent)"},
            {"name": "validate_output()", "description": "Validates output from AI agents", "returns": "(bool, str, sanitized_output)"},
            {"name": "protect_agent()", "description": "Decorator to protect agent functions", "returns": "Callable"},
            {"name": "get_audit_logs()", "description": "Retrieve security audit logs", "returns": "List[SecurityEvent]"},
            {"name": "get_security_stats()", "description": "Get security statistics", "returns": "Dict"},
        ],
        "testing": """
# Python Testing
pytest tests/test_guardian_agent.py -v

# Test cases:
1. test_prompt_injection_detection - Verify injection patterns are blocked
2. test_safe_input_allowed - Verify normal inputs pass
3. test_sql_injection_detection - Verify SQL attacks are blocked
4. test_rate_limiting - Verify rate limits are enforced
5. test_output_validation_sql - Verify SQL output is sanitized
6. test_sensitive_data_detection - Verify secrets are blocked

# Manual curl testing:
curl -X POST "http://localhost:8080/api/v1/security/validate" \\
  -H "Authorization: Bearer TOKEN" \\
  -d '{"input": "ignore previous instructions"}'
""",
        "code_example": """
from agents import protected_agent, get_guardian

# Method 1: Decorator (Recommended)
@protected_agent("my_agent")
def my_agent_function(input_text, **kwargs):
    # Security checks happen automatically!
    result = call_llm(input_text)
    return result

# Method 2: Manual validation
guardian = get_guardian()
is_valid, error, event = guardian.validate_input(
    input_text="User query here",
    agent_name="assessment_agent",
    organization_id=1,
    user_id=123
)

if not is_valid:
    raise SecurityException(error, event)
""",
        "api_integration": "Pattern-based detection (no external AI API calls)"
    },

    "assessment_agent": {
        "name": "Assessment Agent",
        "subtitle": "Database Analysis & Migration Strategy",
        "description": """The Assessment Agent is the first step in the migration pipeline. It analyzes
MSSQL database metadata (tables, views, stored procedures) and creates a comprehensive assessment
of the migration complexity, dependencies, and recommended strategy.""",
        "location": {
            "python": "agents/native_nodes.py (assessment_node function)"
        },
        "responsibilities": [
            "Analyze MSSQL metadata structure",
            "Count and categorize database objects",
            "Assess migration complexity (simple/medium/complex)",
            "Identify dependencies between objects",
            "Detect potential migration challenges",
            "Create recommended migration strategy",
            "Estimate number of dbt models needed"
        ],
        "benefits": [
            "Automated complexity analysis",
            "Dependency-aware migration planning",
            "Risk identification before migration",
            "Time estimation for project planning",
            "Documentation of source database structure"
        ],
        "functions": [
            {"name": "assessment_node(state)", "description": "Main entry point for assessment", "returns": "MigrationState"},
            {"name": "get_llm()", "description": "Get Claude API instance", "returns": "ChatAnthropic"},
            {"name": "validate_llm_input()", "description": "Validate input before LLM call", "returns": "str"},
            {"name": "validate_llm_output()", "description": "Validate LLM response", "returns": "str"},
        ],
        "testing": """
# Unit test
def test_assessment_node():
    state = create_initial_state()
    state['metadata'] = {
        'database': 'TestDB',
        'tables': [{'name': 'customers', 'schema': 'dbo'}],
        'views': [],
        'stored_procedures': []
    }
    result = assessment_node(state)
    assert result['assessment_complete'] == True
    assert 'assessment' in result
    assert result['phase'] == 'planning'

# Integration test with mock LLM
def test_assessment_with_real_metadata():
    with open('mssql_metadata.json') as f:
        metadata = json.load(f)
    state = create_initial_state()
    state['metadata'] = metadata
    result = assessment_node(state)
    assert result['assessment']['tables_count'] > 0
""",
        "code_example": """
from agents import assessment_node, create_initial_state

# Create initial state with metadata
state = create_initial_state()
state['metadata'] = {
    'database': 'MyDatabase',
    'tables': [
        {'name': 'customers', 'schema': 'dbo', 'columns': [...]},
        {'name': 'orders', 'schema': 'dbo', 'columns': [...]}
    ],
    'views': [...],
    'stored_procedures': [...]
}

# Run assessment
result = assessment_node(state)

# Access results
print(f"Complexity: {result['assessment']['complexity']}")
print(f"Estimated models: {result['assessment']['estimated_models']}")
print(f"Challenges: {result['assessment']['challenges']}")
""",
        "api_integration": "Claude API (claude-sonnet-4) via LangChain ChatAnthropic"
    },

    "planner_agent": {
        "name": "Planner Agent",
        "subtitle": "dbt Model Planning & Sequencing",
        "description": """The Planner Agent creates a detailed migration plan based on the assessment.
It generates specifications for each dbt model, determines the optimal migration order based on
dependencies, and initializes the model tracking system.""",
        "location": {
            "python": "agents/native_nodes.py (planner_node function)"
        },
        "responsibilities": [
            "Create dbt model specifications for each object",
            "Apply dbt naming conventions (stg_, int_, fct_, dim_)",
            "Determine migration order (dependency-first)",
            "Assign priority levels to models",
            "Initialize model tracking entries",
            "Set up execution pipeline"
        ],
        "benefits": [
            "Consistent dbt naming conventions",
            "Dependency-aware sequencing",
            "Clear migration roadmap",
            "Progress tracking foundation",
            "Parallelization opportunities identified"
        ],
        "functions": [
            {"name": "planner_node(state)", "description": "Main entry point for planning", "returns": "MigrationState"},
            {"name": "validate_llm_input()", "description": "Validate planning prompt", "returns": "str"},
        ],
        "testing": """
# Unit test
def test_planner_node():
    state = create_initial_state()
    state['assessment'] = {
        'tables_count': 5,
        'complexity': 'medium'
    }
    state['metadata'] = {
        'tables': [{'name': f'table_{i}', 'schema': 'dbo'} for i in range(5)]
    }
    result = planner_node(state)
    assert result['plan_complete'] == True
    assert len(result['models']) == 5
    assert result['current_model_index'] == 0
""",
        "code_example": """
from agents import planner_node

# After assessment is complete
state = assessment_result  # From assessment_node

# Run planner
result = planner_node(state)

# Access migration plan
print(f"Total models: {len(result['models'])}")
for model in result['models']:
    print(f"  - {model['name']}: {model['model_type']}")
    print(f"    Source: {model['source_object']}")
    print(f"    Priority: {model['priority']}")
""",
        "api_integration": "Claude API (claude-sonnet-4) via LangChain ChatAnthropic"
    },

    "executor_agent": {
        "name": "Executor Agent",
        "subtitle": "dbt Model Code Generation",
        "description": """The Executor Agent generates the actual dbt model SQL code. It uses the
source schema information to create properly formatted dbt models with correct macros,
transformations, and best practices.""",
        "location": {
            "python": "agents/native_nodes.py (executor_node function)"
        },
        "responsibilities": [
            "Generate dbt model SQL code",
            "Apply {{ source() }} and {{ ref() }} macros correctly",
            "Include appropriate transformations",
            "Follow dbt best practices",
            "Create model files in correct directories",
            "Handle staging/intermediate/marts organization"
        ],
        "benefits": [
            "Automated code generation",
            "Consistent code style",
            "Proper macro usage",
            "Best practices enforcement",
            "File organization automation"
        ],
        "functions": [
            {"name": "executor_node(state)", "description": "Generate dbt model for current model", "returns": "MigrationState"},
            {"name": "sanitize_sql_output()", "description": "Validate generated SQL", "returns": "str"},
            {"name": "get_current_model()", "description": "Get model to process", "returns": "Dict"},
        ],
        "testing": """
# Unit test
def test_executor_node():
    state = create_test_state_with_plan()
    result = executor_node(state)

    current_model = get_current_model(result)
    assert current_model['status'] == 'in_progress'
    assert current_model['file_path'] is not None
    assert Path(current_model['file_path']).exists()

# Verify generated SQL
def test_generated_sql_quality():
    state = executor_node(create_test_state())
    file_path = get_current_model(state)['file_path']

    with open(file_path) as f:
        sql = f.read()

    assert 'SELECT' in sql
    assert 'source(' in sql or 'ref(' in sql
""",
        "code_example": """
from agents import executor_node, get_current_model
from pathlib import Path

# After planning is complete
state = planner_result  # From planner_node

# Execute current model
result = executor_node(state)

# Check generated file
model = get_current_model(result)
print(f"Generated: {model['file_path']}")

with open(model['file_path']) as f:
    print(f.read())

# Output example:
# {{ config(materialized='view') }}
#
# SELECT
#     customer_id,
#     customer_name,
#     created_at
# FROM {{ source('mssql', 'customers') }}
""",
        "api_integration": "Claude API (claude-sonnet-4) via LangChain ChatAnthropic"
    },

    "tester_agent": {
        "name": "Tester Agent",
        "subtitle": "dbt Model Validation & QA",
        "description": """The Tester Agent validates generated dbt models against quality standards.
It checks for correct macro usage, valid SQL syntax, best practices compliance, and provides
a validation score to determine if the model passes or needs rebuilding.""",
        "location": {
            "python": "agents/native_nodes.py (tester_node function)"
        },
        "responsibilities": [
            "Validate dbt model SQL syntax",
            "Check {{ source() }} and {{ ref() }} usage",
            "Verify transformation appropriateness",
            "Score model quality (0.0 - 1.0)",
            "Identify issues and recommendations",
            "Mark models as completed or failed"
        ],
        "benefits": [
            "Automated quality assurance",
            "Consistent validation standards",
            "Early issue detection",
            "Quality metrics tracking",
            "Rebuild triggering for failures"
        ],
        "functions": [
            {"name": "tester_node(state)", "description": "Validate current model", "returns": "MigrationState"},
            {"name": "validate_llm_input()", "description": "Prepare validation prompt", "returns": "str"},
            {"name": "validate_llm_output()", "description": "Parse validation result", "returns": "str"},
        ],
        "testing": """
# Unit test
def test_tester_node_pass():
    state = create_state_with_valid_model()
    result = tester_node(state)

    model = get_current_model(result)
    assert model['status'] == 'completed'
    assert model['validation_score'] >= 0.7

def test_tester_node_fail():
    state = create_state_with_invalid_model()
    result = tester_node(state)

    model = get_current_model(result)
    assert model['status'] == 'failed'
    assert len(model['errors']) > 0
""",
        "code_example": """
from agents import tester_node, get_current_model

# After executor has run
state = executor_result  # From executor_node

# Run validation
result = tester_node(state)

# Check results
model = get_current_model(result)
if model['status'] == 'completed':
    print(f"Model passed! Score: {model['validation_score']}")
else:
    print(f"Model failed. Issues:")
    for error in model['errors']:
        print(f"  - {error}")
""",
        "api_integration": "Claude API (claude-sonnet-4) via LangChain ChatAnthropic"
    },

    "rebuilder_agent": {
        "name": "Rebuilder Agent",
        "subtitle": "Failed Model Recovery & Repair",
        "description": """The Rebuilder Agent attempts to fix models that failed validation.
It analyzes the errors from previous attempts and generates corrected SQL code.
It implements retry logic with a maximum attempt limit.""",
        "location": {
            "python": "agents/native_nodes.py (rebuilder_node function)"
        },
        "responsibilities": [
            "Analyze validation errors",
            "Generate corrected SQL code",
            "Implement retry logic (max 3 attempts)",
            "Learn from previous failures",
            "Reset model status for retry",
            "Track rebuild attempts"
        ],
        "benefits": [
            "Automatic error recovery",
            "Reduced manual intervention",
            "Learning from mistakes",
            "Configurable retry limits",
            "Improved success rates"
        ],
        "functions": [
            {"name": "rebuilder_node(state)", "description": "Attempt to fix failed model", "returns": "MigrationState"},
            {"name": "sanitize_sql_output()", "description": "Validate corrected SQL", "returns": "str"},
        ],
        "testing": """
# Unit test
def test_rebuilder_node():
    state = create_state_with_failed_model()
    state['max_retries'] = 3

    result = rebuilder_node(state)
    model = get_current_model(result)

    assert model['status'] == 'pending'  # Reset for retry
    assert len(model['errors']) == 0  # Errors cleared

def test_rebuilder_max_retries():
    state = create_state_with_failed_model()
    state['models'][0]['attempts'] = 3

    result = rebuilder_node(state)
    model = get_current_model(result)

    # Should not reset if max retries exceeded
    assert model['status'] == 'failed'
""",
        "code_example": """
from agents import rebuilder_node, get_current_model

# After tester marks model as failed
state = tester_result  # From tester_node with failed model

# Attempt rebuild
if get_current_model(state)['status'] == 'failed':
    result = rebuilder_node(state)

    model = get_current_model(result)
    if model['status'] == 'pending':
        print("Model rebuilt, ready for retry")
        # Will go back through executor -> tester
    else:
        print(f"Max retries exceeded ({model['attempts']} attempts)")
""",
        "api_integration": "Claude API (claude-sonnet-4) via LangChain ChatAnthropic"
    },

    "evaluator_agent": {
        "name": "Evaluator Agent",
        "subtitle": "Final Reporting & Summary",
        "description": """The Evaluator Agent runs at the end of the migration pipeline.
It calculates success rates, gathers statistics, and produces a comprehensive
final report of the migration process.""",
        "location": {
            "python": "agents/native_nodes.py (evaluator_node function)"
        },
        "responsibilities": [
            "Calculate migration success rate",
            "Gather completion statistics",
            "List completed and failed models",
            "Produce final evaluation report",
            "Set migration phase to 'complete'",
            "Log summary information"
        ],
        "benefits": [
            "Clear migration summary",
            "Success metrics tracking",
            "Failed model documentation",
            "Actionable insights",
            "Audit trail completion"
        ],
        "functions": [
            {"name": "evaluator_node(state)", "description": "Generate final evaluation", "returns": "MigrationState"},
        ],
        "testing": """
# Unit test
def test_evaluator_node():
    state = create_completed_migration_state()
    state['completed_count'] = 8
    state['failed_count'] = 2

    result = evaluator_node(state)

    assert result['phase'] == 'complete'
    assert result['evaluation']['success_rate'] == 80.0
    assert len(result['evaluation']['completed_models']) == 8
    assert len(result['evaluation']['failed_models']) == 2
""",
        "code_example": """
from agents import evaluator_node

# After all models processed
state = final_state  # After all executor/tester cycles

# Generate evaluation
result = evaluator_node(state)

# Access final report
eval = result['evaluation']
print("=" * 60)
print("MIGRATION COMPLETE")
print("=" * 60)
print(f"Total Models: {eval['total_models']}")
print(f"Completed: {eval['completed']}")
print(f"Failed: {eval['failed']}")
print(f"Success Rate: {eval['success_rate']:.1f}%")

print("\\nCompleted Models:")
for name in eval['completed_models']:
    print(f"  ✓ {name}")

print("\\nFailed Models:")
for item in eval['failed_models']:
    print(f"  ✗ {item['name']}: {', '.join(item['errors'])}")
""",
        "api_integration": "No API calls (pure computation)"
    }
}


# =============================================================================
# PDF GENERATION
# =============================================================================

def create_pdf_documentation(agents_data: dict, output_dir: str = "docs"):
    """Generate PDF documentation for all agents"""

    os.makedirs(output_dir, exist_ok=True)

    for agent_key, agent in agents_data.items():
        filename = f"{output_dir}/{agent_key}_documentation.pdf"
        doc = SimpleDocTemplate(filename, pagesize=A4,
                               rightMargin=72, leftMargin=72,
                               topMargin=72, bottomMargin=72)

        # Styles
        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(
            name='Title_Custom',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=30,
            alignment=TA_CENTER,
            textColor=colors.HexColor('#1a365d')
        ))
        styles.add(ParagraphStyle(
            name='Subtitle_Custom',
            parent=styles['Normal'],
            fontSize=14,
            spaceAfter=20,
            alignment=TA_CENTER,
            textColor=colors.HexColor('#4a5568')
        ))
        styles.add(ParagraphStyle(
            name='Section',
            parent=styles['Heading2'],
            fontSize=16,
            spaceBefore=20,
            spaceAfter=10,
            textColor=colors.HexColor('#2d3748')
        ))
        styles.add(ParagraphStyle(
            name='Body_Custom',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=12,
            alignment=TA_JUSTIFY,
            leading=14
        ))
        styles.add(ParagraphStyle(
            name='CodeBlock',
            parent=styles['Normal'],
            fontName='Courier',
            fontSize=9,
            backColor=colors.HexColor('#f7fafc'),
            borderPadding=10,
            leftIndent=20,
            spaceAfter=12
        ))

        story = []

        # Title
        story.append(Paragraph(agent['name'], styles['Title_Custom']))
        story.append(Paragraph(agent['subtitle'], styles['Subtitle_Custom']))
        story.append(Spacer(1, 20))

        # Overview
        story.append(Paragraph("Overview", styles['Section']))
        story.append(Paragraph(agent['description'], styles['Body_Custom']))
        story.append(Spacer(1, 10))

        # Location
        story.append(Paragraph("File Location", styles['Section']))
        for lang, path in agent['location'].items():
            story.append(Paragraph(f"<b>{lang.upper()}:</b> {path}", styles['Body_Custom']))
        story.append(Spacer(1, 10))

        # Responsibilities
        story.append(Paragraph("Responsibilities", styles['Section']))
        for resp in agent['responsibilities']:
            story.append(Paragraph(f"• {resp}", styles['Body_Custom']))
        story.append(Spacer(1, 10))

        # Benefits
        story.append(Paragraph("Benefits", styles['Section']))
        for benefit in agent['benefits']:
            story.append(Paragraph(f"✓ {benefit}", styles['Body_Custom']))
        story.append(Spacer(1, 10))

        # Functions
        story.append(Paragraph("Key Functions", styles['Section']))
        func_data = [['Function', 'Description', 'Returns']]
        for func in agent['functions']:
            func_data.append([func['name'], func['description'], func['returns']])

        func_table = Table(func_data, colWidths=[2*inch, 3*inch, 1.5*inch])
        func_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4a5568')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#f7fafc')),
            ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#e2e8f0')),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ]))
        story.append(func_table)
        story.append(Spacer(1, 15))

        # API Integration
        story.append(Paragraph("API Integration", styles['Section']))
        story.append(Paragraph(agent['api_integration'], styles['Body_Custom']))
        story.append(Spacer(1, 10))

        # Code Example
        story.append(PageBreak())
        story.append(Paragraph("Code Example", styles['Section']))
        for line in agent['code_example'].strip().split('\n'):
            story.append(Paragraph(line.replace('<', '&lt;').replace('>', '&gt;'), styles['CodeBlock']))
        story.append(Spacer(1, 15))

        # Testing
        story.append(Paragraph("Testing", styles['Section']))
        for line in agent['testing'].strip().split('\n'):
            story.append(Paragraph(line.replace('<', '&lt;').replace('>', '&gt;'), styles['CodeBlock']))

        # Footer
        story.append(Spacer(1, 30))
        story.append(Paragraph(
            f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')} | DataMigrate AI",
            ParagraphStyle(name='Footer', fontSize=9, textColor=colors.grey, alignment=TA_CENTER)
        ))

        doc.build(story)
        print(f"[OK] Created: {filename}")


# =============================================================================
# WORD GENERATION
# =============================================================================

def create_word_documentation(agents_data: dict, output_dir: str = "docs"):
    """Generate Word documentation for all agents"""

    os.makedirs(output_dir, exist_ok=True)

    for agent_key, agent in agents_data.items():
        filename = f"{output_dir}/{agent_key}_documentation.docx"
        doc = Document()

        # Title
        title = doc.add_heading(agent['name'], 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        subtitle = doc.add_paragraph(agent['subtitle'])
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].italic = True

        doc.add_paragraph()

        # Overview
        doc.add_heading('Overview', level=1)
        doc.add_paragraph(agent['description'])

        # Location
        doc.add_heading('File Location', level=1)
        for lang, path in agent['location'].items():
            p = doc.add_paragraph()
            p.add_run(f"{lang.upper()}: ").bold = True
            p.add_run(path)

        # Responsibilities
        doc.add_heading('Responsibilities', level=1)
        for resp in agent['responsibilities']:
            doc.add_paragraph(resp, style='List Bullet')

        # Benefits
        doc.add_heading('Benefits', level=1)
        for benefit in agent['benefits']:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run('✓ ')
            p.add_run(benefit)

        # Functions Table
        doc.add_heading('Key Functions', level=1)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'

        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Function'
        header_cells[1].text = 'Description'
        header_cells[2].text = 'Returns'

        for func in agent['functions']:
            row = table.add_row().cells
            row[0].text = func['name']
            row[1].text = func['description']
            row[2].text = func['returns']

        doc.add_paragraph()

        # API Integration
        doc.add_heading('API Integration', level=1)
        doc.add_paragraph(agent['api_integration'])

        # Code Example
        doc.add_page_break()
        doc.add_heading('Code Example', level=1)
        code_para = doc.add_paragraph()
        code_run = code_para.add_run(agent['code_example'].strip())
        code_run.font.name = 'Consolas'
        code_run.font.size = Pt(9)

        # Testing
        doc.add_heading('Testing', level=1)
        test_para = doc.add_paragraph()
        test_run = test_para.add_run(agent['testing'].strip())
        test_run.font.name = 'Consolas'
        test_run.font.size = Pt(9)

        # Footer
        doc.add_paragraph()
        footer = doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')} | DataMigrate AI")
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.runs[0].font.size = Pt(9)
        footer.runs[0].font.color.rgb = RGBColor(128, 128, 128)

        doc.save(filename)
        print(f"[OK] Created: {filename}")


# =============================================================================
# COMBINED DOCUMENTATION
# =============================================================================

def create_combined_pdf(agents_data: dict, output_dir: str = "docs"):
    """Generate combined PDF with all agents"""

    os.makedirs(output_dir, exist_ok=True)
    filename = f"{output_dir}/ALL_AGENTS_COMPLETE_DOCUMENTATION.pdf"

    doc = SimpleDocTemplate(filename, pagesize=A4,
                           rightMargin=72, leftMargin=72,
                           topMargin=72, bottomMargin=72)

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='MainTitle', parent=styles['Heading1'],
                             fontSize=28, spaceAfter=40, alignment=TA_CENTER,
                             textColor=colors.HexColor('#1a365d')))
    styles.add(ParagraphStyle(name='AgentTitle', parent=styles['Heading1'],
                             fontSize=22, spaceBefore=30, spaceAfter=20,
                             textColor=colors.HexColor('#2d3748')))
    styles.add(ParagraphStyle(name='Section', parent=styles['Heading2'],
                             fontSize=14, spaceBefore=15, spaceAfter=8,
                             textColor=colors.HexColor('#4a5568')))
    styles.add(ParagraphStyle(name='Body', parent=styles['Normal'],
                             fontSize=10, spaceAfter=8, leading=14))
    styles.add(ParagraphStyle(name='CodeBlock', fontName='Courier', fontSize=8,
                             backColor=colors.HexColor('#f7fafc'),
                             leftIndent=10, spaceAfter=8))

    story = []

    # Cover page
    story.append(Spacer(1, 100))
    story.append(Paragraph("DataMigrate AI", styles['MainTitle']))
    story.append(Paragraph("Agent System Documentation",
                          ParagraphStyle(name='Sub', fontSize=18,
                                        alignment=TA_CENTER,
                                        textColor=colors.HexColor('#4a5568'))))
    story.append(Spacer(1, 50))
    story.append(Paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}",
                          ParagraphStyle(name='Date', fontSize=12,
                                        alignment=TA_CENTER)))
    story.append(PageBreak())

    # Table of Contents
    story.append(Paragraph("Table of Contents", styles['AgentTitle']))
    for i, (key, agent) in enumerate(agents_data.items(), 1):
        story.append(Paragraph(f"{i}. {agent['name']} - {agent['subtitle']}", styles['Body']))
    story.append(PageBreak())

    # Each agent
    for agent in agents_data.values():
        story.append(Paragraph(agent['name'], styles['AgentTitle']))
        story.append(Paragraph(agent['subtitle'],
                              ParagraphStyle(name='AgentSub', fontSize=12,
                                            textColor=colors.grey, spaceAfter=15)))

        story.append(Paragraph("Overview", styles['Section']))
        story.append(Paragraph(agent['description'], styles['Body']))

        story.append(Paragraph("Responsibilities", styles['Section']))
        for r in agent['responsibilities']:
            story.append(Paragraph(f"• {r}", styles['Body']))

        story.append(Paragraph("Benefits", styles['Section']))
        for b in agent['benefits']:
            story.append(Paragraph(f"- {b}", styles['Body']))

        story.append(Paragraph("API Integration", styles['Section']))
        story.append(Paragraph(agent['api_integration'], styles['Body']))

        story.append(PageBreak())

    doc.build(story)
    print(f"[OK] Created: {filename}")


def create_combined_word(agents_data: dict, output_dir: str = "docs"):
    """Generate combined Word document with all agents"""

    os.makedirs(output_dir, exist_ok=True)
    filename = f"{output_dir}/ALL_AGENTS_COMPLETE_DOCUMENTATION.docx"

    doc = Document()

    # Cover
    doc.add_heading('DataMigrate AI', 0)
    subtitle = doc.add_paragraph('Agent System Documentation')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    date_para = doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # TOC
    doc.add_heading('Table of Contents', level=1)
    for i, agent in enumerate(agents_data.values(), 1):
        doc.add_paragraph(f"{i}. {agent['name']} - {agent['subtitle']}")

    doc.add_page_break()

    # Each agent
    for agent in agents_data.values():
        doc.add_heading(agent['name'], level=1)
        sub = doc.add_paragraph(agent['subtitle'])
        sub.runs[0].italic = True

        doc.add_heading('Overview', level=2)
        doc.add_paragraph(agent['description'])

        doc.add_heading('Responsibilities', level=2)
        for r in agent['responsibilities']:
            doc.add_paragraph(r, style='List Bullet')

        doc.add_heading('Benefits', level=2)
        for b in agent['benefits']:
            doc.add_paragraph(f"✓ {b}", style='List Bullet')

        doc.add_heading('API Integration', level=2)
        doc.add_paragraph(agent['api_integration'])

        doc.add_page_break()

    doc.save(filename)
    print(f"[OK] Created: {filename}")


# =============================================================================
# MAIN
# =============================================================================

if __name__ == "__main__":
    print("\n" + "=" * 60)
    print("DataMigrate AI - Agent Documentation Generator")
    print("=" * 60 + "\n")

    output_dir = "docs"

    print("Generating individual agent documentation...\n")

    # Individual PDFs
    print("Creating PDF files...")
    create_pdf_documentation(AGENTS_DATA, output_dir)

    print("\nCreating Word files...")
    create_word_documentation(AGENTS_DATA, output_dir)

    print("\nCreating combined documentation...")
    create_combined_pdf(AGENTS_DATA, output_dir)
    create_combined_word(AGENTS_DATA, output_dir)

    print("\n" + "=" * 60)
    print("Documentation generation complete!")
    print(f"Files saved to: {os.path.abspath(output_dir)}/")
    print("=" * 60 + "\n")

    print("Generated files:")
    for f in os.listdir(output_dir):
        if f.endswith(('.pdf', '.docx')):
            print(f"  - {f}")
