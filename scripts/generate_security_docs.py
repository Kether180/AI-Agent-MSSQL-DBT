"""
Script to generate Security Documentation in Word (.docx) and PDF formats
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from datetime import datetime

def set_cell_shading(cell, fill_color):
    """Set cell background color"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)

def add_table_with_header(doc, headers, rows, header_color='2E5090'):
    """Add a formatted table with colored header"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'

    # Header row
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        header_cells[idx].text = header
        set_cell_shading(header_cells[idx], header_color)
        for paragraph in header_cells[idx].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(10)

    # Data rows
    for row_data in rows:
        row = table.add_row()
        for idx, cell_text in enumerate(row_data):
            row.cells[idx].text = str(cell_text)
            for paragraph in row.cells[idx].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    return table

def create_security_document():
    doc = Document()

    # Set up styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Title
    title = doc.add_heading('DataMigrate AI', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading('Security Documentation', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Document info
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run(f'Version: 1.0 | Date: {datetime.now().strftime("%B %Y")} | Classification: Confidential').italic = True

    doc.add_paragraph()
    doc.add_page_break()

    # Table of Contents
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        '1. Executive Summary',
        '2. Security Architecture Overview',
        '3. Authentication & Authorization',
        '4. Data Encryption',
        '5. Input Validation & Sanitization',
        '6. Rate Limiting & DDoS Protection',
        '7. Audit Logging & Monitoring',
        '8. Network Security',
        '9. Compliance & Standards',
        '10. Risk Assessment',
        '11. Incident Response',
        '12. Security Best Practices for Clients'
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Number')

    doc.add_page_break()

    # 1. Executive Summary
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        'DataMigrate AI is designed with security as a foundational principle. Our platform facilitates '
        'the migration of data from Microsoft SQL Server to modern data warehouse platforms (Snowflake, '
        'BigQuery, Databricks, etc.) with comprehensive security measures to protect sensitive business '
        'data throughout the migration lifecycle.'
    )

    doc.add_heading('Key Security Highlights', level=2)
    highlights = [
        'AES-256-GCM Encryption for database credentials at rest',
        'Bcrypt Password Hashing for user authentication with configurable cost factor',
        'JWT-Based Authentication with configurable expiration and HMAC-SHA256 signing',
        'Multi-Layer Security Agent (Guardian) with real-time threat detection',
        'Comprehensive Audit Logging for all security-relevant operations',
        'Rate Limiting with sliding window algorithm to prevent abuse',
        'Input Validation with pattern detection for SQL injection, XSS, and prompt injection attacks'
    ]
    for h in highlights:
        doc.add_paragraph(h, style='List Bullet')

    doc.add_page_break()

    # 2. Security Architecture
    doc.add_heading('2. Security Architecture Overview', level=1)

    doc.add_heading('2.1 Defense in Depth', level=2)
    doc.add_paragraph(
        'DataMigrate AI implements a multi-layered security architecture with five distinct security layers:'
    )

    layers = [
        ('Layer 1: Network Security', 'CORS Policy, HTTPS/TLS Encryption, Rate Limiting'),
        ('Layer 2: Guardian Security Agent', 'Request validation, Pattern detection, Real-time threat blocking'),
        ('Layer 3: Authentication & Authorization', 'JWT Token validation, Role-based access control, Organization isolation'),
        ('Layer 4: Application Security', 'Input validation, Parameterized queries, Secure credential handling'),
        ('Layer 5: Data Security', 'AES-256-GCM encryption, Bcrypt hashing, Audit logging')
    ]
    add_table_with_header(doc, ['Security Layer', 'Components'], layers)

    doc.add_paragraph()

    doc.add_heading('2.2 Security Components', level=2)
    components = [
        ('Encryption Service', 'AES-256-GCM', 'Encrypts database credentials at rest'),
        ('Password Hashing', 'Bcrypt (cost factor 10)', 'Secure user password storage'),
        ('Authentication', 'JWT (HS256)', 'Stateless session management'),
        ('Security Agent', 'Guardian', 'Real-time threat detection and blocking'),
        ('Pattern Detector', 'Regex-based', 'SQL injection, XSS, command injection detection'),
        ('Rate Limiter', 'Sliding Window', 'DDoS and abuse prevention'),
        ('Audit Logger', 'PostgreSQL-backed', 'Security event tracking')
    ]
    add_table_with_header(doc, ['Component', 'Technology', 'Purpose'], components)

    doc.add_page_break()

    # 3. Authentication & Authorization
    doc.add_heading('3. Authentication & Authorization', level=1)

    doc.add_heading('3.1 User Authentication', level=2)
    doc.add_paragraph('Password Storage:')
    doc.add_paragraph('All user passwords are hashed using bcrypt with a cost factor of 10. Passwords are never stored in plaintext. Password hashing uses cryptographically secure random salt.', style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('JWT Token Security:')
    jwt_features = [
        'Algorithm: HMAC-SHA256 (HS256)',
        'Configurable expiration (default: 24 hours)',
        'Secret key stored in environment variable',
        'Token transmitted via Authorization header with Bearer scheme'
    ]
    for f in jwt_features:
        doc.add_paragraph(f, style='List Bullet')

    doc.add_heading('3.2 Authorization Model', level=2)
    doc.add_paragraph('Role-Based Access Control (RBAC):')
    roles = [
        ('User', 'CRUD on own resources, view own organization data'),
        ('Admin', 'All user permissions + user management, security logs'),
        ('Super Admin', 'Platform-wide administration')
    ]
    add_table_with_header(doc, ['Role', 'Permissions'], roles)

    doc.add_paragraph()
    doc.add_paragraph('Organization Isolation:', style='Heading 4')
    doc.add_paragraph('All data is scoped to the user\'s organization. Database queries include organization_id filtering. Cross-organization data access is blocked at the application layer.')

    doc.add_heading('3.3 Password Reset Flow', level=2)
    reset_steps = [
        'User requests password reset via email',
        'Cryptographically secure 32-byte token generated using crypto/rand',
        'Token stored with 1-hour expiration',
        'Token can only be used once (marked as used after reset)',
        'Previous tokens for user are invalidated upon new request'
    ]
    for idx, step in enumerate(reset_steps, 1):
        doc.add_paragraph(f'{idx}. {step}')

    doc.add_page_break()

    # 4. Data Encryption
    doc.add_heading('4. Data Encryption', level=1)

    doc.add_heading('4.1 Encryption at Rest', level=2)
    doc.add_paragraph('Database Credentials Encryption:')
    doc.add_paragraph('DataMigrate AI uses AES-256-GCM (Galois/Counter Mode) for encrypting stored database connection passwords.')

    doc.add_paragraph()
    doc.add_paragraph('Technical Specifications:')
    specs = [
        ('Algorithm', 'AES-256-GCM'),
        ('Key Size', '256 bits (32 bytes)'),
        ('Nonce Size', '96 bits (12 bytes)'),
        ('Authentication Tag', '128 bits (16 bytes)'),
        ('Key Storage', 'Environment variable (base64-encoded)')
    ]
    add_table_with_header(doc, ['Parameter', 'Value'], specs)

    doc.add_paragraph()
    doc.add_paragraph('Encryption Process:')
    enc_steps = [
        'Generate random 12-byte nonce using crypto/rand',
        'Encrypt plaintext with AES-256-GCM',
        'Prepend nonce to ciphertext',
        'Encode result as base64 for storage'
    ]
    for idx, step in enumerate(enc_steps, 1):
        doc.add_paragraph(f'{idx}. {step}')

    doc.add_heading('4.2 Encryption in Transit', level=2)
    doc.add_paragraph('TLS/HTTPS:')
    tls_features = [
        'All API endpoints should be served over HTTPS',
        'TLS 1.2+ recommended',
        'Strong cipher suites enforced'
    ]
    for f in tls_features:
        doc.add_paragraph(f, style='List Bullet')

    doc.add_heading('4.3 Sensitive Data Handling', level=2)
    doc.add_paragraph('Data Never Logged:')
    never_logged = ['Passwords', 'API keys', 'Database connection strings', 'JWT tokens', 'Credit card information']
    for item in never_logged:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # 5. Input Validation
    doc.add_heading('5. Input Validation & Sanitization', level=1)

    doc.add_heading('5.1 Guardian Security Agent', level=2)
    doc.add_paragraph('The Guardian Agent provides centralized security middleware that inspects all incoming requests.')

    doc.add_paragraph('Request Validation:')
    validation_items = [
        'Maximum body size: 1MB',
        'Content-Type validation',
        'Query parameter inspection',
        'JSON schema validation'
    ]
    for item in validation_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('5.2 Pattern Detection', level=2)

    doc.add_paragraph('SQL Injection Protection:', style='Heading 4')
    sql_patterns = [
        ('SQL Statements', 'SELECT...FROM, INSERT INTO, DROP TABLE'),
        ('UNION Injection', 'UNION SELECT'),
        ('Boolean Injection', 'OR 1=1, AND 1=1'),
        ('Comment Injection', '--, /* */'),
        ('Time-Based', 'SLEEP(), WAITFOR DELAY, BENCHMARK()'),
        ('System Access', 'xp_cmdshell, INFORMATION_SCHEMA')
    ]
    add_table_with_header(doc, ['Pattern Type', 'Examples Blocked'], sql_patterns)

    doc.add_paragraph()
    doc.add_paragraph('XSS Protection:', style='Heading 4')
    xss_patterns = [
        ('Script Tags', '<script>, <script src=...>'),
        ('Event Handlers', 'onclick=, onload=, onerror='),
        ('JavaScript Protocol', 'javascript:'),
        ('Dangerous Elements', '<iframe>, <embed>, <object>'),
        ('Data URIs', 'data:text/html'),
        ('SVG Attacks', '<svg onload=...>')
    ]
    add_table_with_header(doc, ['Pattern Type', 'Examples Blocked'], xss_patterns)

    doc.add_paragraph()
    doc.add_paragraph('Prompt Injection Protection (AI-Specific):', style='Heading 4')
    prompt_patterns = [
        ('Instruction Override', 'ignore previous instructions'),
        ('Role Change', 'you are now, act as if you'),
        ('System Reveal', 'reveal your system prompt'),
        ('Marker Injection', 'HUMAN:, ASSISTANT:, <|im_start|>')
    ]
    add_table_with_header(doc, ['Pattern Type', 'Examples Blocked'], prompt_patterns)

    doc.add_page_break()

    # 6. Rate Limiting
    doc.add_heading('6. Rate Limiting & DDoS Protection', level=1)

    doc.add_heading('6.1 Rate Limiting Configuration', level=2)
    doc.add_paragraph('Default Limits:')
    limits = [
        ('Per Second (Burst)', '50 requests', 'Prevent rapid-fire attacks'),
        ('Per Minute', '300 requests', 'Normal operation limit'),
        ('Per Hour', '3,000 requests', 'Sustained abuse prevention')
    ]
    add_table_with_header(doc, ['Window', 'Limit', 'Purpose'], limits)

    doc.add_paragraph()
    doc.add_paragraph('Blocking Behavior:')
    blocking = [
        '3 burst violations result in 1 minute block',
        '5 per-minute violations result in 1 minute block',
        'Automatic cleanup of stale entries every 10 minutes'
    ]
    for b in blocking:
        doc.add_paragraph(b, style='List Bullet')

    doc.add_heading('6.2 Rate Limit Algorithm', level=2)
    doc.add_paragraph('DataMigrate AI uses a sliding window rate limiting algorithm:')
    algo_steps = [
        'Track request timestamps per IP+endpoint combination',
        'Count requests in rolling time windows',
        'Apply progressively stricter limits',
        'Block repeat offenders temporarily'
    ]
    for idx, step in enumerate(algo_steps, 1):
        doc.add_paragraph(f'{idx}. {step}')

    doc.add_page_break()

    # 7. Audit Logging
    doc.add_heading('7. Audit Logging & Monitoring', level=1)

    doc.add_heading('7.1 Events Logged', level=2)
    doc.add_paragraph('Security Events:')
    sec_events = [
        'Authentication attempts (success/failure)',
        'Password changes',
        'Rate limit violations',
        'Suspicious pattern detections',
        'Blocked requests'
    ]
    for e in sec_events:
        doc.add_paragraph(e, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('Audit Events:')
    audit_events = [
        'Resource creation/modification/deletion',
        'Database connection tests',
        'Migration operations',
        'API key management'
    ]
    for e in audit_events:
        doc.add_paragraph(e, style='List Bullet')

    doc.add_heading('7.2 Severity Levels', level=2)
    severities = [
        ('Critical', 'Immediate threat', 'Prompt injection, command injection'),
        ('High', 'Serious security issue', 'SQL injection, XSS attempts'),
        ('Warning', 'Potential issue', 'Rate limit exceeded, failed auth'),
        ('Info', 'Normal operation', 'Successful requests')
    ]
    add_table_with_header(doc, ['Level', 'Description', 'Examples'], severities)

    doc.add_heading('7.3 Monitoring Endpoints', level=2)
    endpoints = [
        'GET /api/v1/security/audit-logs - Retrieve audit logs',
        'GET /api/v1/security/stats - Security statistics',
        'GET /api/v1/security/dashboard - Security dashboard data',
        'GET /metrics - Prometheus metrics endpoint'
    ]
    for e in endpoints:
        doc.add_paragraph(e, style='List Bullet')

    doc.add_page_break()

    # 8. Network Security
    doc.add_heading('8. Network Security', level=1)

    doc.add_heading('8.1 CORS Configuration', level=2)
    doc.add_paragraph('Configurable Allowed Origins:')
    cors_items = [
        'Origins specified via ALLOWED_ORIGINS environment variable',
        'Comma-separated list of allowed domains',
        'Only listed origins can make cross-origin requests'
    ]
    for c in cors_items:
        doc.add_paragraph(c, style='List Bullet')

    doc.add_heading('8.2 Recommended Network Configuration', level=2)
    doc.add_paragraph('Production Environment:')
    prod_config = [
        'Deploy behind a reverse proxy (nginx, Cloudflare)',
        'Enable TLS 1.2+ only',
        'Configure firewall rules for database access',
        'Use private networking for backend services',
        'Enable DDoS protection at edge'
    ]
    for p in prod_config:
        doc.add_paragraph(p, style='List Bullet')

    doc.add_page_break()

    # 9. Compliance
    doc.add_heading('9. Compliance & Standards', level=1)

    doc.add_heading('9.1 Security Standards Alignment', level=2)
    doc.add_paragraph('DataMigrate AI aligns with the following security frameworks:')
    standards = [
        ('OWASP Top 10', 'SQL Injection, XSS, Auth, Sensitive Data, Security Misconfiguration'),
        ('SOC 2 Type II', 'Access Controls, Encryption, Audit Logging, Monitoring'),
        ('GDPR', 'Data Protection, Access Controls, Audit Trails'),
        ('ISO 27001', 'Information Security Management')
    ]
    add_table_with_header(doc, ['Standard', 'Coverage'], standards)

    doc.add_heading('9.2 Data Protection Measures', level=2)
    doc.add_paragraph('Customer Data:')
    data_protection = [
        'Organization-level data isolation',
        'No cross-tenant data access',
        'Encrypted credentials at rest',
        'Secure credential transmission'
    ]
    for d in data_protection:
        doc.add_paragraph(d, style='List Bullet')

    doc.add_page_break()

    # 10. Risk Assessment
    doc.add_heading('10. Risk Assessment', level=1)

    doc.add_heading('10.1 Risk Matrix', level=2)
    risks = [
        ('SQL Injection', 'Low', 'Critical', 'Pattern detection, parameterized queries', 'Low'),
        ('XSS Attack', 'Low', 'High', 'Input sanitization, CSP headers', 'Low'),
        ('Credential Theft', 'Low', 'Critical', 'AES-256 encryption, bcrypt hashing', 'Low'),
        ('Brute Force Attack', 'Medium', 'Medium', 'Rate limiting, account lockout', 'Low'),
        ('DDoS Attack', 'Medium', 'High', 'Rate limiting, CDN integration', 'Medium'),
        ('Session Hijacking', 'Low', 'High', 'Short-lived JWT, HTTPS only', 'Low'),
        ('Prompt Injection', 'Medium', 'Medium', 'Pattern detection, input validation', 'Low'),
        ('Data Breach', 'Low', 'Critical', 'Encryption, access controls, auditing', 'Low')
    ]
    add_table_with_header(doc, ['Risk', 'Likelihood', 'Impact', 'Mitigation', 'Residual'], risks)

    doc.add_heading('10.2 Security Testing', level=2)
    doc.add_paragraph('Recommended Testing:')
    testing = [
        'Penetration testing (annual)',
        'Vulnerability scanning (weekly)',
        'Code security review (per release)',
        'Dependency scanning (continuous)'
    ]
    for t in testing:
        doc.add_paragraph(t, style='List Bullet')

    doc.add_page_break()

    # 11. Incident Response
    doc.add_heading('11. Incident Response', level=1)

    doc.add_heading('11.1 Incident Classification', level=2)
    incidents = [
        ('P1 - Critical', 'Active breach, data exposure', 'Immediate', 'Data leak, successful attack'),
        ('P2 - High', 'Active attack attempt', '1 hour', 'Multiple blocked attacks'),
        ('P3 - Medium', 'Security anomaly', '4 hours', 'Unusual access patterns'),
        ('P4 - Low', 'Minor security event', '24 hours', 'Failed login attempts')
    ]
    add_table_with_header(doc, ['Level', 'Description', 'Response Time', 'Examples'], incidents)

    doc.add_heading('11.2 Response Procedures', level=2)

    doc.add_paragraph('Immediate Actions:', style='Heading 4')
    immediate = [
        'Identify and contain the threat',
        'Preserve audit logs and evidence',
        'Notify security team',
        'Assess impact scope'
    ]
    for idx, a in enumerate(immediate, 1):
        doc.add_paragraph(f'{idx}. {a}')

    doc.add_paragraph()
    doc.add_paragraph('Investigation:', style='Heading 4')
    investigation = [
        'Review audit logs',
        'Analyze attack patterns',
        'Identify affected resources',
        'Determine root cause'
    ]
    for idx, i in enumerate(investigation, 1):
        doc.add_paragraph(f'{idx}. {i}')

    doc.add_paragraph()
    doc.add_paragraph('Recovery:', style='Heading 4')
    recovery = [
        'Remove threat/vulnerability',
        'Restore affected systems',
        'Reset compromised credentials',
        'Notify affected users'
    ]
    for idx, r in enumerate(recovery, 1):
        doc.add_paragraph(f'{idx}. {r}')

    doc.add_page_break()

    # 12. Best Practices
    doc.add_heading('12. Security Best Practices for Clients', level=1)

    doc.add_heading('12.1 Account Security', level=2)
    account_practices = [
        'Use strong, unique passwords (12+ characters)',
        'Enable organization-level access controls',
        'Review user access regularly',
        'Remove inactive accounts promptly'
    ]
    for p in account_practices:
        doc.add_paragraph(p, style='List Bullet')

    doc.add_heading('12.2 API Key Management', level=2)
    api_key_practices = [
        'Rotate API keys regularly (90 days recommended)',
        'Use minimum necessary permissions',
        'Monitor API key usage',
        'Revoke unused keys immediately'
    ]
    for p in api_key_practices:
        doc.add_paragraph(p, style='List Bullet')

    doc.add_heading('12.3 Database Connection Security', level=2)
    db_practices = [
        'Use dedicated service accounts for migrations',
        'Grant minimum required permissions',
        'Use encrypted connections',
        'Rotate credentials after migration completion'
    ]
    for p in db_practices:
        doc.add_paragraph(p, style='List Bullet')

    doc.add_heading('12.4 Network Configuration', level=2)
    network_practices = [
        'Whitelist DataMigrate AI IP addresses',
        'Use VPN/private endpoints where possible',
        'Monitor connection attempts',
        'Log all database access'
    ]
    for p in network_practices:
        doc.add_paragraph(p, style='List Bullet')

    doc.add_page_break()

    # Appendix A
    doc.add_heading('Appendix A: Security Configuration Checklist', level=1)

    doc.add_heading('Production Deployment', level=2)
    checklist = [
        'Set strong JWT_SECRET (32+ characters)',
        'Set ENCRYPTION_KEY for AES-256 (32 bytes, base64-encoded)',
        'Configure ALLOWED_ORIGINS for your domains',
        'Set ENVIRONMENT=production',
        'Enable HTTPS/TLS only',
        'Configure firewall rules',
        'Set up monitoring and alerting',
        'Configure backup encryption',
        'Review rate limiting thresholds',
        'Test incident response procedures'
    ]
    for c in checklist:
        p = doc.add_paragraph()
        p.add_run('\u2610 ').font.size = Pt(12)
        p.add_run(c)

    doc.add_heading('Required Environment Variables', level=2)
    env_vars = [
        'JWT_SECRET - Secure secret for JWT signing (32+ characters)',
        'ENCRYPTION_KEY - Base64-encoded 32-byte key for AES-256',
        'ALLOWED_ORIGINS - Comma-separated allowed CORS origins',
        'ENVIRONMENT - production, staging, or development',
        'DATABASE_URL - PostgreSQL connection string with SSL'
    ]
    for e in env_vars:
        doc.add_paragraph(e, style='List Bullet')

    doc.add_page_break()

    # Appendix B
    doc.add_heading('Appendix B: Generating Encryption Keys', level=1)

    doc.add_paragraph('Using OpenSSL:')
    doc.add_paragraph('openssl rand -base64 32', style='Quote')

    doc.add_paragraph()
    doc.add_paragraph('This command generates a cryptographically secure 32-byte random key and encodes it in base64 format for use as the ENCRYPTION_KEY environment variable.')

    doc.add_page_break()

    # Document Control
    doc.add_heading('Document Control', level=1)
    doc_control = [
        ('1.0', datetime.now().strftime('%B %Y'), 'DataMigrate AI Team', 'Initial release')
    ]
    add_table_with_header(doc, ['Version', 'Date', 'Author', 'Changes'], doc_control)

    doc.add_paragraph()
    doc.add_paragraph()

    # Contact Information
    doc.add_heading('Contact Information', level=2)
    contacts = doc.add_paragraph()
    contacts.add_run('Security Team: ').bold = True
    contacts.add_run('security@datamigrate.ai\n')
    contacts.add_run('Support: ').bold = True
    contacts.add_run('support@datamigrate.ai\n')
    contacts.add_run('Emergency: ').bold = True
    contacts.add_run('+1-XXX-XXX-XXXX (24/7)')

    doc.add_paragraph()
    disclaimer = doc.add_paragraph()
    disclaimer.italic = True
    disclaimer.add_run('This document contains confidential security information. Distribution is restricted to authorized personnel and clients under NDA.')

    return doc

def main():
    # Create the security document
    print("Generating Security Documentation...")
    doc = create_security_document()

    # Save paths
    base_path = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    docs_path = os.path.join(base_path, 'docs')

    # Ensure docs directory exists
    os.makedirs(docs_path, exist_ok=True)

    # Save Word document
    docx_path = os.path.join(docs_path, 'SECURITY_DOCUMENTATION.docx')
    doc.save(docx_path)
    print(f"Word document saved: {docx_path}")

    print("\nDone! Security documentation generated successfully.")
    print(f"\nFiles created:")
    print(f"  - {docx_path}")
    print(f"\nNote: For PDF conversion, open the .docx file in Microsoft Word and export as PDF,")
    print(f"      or use a tool like LibreOffice, or an online converter.")

if __name__ == '__main__':
    main()
